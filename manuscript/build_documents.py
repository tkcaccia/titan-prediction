#!/usr/bin/env python3
"""Build the JTM manuscript, supplement and point-by-point response from results."""
from __future__ import annotations

import csv
import os
import shutil
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
REFERENCE = ROOT / "data" / "reference"
FIGURES = ROOT / "figures"
OUT = ROOT / "manuscript"
OUT.mkdir(exist_ok=True)
REPO = os.environ.get(
    "TITAN_REPOSITORY_URL", "https://github.com/tkcaccia/titan-prediction"
)
MODEL_REPO = os.environ.get(
    "TITAN_MODEL_REPOSITORY_URL", "https://github.com/tkcaccia/TITANPred"
)
MANUSCRIPT_TITLE = (
    "A systematic patient-level benchmark and reusable model resource for molecular "
    "and derived immune-feature prediction from pretrained TITAN whole-slide representations across "
    "32 TCGA cancers"
)


def rows(name):
    with (TABLES / name).open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def optional_rows(name):
    path = TABLES / name
    if not path.exists():
        return []
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def reference_rows(name):
    with (REFERENCE / name).open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fnum(x, digits=3):
    try:
        return f"{float(x):.{digits}f}"
    except (TypeError, ValueError):
        return "NA"


def ival(x):
    try:
        return int(float(x))
    except (TypeError, ValueError):
        return 0


def setup(doc, title=None):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 2.0
    for name, size, color in [("Title", 20, "17365D"), ("Heading 1", 15, "17365D"),
                              ("Heading 2", 12, "2F5597"), ("Heading 3", 10.5, "2F5597")]:
        st = styles[name]; st.font.name = "Arial"; st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color); st.font.bold = True
        st.paragraph_format.line_spacing = 2.0
    styles["Caption"].font.name = "Arial"; styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True
    header = sec.header.paragraphs[0]
    header.text = title or "TITAN patient-level prediction benchmark"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    # Journal of Translational Medicine requires continuous line numbering.
    sect_pr = sec._sectPr
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)
    return doc


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def add_table(doc, headers, data, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = str(h)
        for r in cell.paragraphs[0].runs: r.font.bold = True; r.font.size = Pt(8.5)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    for row in data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        row_pr.append(no_split)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths): row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_figure(doc, filename, caption, width=6.35):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    c = doc.add_paragraph(caption, style="Caption"); c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_landscape_section(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sec._sectPr.append(line_numbers)
    return sec


def add_portrait_section(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sec._sectPr.append(line_numbers)
    return sec


def add_labelled(doc, label, text):
    p = doc.add_paragraph(); r = p.add_run(label + " "); r.bold = True; p.add_run(text)


cohort = rows("patient_cohort_summary.csv")
cohort_rows = rows("patient_cohort_summary.csv")
cohort_patient = cohort_rows
continuous = rows("continuous_screen.csv")
binary = rows("binary_screen.csv")
site_c = rows("continuous_site_grouped_sensitivity.csv")
site_b = rows("binary_site_grouped_sensitivity.csv")
pool_c = rows("continuous_slide_pooling_sensitivity.csv")
pool_b = rows("binary_slide_pooling_sensitivity.csv")
pls2 = rows("pls1_vs_pls2_inflammation_summary.csv")
lit = rows("prior_mutation_literature_pair_summary.csv")
lit_accuracy = rows("prior_mutation_accuracy_comparison.csv")
mutation_novelty = rows("supported_mutation_novelty.csv")
mutation_literature_audit = rows("supported_mutation_literature_audit.csv")
site_retention_summary = rows("site_grouped_retention_summary.csv")
site_threshold_failures = rows("site_grouped_models_below_effect_threshold.csv")
site_fold_composition = rows("site_grouped_fold_composition_summary.csv")
site_fold_details = rows("site_grouped_outer_fold_composition.csv")
site_predictability = rows("tissue_source_site_predictability_summary.csv")
ridge_comparison = rows("pls_vs_ridge_representative_models.csv")
ridge_jobs = rows("pls_vs_ridge_representative_jobs.csv")
ridge_stratified = rows("pls_vs_ridge_representative_stratified_summary.csv")
external_locked_targets = reference_rows("external_validation_locked_targets.csv")
ridge_summary = rows("pls_vs_ridge_representative_summary.csv")
permutation_uncertainty = rows("permutation_monte_carlo_uncertainty.csv")
multiplicity_summary = rows("multiplicity_sensitivity_summary.csv")
multiplicity_by_endpoint = rows("multiplicity_sensitivity_by_endpoint.csv")
targeted_permutation = optional_rows("targeted_permutation_refinement.csv")
highlighted = optional_rows("highlighted_model_performance.csv")
binary_reliability = optional_rows("binary_class_reliability_summary.csv")
binary_class_sensitivity = optional_rows("binary_minimum_class_sensitivity.csv")
binary_limited = optional_rows("binary_limited_evidence_models.csv")
binary_fold_counts = optional_rows("binary_outer_fold_class_counts.csv")
binary_component_summary = optional_rows("binary_selected_component_distribution.csv")
binary_learning = optional_rows("binary_limited_evidence_learning_curve_summary.csv")
slide_coverage = optional_rows("slide_report_coverage_audit.csv")
slide_multiplicity = optional_rows("patient_slide_multiplicity_by_cancer.csv")
mutation_coverage = optional_rows("mutation_coverage_audit.csv")
mutation_eligibility = optional_rows("mutation_target_eligibility_audit.csv")
molecular_coverage = optional_rows("molecular_source_coverage_audit.csv")
participant_characteristics = optional_rows("participant_characteristics_by_cancer.csv")
coad_examples = optional_rows("coad_package_examples.csv")
coad_example_predictions = optional_rows("coad_package_example_predictions.csv")
endpoint_dictionary = rows("endpoint_dictionary.csv")
endpoint_definitions = rows("endpoint_definition_dictionary.csv")
endpoint_dictionary_summary = rows("endpoint_dictionary_summary.csv")
morphology_context = rows("morphology_context_examples.csv")
morphology_context_summary = rows("morphology_context_model_summary.csv")
literature_landscape = []
with (ROOT / "data" / "reference" / "prior_histology_model_landscape.csv").open(
    encoding="utf-8-sig", newline=""
) as f:
    literature_landscape = list(csv.DictReader(f))
with (ROOT / "data" / "reference" / "pan_cancer_benchmark_comparison.csv").open(
    encoding="utf-8-sig", newline=""
) as f:
    pan_cancer_comparison = list(csv.DictReader(f))
with (ROOT / "data" / "reference" / "tripod_ai_reporting_map.csv").open(
    encoding="utf-8-sig", newline=""
) as f:
    tripod_map = list(csv.DictReader(f))

def tier_counts(data):
    ans = defaultdict(Counter)
    for r in data: ans[r["family"]][r["tier"]] += 1
    return ans


ctier = tier_counts(continuous); btier = tier_counts(binary)
supported_c = [r for r in continuous if r["tier"] in ("A", "B")]
supported_b = [r for r in binary if r["tier"] in ("A", "B")]
binary_standard = [
    r for r in binary_reliability
    if r.get("model_evidence_tier") == "standard_internal_evidence"
]
binary_limited_reliability = [
    r for r in binary_reliability
    if r.get("model_evidence_tier") == "exploratory_limited_evidence"
]
binary_reliability_by_key = {
    (r.get("family"), r.get("tumor_type"), r.get("endpoint")): r
    for r in binary_reliability
}
low_prevalence_mutation_fusion = [
    r for r in binary_reliability
    if r.get("family") in ("driver_mutation", "fusion")
    and float(r.get("observed_tcga_prevalence") or 1) < 0.20
]
binary_sensitivity_20 = next(
    (r for r in binary_class_sensitivity if r.get("minimum_per_class") == "20"), {}
)
binary_sensitivity_50 = next(
    (r for r in binary_class_sensitivity if r.get("minimum_per_class") == "50"), {}
)
coad_supported_c = [r for r in supported_c if r["tumor_type"] == "COAD"]
coad_supported_b = [r for r in supported_b if r["tumor_type"] == "COAD"]
global_supported_c = [r for r in supported_c if float(r["q_value_global"]) < 0.05]
global_supported_b = [r for r in supported_b if float(r["q_value_global"]) < 0.05]
global_mutation_b = [r for r in global_supported_b if r["family"] == "driver_mutation"]
combined_multiplicity = next(
    (r for r in multiplicity_summary if r.get("outcome_type") == "combined"), {}
)
zero_999 = [
    r for r in permutation_uncertainty
    if r.get("permutations") == "999"
    and r.get("permutation_exceedances") == "0"
    and r.get("permutation_stopped_early") == "FALSE"
]
zero_999_upper = max(
    (float(r["p_mc_upper_95"]) for r in zero_999 if r.get("p_mc_upper_95")),
    default=float("nan"),
)
targeted_zero = [r for r in targeted_permutation if r.get("zero_exceedances") == "TRUE"]
targeted_p_min = min(
    (float(r["refined_p_9999"]) for r in targeted_permutation),
    default=float("nan"),
)
targeted_p_max = max(
    (float(r["refined_p_9999"]) for r in targeted_permutation),
    default=float("nan"),
)
targeted_p_summary = (
    f"all {fnum(targeted_p_min, 4)}"
    if targeted_permutation and abs(targeted_p_max - targeted_p_min) < 1e-12
    else f"{fnum(targeted_p_min, 4)}–{fnum(targeted_p_max, 4)}"
)
top_c = sorted(supported_c, key=lambda r: float(r["q2"]), reverse=True)
top_b = sorted(supported_b, key=lambda r: float(r["balanced_accuracy"]), reverse=True)
global_abstract_text = (
    f'{len(global_supported_c)} continuous and {len(global_mutation_b)} cancer–mutation '
    f'pairs passed the stricter across-cancer correction'
)

if highlighted:
    highlighted_continuous = [
        r for r in highlighted if r.get("outcome_type") == "continuous"
    ]
    abstract_continuous = max(
        highlighted_continuous,
        key=lambda r: float(r.get("q2") or "-inf"),
        default=None,
    )
    abstract_binary = next(
        (r for r in highlighted if r.get("outcome_type") == "binary"), None
    )
else:
    abstract_continuous = abstract_binary = None

if abstract_continuous and abstract_binary:
    abstract_metric_text = (
        f'{abstract_continuous["tumor_type"]}–{abstract_continuous["endpoint"]} '
        f'(five-repeat mean Q² {fnum(abstract_continuous.get("q2"))}, selection-conditioned 95% patient-resampling interval '
        f'{fnum(abstract_continuous.get("q2_ci_low"))}–{fnum(abstract_continuous.get("q2_ci_high"))}; RMSE '
        f'{fnum(abstract_continuous.get("rmse"))}, Spearman '
        f'{fnum(abstract_continuous.get("spearman"))}) and '
        f'{abstract_binary["tumor_type"]}–{abstract_binary["endpoint"]} '
        f'(five-repeat mean sensitivity {fnum(abstract_binary.get("sensitivity"))}, specificity '
        f'{fnum(abstract_binary.get("specificity"))}, balanced accuracy '
        f'{fnum(abstract_binary.get("balanced_accuracy"))}, selection-conditioned 95% patient-resampling interval '
        f'{fnum(abstract_binary.get("balanced_accuracy_ci_low"))}–{fnum(abstract_binary.get("balanced_accuracy_ci_high"))}; AUROC '
        f'{fnum(abstract_binary.get("auc"))}, PR-AUC {fnum(abstract_binary.get("pr_auc"))})'
    )
else:
    abstract_metric_text = (
        f'{top_c[0]["tumor_type"]}–{top_c[0]["endpoint"]} '
        f'(Q² {fnum(top_c[0]["q2"])}) and '
        f'{top_b[0]["tumor_type"]}–{top_b[0]["endpoint"]} '
        f'(balanced accuracy {fnum(top_b[0]["balanced_accuracy"])})'
    )

meta = rows("patient_cohort_summary.csv")
# The summary CSV contains one row per patient. Recompute cohort totals here.
n_patients = len(meta); n_slides = sum(ival(r.get("n_slides")) for r in meta)
n_multi = sum(ival(r.get("n_slides")) > 1 for r in meta)
n_cancers = len({r.get("tumor_type") for r in meta if r.get("tumor_type")})
n_missing_cancer = sum(not r.get("tumor_type") for r in meta)
n_exact_reports = sum(ival(r.get("exact_report_matched_slides")) for r in slide_coverage)
n_unmatched_reports = sum(ival(r.get("unmatched_slides")) for r in slide_coverage)
n_multiple_samples = sum(ival(r.get("patients_with_multiple_primary_sample_barcodes")) for r in slide_multiplicity)
n_mc3_profiled = sum(ival(r.get("matched_profiled_patients")) for r in mutation_coverage)
n_mc3_missing = sum(ival(r.get("embedding_patients_without_mc3_profile")) for r in mutation_coverage)
n_mutation_eligible = sum(r.get("eligibility") == "eligible" for r in mutation_eligibility)
n_mutation_ineligible = len(mutation_eligibility) - n_mutation_eligible
participant_overall = next(
    (r for r in participant_characteristics if r.get("tumor_type") == "Overall"),
    {},
)
source_coverage_summary = []
for source in sorted({r.get("source") for r in molecular_coverage if r.get("source")}):
    sr = [r for r in molecular_coverage if r.get("source") == source]
    source_coverage_summary.append((
        source,
        sum(ival(r.get("cohort_patients")) for r in sr),
        sum(ival(r.get("covered_patients")) for r in sr),
        sum(ival(r.get("missing_patients")) for r in sr),
        sum(ival(r.get("patients_with_multiple_primary_aliquots")) for r in sr),
        max([ival(r.get("maximum_primary_aliquots")) for r in sr] or [0]),
        sr[0].get("aggregation_rule", ""),
    ))

measurement_class_order = (
    "directly observed genomic alteration",
    "sequencing-derived continuous burden",
    "computationally inferred immune-cell fraction",
    "transcriptomic signature",
    "pathology-associated quantity",
    "composite genomic-context score",
)
endpoint_class_summary = []
for measurement_class in measurement_class_order:
    class_rows = [
        r for r in endpoint_dictionary
        if r.get("measurement_class") == measurement_class
    ]
    endpoint_class_summary.append((
        measurement_class,
        len(class_rows),
        len({
            (r.get("outcome_type"), r.get("family"), r.get("endpoint"))
            for r in class_rows
        }),
        sum(r.get("same_histology_modality") == "TRUE" for r in class_rows),
    ))

definition_group_summary = []
for group in sorted({r.get("definition_group") for r in endpoint_dictionary}):
    group_tests = [r for r in endpoint_dictionary if r.get("definition_group") == group]
    representative = group_tests[0]
    definition_group_summary.append((
        group,
        representative.get("measurement_class"),
        len(group_tests),
        len({
            (r.get("outcome_type"), r.get("family"), r.get("endpoint"))
            for r in group_tests
        }),
        representative.get("source_modality"),
        representative.get("direct_vs_inferred"),
        representative.get("derivation_algorithm"),
        representative.get("original_scale"),
        representative.get("equivalence_caveat"),
    ))
endpoint_class_result_text = "; ".join(
    f"{tests:,} {measurement_class}"
    for measurement_class, tests, _, _ in endpoint_class_summary
)

families = {
    "thorsson": "Thorsson immune/genomic-context",
    "aneuploidy": "aneuploidy",
    "fusion": "fusion",
    "microsatellite_instability": "MSI",
    "microsatellite_instability_sensitivity": "strict MSI sensitivity",
    "oncogenic_pathway": "oncogenic pathway",
    "driver_mutation": "cancer-gene mutation",
}

family_table = []
for fam in sorted(set(r["family"] for r in continuous + binary)):
    crows = [r for r in continuous if r["family"] == fam]
    brows = [r for r in binary if r["family"] == fam]
    if crows:
        counts = ctier[fam]
        family_table.append((families.get(fam, fam), "continuous",
                             len(crows), len({r["tumor_type"] for r in crows}),
                             counts["A"], counts["B"], counts["C"]))
    if brows:
        counts = btier[fam]
        family_table.append((families.get(fam, fam), "binary",
                             len(brows), len({r["tumor_type"] for r in brows}),
                             counts["A"], counts["B"], counts["C"]))


def examples(data, metric, n=8, include_family=False):
    ordered = sorted(data, key=lambda r: float(r[metric]), reverse=True)[:n]
    labels = {
        "driver_mutation": "mutation",
        "oncogenic_pathway": "pathway",
        "fusion": "fusion",
        "aneuploidy": "aneuploidy",
        "microsatellite_instability": "MSI",
        "microsatellite_instability_sensitivity": "strict MSI",
    }
    return "; ".join(
        f'{r["tumor_type"]}–{r["endpoint"]}'
        f'{" [" + labels.get(r["family"], r["family"]) + "]" if include_family else ""}'
        f' ({fnum(r[metric])})'
        for r in ordered
    )


def category(tier):
    return {"A": "within-cancer prespecified screening tier A",
            "B": "within-cancer prespecified screening tier B",
            "C": "screen-negative"}.get(tier, tier or "not assessed")


def cancer_result_text(code):
    crows = [r for r in supported_c if r["tumor_type"] == code]
    brows = [r for r in supported_b if r["tumor_type"] == code]
    parts = []
    if crows:
        parts.append(
            f'{len(crows)} continuous candidates: '
            + examples(crows, "q2", min(5, len(crows)))
        )
    else:
        parts.append("no continuous pair was screen-positive")
    if brows:
        parts.append(
            f'{len(brows)} binary candidates: '
            + examples(brows, "balanced_accuracy", min(6, len(brows)),
                       include_family=True)
        )
    else:
        parts.append("no binary pair was screen-positive")
    return "; ".join(parts)


def median(values):
    z = sorted(float(v) for v in values if v not in ("", "NA") and v is not None)
    if not z: return float("nan")
    m = len(z)//2
    return z[m] if len(z)%2 else (z[m-1]+z[m])/2


binary_learning_medians = {}
for fraction in ("0.5", "0.50", "0.75", "1", "1.0", "1.00"):
    matching = [r for r in binary_learning if r.get("training_fraction") == fraction]
    if matching:
        binary_learning_medians[float(fraction)] = {
            "balanced_accuracy": median(r.get("balanced_accuracy_mean") for r in matching),
            "auc": median(r.get("auc_mean") for r in matching),
            "pr_auc": median(r.get("pr_auc_mean") for r in matching),
        }
binary_learning_50 = binary_learning_medians.get(0.5, {})
binary_learning_75 = binary_learning_medians.get(0.75, {})
binary_learning_100 = binary_learning_medians.get(1.0, {})
binary_min_outer_positive = min(
    (ival(r.get("minimum_outer_test_positive")) for r in binary_reliability),
    default=0,
)
binary_min_outer_negative = min(
    (ival(r.get("minimum_outer_test_negative")) for r in binary_reliability),
    default=0,
)
limited_min_inner_training_positive = min(
    (ival(r.get("minimum_inner_training_positive")) for r in binary_limited_reliability),
    default=0,
)
limited_min_inner_training_negative = min(
    (ival(r.get("minimum_inner_training_negative")) for r in binary_limited_reliability),
    default=0,
)
limited_min_inner_validation_positive = min(
    (ival(r.get("minimum_inner_validation_positive")) for r in binary_limited_reliability),
    default=0,
)
limited_min_inner_validation_negative = min(
    (ival(r.get("minimum_inner_validation_negative")) for r in binary_limited_reliability),
    default=0,
)


site_deltas = [r["delta"] for r in site_c + site_b if r.get("feasible") == "TRUE" and r.get("delta")]
pool_deltas = [r["delta_first_minus_mean"] for r in pool_c + pool_b if r.get("delta_first_minus_mean")]
lit_counts = Counter(r["current_status"] for r in lit)
recovered_reports = [r for r in lit_accuracy if r["current_status"].startswith("recovered_tier_")]
prior_supported_mutations = [
    r for r in mutation_literature_audit
    if r["evidence_class"].startswith("previously supported")
]
prior_evaluated_not_supported = [
    r for r in mutation_literature_audit
    if r["evidence_class"].startswith("previously evaluated")
]
not_identified_mutations = [
    r for r in mutation_literature_audit
    if r["evidence_class"].startswith("not identified")
]
site_combined = next(
    (r for r in site_retention_summary if r.get("outcome_type") == "combined"), {}
)
site_failure_examples = sorted(
    site_threshold_failures,
    key=lambda r: float(r.get("site_grouped_metric") or 999),
)[:12]
site_named_keys = [
    ("binary", "READ", "APC"),
    ("binary", "COAD", "APC"),
    ("binary", "PCPG", "Cell Cycle"),
    ("binary", "LIHC", "CTNNB1"),
    ("continuous", "OV", "TCR Shannon"),
    ("continuous", "COAD", "SNV Neoantigens"),
]
site_named_failures = []
for outcome_type, cancer, endpoint in site_named_keys:
    match = next((
        r for r in site_threshold_failures
        if r.get("outcome_type") == outcome_type
        and r.get("tumor_type") == cancer and r.get("endpoint") == endpoint
    ), None)
    if match:
        site_named_failures.append(match)
site_predictability_eligible = [
    r for r in site_predictability
    if r.get("eligible") == "TRUE" and not r.get("error")
]
site_predictability_ineligible = [
    r for r in site_predictability if r.get("eligible") != "TRUE"
]
site_predictability_ranked = sorted(
    site_predictability_eligible,
    key=lambda r: float(r.get("normalized_macro_balanced_accuracy") or -1),
    reverse=True,
)
ridge_better = [r for r in ridge_comparison if r.get("selected_method") == "ridge"]
pls_better = [r for r in ridge_comparison if r.get("selected_method") == "PLS"]
baseline_uncertain = [
    r for r in ridge_comparison if r.get("selected_method") == "uncertain"
]
binary_baseline_summary = next(
    (r for r in ridge_summary if r.get("outcome_type") == "binary"), {}
)
continuous_baseline_summary = next(
    (r for r in ridge_summary if r.get("outcome_type") == "continuous"), {}
)
binary_secondary_ridge = [
    r for r in ridge_comparison
    if r.get("outcome_type") == "binary" and float(r.get("delta_secondary_ci_low")) > 0
]
binary_secondary_pls = [
    r for r in ridge_comparison
    if r.get("outcome_type") == "binary" and float(r.get("delta_secondary_ci_high")) < 0
]
binary_secondary_uncertain = [
    r for r in ridge_comparison
    if r.get("outcome_type") == "binary"
    and not (float(r.get("delta_secondary_ci_low")) > 0
             or float(r.get("delta_secondary_ci_high")) < 0)
]
continuous_secondary_ridge = [
    r for r in ridge_comparison
    if r.get("outcome_type") == "continuous"
    and float(r.get("delta_secondary_ci_low")) > 0
]
continuous_secondary_pls = [
    r for r in ridge_comparison
    if r.get("outcome_type") == "continuous"
    and float(r.get("delta_secondary_ci_high")) < 0
]
continuous_secondary_uncertain = [
    r for r in ridge_comparison
    if r.get("outcome_type") == "continuous"
    and not (float(r.get("delta_secondary_ci_low")) > 0
             or float(r.get("delta_secondary_ci_high")) < 0)
]
ridge_binary = [r for r in ridge_comparison if r.get("outcome_type") == "binary"]
ridge_continuous = [
    r for r in ridge_comparison if r.get("outcome_type") == "continuous"
]
ridge_screen_positive = [
    r for r in ridge_comparison if r.get("screen_tier") in {"A", "B"}
]
binary_primary_ridge_labels = " and ".join(
    f'{r.get("tumor_type")}–{r.get("endpoint")}'
    for r in ridge_better if r.get("outcome_type") == "binary"
)
continuous_primary_ridge_labels = ", ".join(
    f'{r.get("tumor_type")}–{r.get("endpoint")}'
    for r in ridge_better if r.get("outcome_type") == "continuous"
)


def prior_current_text(records):
    return "; ".join(
        f'{r["cancer"]}–{r["gene"]}: prior {r["prior_metric"]}, '
        f'current balanced accuracy {fnum(r["current_balanced_accuracy"])} '
        f'(q={fnum(r["current_q"])})'
        for r in records
    )


def site_failure_text(records):
    parts = []
    for r in records:
        metric = "balanced accuracy" if r.get("outcome_type") == "binary" else "Q²"
        parts.append(
            f'{r.get("tumor_type")}–{r.get("endpoint")} ({metric} '
            f'{fnum(r.get("original_metric"))} to {fnum(r.get("site_grouped_metric"))})'
        )
    return "; ".join(parts)


def selected_prior_records(keys):
    selected = []
    for cancer, gene, study in keys:
        match = next((
            r for r in lit_accuracy
            if r["cancer"] == cancer and r["gene"] == gene and r["study"] == study
        ), None)
        if match:
            selected.append(match)
    return selected


def benchmark_family_label(value):
    return {
        "driver_mutation": "mut.",
        "microsatellite_instability_sensitivity": "MSI",
        "oncogenic_pathway": "pathway",
        "thorsson": "immune",
    }.get(value, value)


def benchmark_type_label(value):
    return {"binary": "bin.", "continuous": "cont."}.get(value, value)


def benchmark_metric_label(value):
    return {
        "AUROC (threshold-independent)": "AUROC",
        "balanced accuracy (inner-CV thresholds for both)": "BA",
        "Spearman correlation": "Spearman",
        "Q2": "Q²",
    }.get(value, value)


def benchmark_interpretation(value):
    return "uncertain" if "difference uncertain" in value else value


previously_reported_not_supported = selected_prior_records([
    ("BRCA", "TP53", "Kather2020"),
    ("UCEC", "TP53", "Loeffler2022"),
    ("LIHC", "CTNNB1", "Kather2020"),
    ("PAAD", "KRAS", "Kather2020"),
])


doc = setup(Document(), "TITAN patient-level benchmark and model resource")
p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(MANUSCRIPT_TITLE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors = [
    ("Aamilah Ismail", "3,*"),
    ("Martin Ocharo", "1,2,*"),
    ("Moussa Kassim", "1,2"),
    ("Dalia Ahmed", "1"),
    ("Brendon Price", "5"),
    ("Dupe Ojo", "1"),
    ("Ekene Emmanuel Nweke", "4"),
    ("Silvano Piazza", "6,8"),
    ("Dinesh Gupta", "7"),
    ("Stefano Cacciatore", "1,2,†"),
]
for i, (name, markers) in enumerate(authors):
    if i:
        p.add_run(", ")
    p.add_run(name).bold = True
    marker_run = p.add_run(markers)
    marker_run.font.superscript = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("1 ").bold = True
p.add_run("Bioinformatics Unit, International Centre for Genetic Engineering and Biotechnology (ICGEB), Cape Town 7925, South Africa")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("2 ").bold = True
p.add_run("Department of Integrative Biomedical Sciences, Institute of Infectious Disease & Molecular Medicine (IDM), University of Cape Town, Cape Town 7925, South Africa")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("3 ").bold = True
p.add_run("Division of Engineering, New York University Abu Dhabi, Abu Dhabi, United Arab Emirates")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("4 ").bold = True
p.add_run("Department of Surgery, Faculty of Health Sciences, University of the Witwatersrand, Johannesburg, Gauteng, South Africa")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("5 ").bold = True
p.add_run("Division of Anatomical Pathology, University of Cape Town and National Health Laboratory Service, Observatory, Cape Town, South Africa")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("6 ").bold = True
p.add_run("Computational Biology Group, International Centre for Genetic Engineering and Biotechnology (ICGEB), Trieste, Italy")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("7 ").bold = True
p.add_run("Translational Bioinformatics Group, International Centre for Genetic Engineering and Biotechnology (ICGEB), Aruna Asaf Ali Marg, New Delhi 110067, India")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("8 ").bold = True
p.add_run("Bioinformatics Facility, Department of Cellular, Computational and Integrative Biology - CIBIO, University of Trento, Trento, Italy")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("* These authors contributed equally.").italic = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("† Corresponding author: Stefano Cacciatore (stefano.cacciatore@icgeb.org).").italic = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Author emails: aamilahismail@gmail.com; moussa.kassim@icgeb.org; dalia.ahmed@icgeb.org; dupe.ojo@icgeb.org; stefano.cacciatore@icgeb.org")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Research article — Molecular Pathology | Journal of Translational Medicine")

doc.add_heading("Abstract", level=1)
add_labelled(doc, "Background.", "Pan-cancer histology studies span molecular endpoints: Arslan et al. trained 12,093 models for 4,031 biomarkers across 32 TCGA cancers. We evaluated TITAN as a patient-level benchmark and model resource, without claiming the first pan-cancer screen or unprecedented mutation targets.")
add_labelled(doc, "Methods.", f"We mean-pooled {n_slides:,} diagnostic slide embeddings into {n_patients:,} patient vectors across {n_cancers} TCGA cancers; {n_multi:,} patients contributed multiple slides. We tested {len(continuous):,} continuous and {len(binary):,} binary cancer–endpoint pairs using patient-level nested PLS1 or PLS–LDA. Targets were classified as direct alteration calls, sequencing-derived burdens, inferred immune-cell fractions, transcriptomic signatures, pathology-derived quantities or composite scores. Models meeting the prespecified screening-statistic threshold underwent 999 complete-process permutations with conservative stopping, exact Monte Carlo intervals and Benjamini–Hochberg corrections locally, across cancers and atlas-wide. Eight prespecified models were extended to 9,999 permutations.")
add_labelled(doc, "Results.", f"Screening identified {len(supported_c):,} continuous and {len(supported_b):,} binary candidates in prespecified screening tiers A or B. A stricter ≥50-per-class sensitivity retained {len(binary_standard)}/{len(supported_b)} binary candidates; {len(binary_limited_reliability)} smaller-class models were exploratory and excluded from default inference. Highlighted repeated-CV results included {abstract_metric_text}. Prior studies supported {len(prior_supported_mutations)}/{len(mutation_literature_audit)} screen-positive cancer–gene pairs. {ival(combined_multiplicity.get('atlas_wide_pass'))}/{ival(combined_multiplicity.get('within_cancer_family_candidates'))} local candidates passed single BH correction across all {ival(combined_multiplicity.get('eligible_tests')):,} eligible tests. Tissue-source-site grouping reduced {ival(site_combined.get('below_threshold_models'))}/{ival(site_combined.get('screen_positive_models'))} candidates ({fnum(site_combined.get('below_threshold_percent'), 1)}%) below the original screening threshold. In a 47-target metadata-stratified benchmark selected without PLS performance, median ridge-minus-PLS differences were {fnum(binary_baseline_summary.get('median_delta_ridge_minus_pls'))} AUROC and {fnum(continuous_baseline_summary.get('median_delta_ridge_minus_pls'))} Q²; paired intervals favoured ridge for {sum(r.get('outcome_type') == 'binary' for r in ridge_better)}/{len(ridge_binary)} binary and {sum(r.get('outcome_type') == 'continuous' for r in ridge_better)}/{len(ridge_continuous)} continuous targets, and PLS for none.")
add_labelled(doc, "Conclusions.", "This systematic patient-level benchmark defines internally derived TCGA estimates under one fixed pretrained TITAN representation and constructs compact PLS reference models for future locked evaluation. PLS is the prespecified reference method, not an empirically optimal algorithm; the representative comparison more often favoured ridge. No independent performance evidence or basis for clinical application is presented.")
add_labelled(doc, "Trial registration.", "Not applicable.")
doc.add_paragraph("Keywords: computational pathology; whole-slide imaging; foundation model; mutation; inflammation; microsatellite instability; gene fusion; aneuploidy; partial least squares; linear discriminant analysis")

doc.add_heading("Background", level=1)
doc.add_paragraph("Routine haematoxylin-and-eosin sections reflect phenotypic consequences of tumour genotype and the immune microenvironment. Coudray and colleagues established mutation prediction from lung histology [1]. Subsequent work predicted microsatellite instability, including externally validated colorectal models [2,3]; extended mutation and multi-omic screening across TCGA cancers [4,5,8,11,13]; inferred RNA expression [6,12]; identified tissue-source-site bias [7]; detected gene fusions [9,10]; estimated homologous-recombination deficiency [14]; and characterised tumour-microenvironment phenotypes [15]. Collectively, these studies establish biological plausibility while also showing that performance depends on endpoint, disease, cohort and validation design.")
doc.add_paragraph("Breadth alone is not the gap. Fu et al. analysed 17,355 slides across 28 cancers [4], Kather et al. applied one workflow to more than 5,000 patients across 14 cancers [5], Saldanha et al. externally tested mutation models across seven matched TCGA and CPTAC cancers [11], and Arslan et al. trained 12,093 models for 4,031 genomic, transcriptomic, proteomic and clinical biomarkers in 8,890 TCGA patients across the same 32 cancers [13]. The present study consequently does not claim the first pan-cancer molecular screen, a larger endpoint catalogue, or general mutation-target novelty.")
doc.add_paragraph("The narrower unresolved gap is a reproducible benchmark that applies one fixed pretrained representation to every target; deterministically aggregates slides before outcome matching; performs nested validation with the patient as the unit; preserves tested-negative and sample-size-ineligible results; reports target-level tissue-source-site sensitivity; models selected derived immune and genomic-context features continuously; and constructs compact fitted predictors with research-use provenance metadata. These design and resource features—not unprecedented endpoint classes—define the contribution of this study.")
doc.add_paragraph("TITAN is a multimodal whole-slide foundation model pretrained on 335,645 slides with visual self-supervision and vision–language alignment [16]. The published Mass-340K pretraining corpus explicitly excluded TCGA and PANDA; TCGA was instead used for downstream evaluation of the pretrained model. The present work therefore has no reported TCGA pretraining overlap, but it is a secondary analysis of a cohort previously used to benchmark TITAN and is not independent external validation. We use its fixed 768-dimensional slide representation with PLS as one prespecified supervised linear reference; the study does not assume that PLS is algorithmically optimal [17,18].")
doc.add_paragraph("A practical feature of the fitted analysis is model portability. PLS, PLS–LDA and ridge predictors can all be represented by compact learned preprocessing and coefficient objects and applied without release of patient-level training embeddings or outcomes. Portability therefore motivates comparison within the linear-model family but is not a unique advantage of PLS. It supports external research testing while minimising distribution of patient-level data, but does not itself establish privacy, licensing compatibility or transportability.")
doc.add_paragraph("The primary question was cancer-specific: which individual mutations, derived inflammatory features, oncogenic pathways, MSI phenotypes, aneuploidy measures and fusions are predictable within each cancer? Molecular subtype association was not analysed. One-at-a-time PLS1 and PLS–LDA formed the prespecified reference screen because they preserve outcome-specific denominators and yield one independently usable model per endpoint. Joint PLS2 was evaluated secondarily only for coherent continuous inflammatory panels; its complete-case panel, shared latent structure and multiresponse estimand are not interchangeable with the endpoint-specific atlas.")
add_portrait_section(doc)

doc.add_heading("Methods", level=1)
doc.add_heading("Study design, slides and patient unit", level=2)
doc.add_paragraph(f"The published TITAN TCGA table contained {n_slides:,} eligible primary-tumour diagnostic slides (TCGA sample type 01 and –DX filename) from {n_patients:,} participants. The {n_multi:,} participants with multiple eligible slides contributed one patient vector obtained by feature-wise arithmetic mean before outcomes were joined. Patient—not slide—was the independent cross-validation unit. TITAN's TCGA-Slide-Reports.csv matched {n_exact_reports:,} selected slides exactly by filename; {n_unmatched_reports:,} slides lacked an exact report row. Report metadata were used only to audit identifiers, project/cancer provenance and resection-site annotations, and no report text entered a model. The two-character tissue-source-site code used for grouped validation was derived directly from the TCGA participant barcode. {n_missing_cancer:,} participants without a resolvable cancer label did not enter cancer-specific modelling. No participant had eligible slides from more than one primary sample barcode.")
doc.add_paragraph("All models were stratified by cancer type. No model learned pan-cancer differences, and no patient could occur in more than one validation fold. The intended use was discovery-stage prioritisation rather than diagnosis, treatment selection or replacement of molecular testing.")
doc.add_paragraph("All available eligible TCGA participants were used; no formal power calculation was performed. Minimum outcome-specific denominators were prespecified to support nested folds. Treatments were not modelled because the endpoints were contemporaneous molecular or derived immune/genomic features rather than prognosis or treatment response. Available demographic fields, subgroup sizes and endpoint-specific missingness did not provide a basis for representative subgroup performance or fairness evaluation; this is considered a limitation rather than evidence of equivalent performance across groups.")
doc.add_paragraph("Reporting was audited against TRIPOD+AI [28]. Participant characteristics were linked from the TCGA Clinical Data Resource [29] and summarized overall and by cancer. These descriptive fields were not supplied to prediction models, and no subgroup performance comparison was prespecified.")
add_figure(doc, "Figure1_patient_first_workflow.png", "Figure 1. Patient-first study design. Eligible diagnostic slides are mean-pooled before outcome matching; all model selection and evaluation occur at patient level within cancer type.")

doc.add_heading("Predictors and outcomes", level=2)
doc.add_paragraph("Predictors were the 768 fixed pretrained TITAN dimensions released for the TCGA slides [16]. The official gated artifact, TCGA_TITAN_features.pkl, was downloaded from https://huggingface.co/MahmoodLab/TITAN/blob/main/TCGA_TITAN_features.pkl after accepting the upstream terms and converted with the repository script to filename,titan_000,…,titan_767 CSV format; its source and converted-file SHA-256 hashes are retained in the provenance record. No slide pixels were reprocessed and TITAN weights were not fine-tuned in this study. Continuous outcomes comprised 39 immune/inflammatory measures and 11 genomic-context scores from Thorsson et al. [19], three aneuploidy burdens from Taylor et al. [20], log-transformed fusion burden from Gao et al. [21], and MANTIS/MSIsensor scores from Bonneville et al. and the cBioPortal TCGA PanCancer Atlas files [22,23]. Binary outcomes comprised qualifying protein-altering PASS mutations in tissue-specific consensus cancer genes from MC3 and Bailey et al. [24,25], ten oncogenic-pathway alteration indicators [26], genome doubling [20], any called fusion and eligible recurrent fusion pairs [21], and MSI-H definitions at MANTIS >0.4 and a strict >0.6 sensitivity threshold [22]. Protein-altering classes were missense, nonsense, nonstop, splice-site and translation-start mutations plus frameshift and in-frame insertions or deletions; their inclusion is recorded in a cancer-level variant-class audit. A mutation target denotes a qualifying alteration in a cancer gene; it does not assert that every allele is a functionally validated driver.")
doc.add_paragraph("The Thorsson Nonsilent Mutation Rate, Silent Mutation Rate, SNV Neoantigens, Indel Neoantigens and Number of Segments variables, and the Gao fusion burden, were analysed as log(1+x); all other continuous endpoints retained their source scale. Reported RMSE values therefore use these analysed units.")
doc.add_paragraph("The Thorsson umbrella does not denote a single immune assay. Its targets include CIBERSORT relative leukocyte fractions inferred by deconvolving bulk RNA sequencing [37], a leukocyte fraction inferred from DNA methylation, bulk-RNA gene-set signatures, BCR/TCR repertoire quantities reconstructed from RNA sequencing, in-silico neoantigen burdens, purity- and copy-number-derived genomic-context scores, and TIL Regional Fraction inferred from H&E by the Saltz deep-learning workflow [19,38]. The last target is a same-histology-modality concordance task rather than cross-modal molecular prediction. None of these inferred quantities is treated as equivalent to flow cytometry, immunohistochemistry, a directly counted immune-cell assay or a clinically certified biomarker. A 2,073-row endpoint dictionary records source modality, direct/inferred status, derivation, source scale, transformation, outcome-specific missingness, expected measurement error, biological interpretation and an assay-equivalence caveat for every modelled target; its 194 unique definitions are supplied separately.")
doc.add_paragraph(f"Molecular missingness was retained rather than converted to a negative label, and no outcome imputation was performed. Mutation status was defined only among {n_mc3_profiled:,} embedding patients matched to an MC3 primary-tumour profile; {n_mc3_missing:,} embedding patients without such a profile were excluded from mutation denominators. Within this profiled set, wild type meant no qualifying PASS protein-altering variant in the specified gene. Fusion-negative status was assigned only within the Gao study sample list. Multiple primary aliquots, if present, were collapsed at patient level by any alteration for binary mutation/pathway/fusion endpoints and by the prespecified mean or any-positive rule for continuous or binary instability endpoints. The source-level coverage table records covered patients, missing patients, aliquot multiplicity and aggregation for every cancer and source; no source contained multiple primary aliquots among matched embedding patients in this release.")
doc.add_paragraph("Continuous pairs required at least 50 non-missing patients. Binary screening eligibility required at least 20 positive and 20 negative patients. Eligibility was decided before modelling and every eligible test remained in its endpoint-family multiplicity denominator. Because this lower binary limit can produce sparse nested folds, we additionally repeated the eligibility and candidate-retention summary at 50 patients per class. Screen-positive models below that stricter threshold were retained in the complete atlas as exploratory/limited evidence, but were separated from standard-evidence models and excluded from the default inference interface.")

doc.add_heading("PLS1 regression and PLS–LDA classification", level=2)
doc.add_paragraph("For each cancer–endpoint pair, five outer folds estimated performance and five inner folds selected 1–10 PLS components. Scaling and component selection were learned from training patients only. All primary, permutation, repeated, sensitivity, PLS1–PLS2 and final-model fits used CPU rSVD with 10 oversampling vectors, two power iterations and explicit seeds. Continuous performance was out-of-fold Q², with RMSE and Spearman correlation secondary. Binary models supplied PLS latent scores to ridge-stabilised LDA; balanced accuracy was primary, with AUROC and precision–recall area under the curve (PR-AUC) calculated from continuous LDA scores in repeated validation.")
doc.add_paragraph("Q² was 1−Σ(y−ŷ)²/Σ(y−ȳ)² over outer-fold predictions, and RMSE was the square root of the mean squared prediction error. For binary outcomes, class 1 denoted altered or positive and class 0 denoted wild type or negative within the outcome-specific covered denominator. Sensitivity was the class-1 true-positive rate, specificity the class-0 true-negative rate, and balanced accuracy their arithmetic mean. AUROC used the continuous class-1-minus-class-0 LDA discriminant score. PR-AUC was non-interpolated average precision, so its no-skill reference equals outcome prevalence. Positive and negative predictive values (PPV and NPV) were calculated from held-out class calls at the observed outcome prevalence in the analysed TCGA cohort; they are cohort-specific descriptive estimates, not calibrated probabilities or transportable operating characteristics. Spearman correlation used observed and held-out continuous predictions.")
doc.add_paragraph("As a secondary algorithmic benchmark, targets were selected from all 2,073 eligible cancer–endpoint tests without using PLS performance. Continuous and binary sampling frames were divided into empirical sample-size terciles; binary endpoints were additionally divided into empirical terciles of minority-class fraction. Within each non-empty outcome-family × sample-size cell—and binary family × size × imbalance cell—the endpoint with the lowest salted SHA-256 rank was selected. The fixed rule yielded 12 continuous and 35 binary targets, including 14 screen-positive and 33 screen-negative endpoints. The complete sampling frame, salt, hashes and selected jobs are released. This is a representative metadata-stratified comparison, not an atlas-wide fit of ridge to every endpoint.")
doc.add_paragraph("Each selected target was fitted across five repeated nested partitions using PLS1 or PLS–LDA and ridge-penalised Gaussian or logistic regression. Within every repeat, both methods used identical patients, outer folds and inner folds. For continuous outcomes, both PLS component count and ridge penalty maximised pooled inner out-of-fold Q²; maximising Q² is equivalent to minimising pooled squared error because the validation outcome denominator is fixed. For binary outcomes, each method's operating threshold maximised balanced accuracy on its own inner out-of-fold continuous scores, without access to the outer test fold. Threshold-independent AUROC was the primary binary comparison metric and symmetrically thresholded balanced accuracy was secondary; continuous comparisons used Q² and Spearman correlation. For repeat r, d_r=M_r(ridge)−M_r(PLS) and Δ=(1/5)Σ_r d_r. Two thousand paired patient-resampling replicates sampled patients with replacement, retained both methods and all five matched predictions, and recalculated d_r and Δ; percentile 2.5 and 97.5 formed the interval. The interval does not repeat metadata-stratified target selection, generate new partitions or refit either algorithm. Ridge was chosen as a simple portable linear baseline; portability is not unique to PLS.")
doc.add_paragraph("No class under-sampling, over-sampling or synthetic augmentation was used. The prespecified atlas—not the secondary method benchmark—retained the fitted LDA class rule with observed training-fold class priors and returned a class plus an uncalibrated discriminant score, not a probability or clinical-risk threshold. The symmetric benchmark applied its nested threshold rule to both PLS–LDA and ridge and did not replace the primary atlas predictions or fitted models. Outcome tables were constructed independently of TITAN feature processing and joined only after patient-level aggregation.")
doc.add_paragraph("Nested performance was first checkpointed for the complete screen. Only models with Q²≥0.20 or chance-corrected balanced accuracy≥0.20 entered permutation testing. For every permutation, patient labels were reassigned and the complete modelling process was repeated: training-fold centering, five-fold inner selection of 1–10 components, outer-fold refitting and held-out prediction. Thus neither scaling parameters, component counts nor outer predictions were reused from the observed-label fit. A 99-permutation checkpoint was followed by refinement toward 999 permutations. During final refinement, evaluation stopped conservatively after 49 exceedances because raw p<0.05 was then impossible even if every remaining null statistic were less extreme; stopped endpoints were assigned p=1. Early-stopped jobs resumed from the exact number attempted, so no permutation indices were skipped. Regression extremeness used lower RMSE, whereas classification used higher balanced accuracy. Completed finite p values were (b+1)/(B+1), giving minimum resolution 0.001. Exact two-sided 95% Clopper–Pearson intervals for the underlying null exceedance probability quantify Monte Carlo uncertainty; with zero exceedances among 999 permutations this interval is 0–0.003686.")
doc.add_paragraph("Because each disease defines a separate prediction and intended-use population, Benjamini–Hochberg FDR was controlled separately for continuous and binary screens within cancer and prespecified endpoint family [27]. Two stricter sensitivities were retained: BH across cancers within the same outcome type and family, and a single BH correction across all eligible continuous and binary cancer–endpoint tests in the atlas. Eight models were locked before high-resolution results were generated: four leading continuous immune programmes and four leading binary claims spanning mutation, strict MSI and fusion endpoints. Their saved first 999 permutations were extended without early stopping to 9,999 complete-process permutations. This targeted analysis improves Monte Carlo resolution for leading claims but was not substituted into the prespecified 999-permutation FDR screen. The screening statistic is Q² for continuous outcomes and 2×balanced accuracy−1 for binary outcomes. Screen-positive candidates were assigned to prespecified screening tier A (primary q<0.05 and screening statistic≥0.40) or prespecified screening tier B (primary q<0.05 and screening statistic from 0.20 to <0.40). For binary outcomes these cut-offs correspond to balanced accuracy≥0.70 and 0.60 to <0.70, respectively; for continuous outcomes they correspond to Q²≥0.40 and 0.20 to <0.40. Screen-negative tested pairs and pairs ineligible by sample-size rules were reported separately. Models are ordered by the outcome-appropriate predictive metric, with repeated and site-grouped stability reported alongside; tied permutation q-values are not used for ranking. The tiers are prespecified prioritisation rules, not established clinical or statistical effect categories, and their numerical values are not interpreted as directly commensurate across outcome types.")

doc.add_heading("Robustness and saved models", level=2)
doc.add_paragraph("The initial nested-CV estimate used for candidate screening, permutation testing and tier assignment is termed the primary screening estimate. Only endpoints passing that initial atlas screen entered five additional independently seeded nested-CV partitions; their reported repeated-validation estimate is the arithmetic mean of the five repeat-specific metrics. The two estimates answer different questions and need not be numerically identical because they use different fold partitions and rSVD seeds. The primary screening estimate preserves the prespecified selection and multiplicity path, whereas the five-repeat estimate describes post-selection resampling stability; it does not replace or retroactively alter screen membership. For example, THYM–Th17 had primary screening Q²=0.638 and five-repeat mean Q²=0.594. Each repeat seed governed both fold construction and rSVD. Sensitivity, specificity, balanced accuracy and AUROC were calculated within each repeat for binary endpoints; Q², RMSE and Spearman correlation were calculated within each repeat for continuous endpoints. LDA scores from independently fitted repeats were never pooled onto an assumed common scale. For highlighted models, 1,000 patient-cluster bootstrap resamples retained all five existing held-out predictions for each sampled patient, recalculated the metric within repeat and then averaged the five repeat metrics; the 2.5th and 97.5th percentiles formed the reported 95% selection-conditioned patient-resampling interval for repeated out-of-fold predictions. This interval represents patient sampling variability conditional on the five fitted nested-CV prediction sets. It includes neither the initial screen and endpoint highlighting nor new partition generation, scaling or component reselection, model refitting inside bootstrap replicates, winner's-curse correction, or external cohort, site, scanner and population variation. Tissue-source-site-grouped validation assigned all patients sharing a two-character TCGA tissue-source-site code to the same outer fold. The same site constraint was passed to every inner component-selection split within the outer training set; thus neither outer performance estimation nor inner tuning divided a tissue-source site across folds. For every model and outer fold, we recorded test and training patients, sites, and—where applicable—positive and negative cases. A sensitivity analysis replaced each mean-pooled vector with the lexicographically first eligible diagnostic slide while preserving patients and seeds. Full-data research models were tuned by ten-fold CV and saved with feature order and checksum, training ranges, aggregation rule, endpoint transformation and output units, class coding and priors, decision rule, calibration status, class counts, exact software version and commit, computation backend, rSVD controls, external-validation status and intended-use metadata.")
doc.add_paragraph("For every screen-positive binary model, we recorded positive and negative counts in all 25 repeated outer test folds and reconstructed the exact five-fold inner component-selection partitions within every outer training set. We summarized the selected-component distribution across the resulting 25 fits. Repeat stability was quantified using pairwise Spearman correlations of patient-level standardized held-out scores and agreement of held-out class calls across repeats. For the limited-evidence subset, learning curves retained each repeat's original outer test fold while fitting on stratified 50%, 75% and 100% subsets of the outer training data; the five repeats were analysed separately. This design isolates training-size sensitivity without moving patients between training and evaluation. These analyses do not convert sparse internal validation into external evidence.")
doc.add_paragraph("To quantify whether TITAN itself retained submitting-site information, we treated tissue-source site as a multiclass outcome within each cancer. This dedicated confounding analysis retained sites with at least 10 patients, required at least two eligible sites per cancer, and used five independently seeded nested five-fold PLS–LDA validations with the same component range and rSVD controls. The primary metric was multiclass macro balanced accuracy (mean per-site recall); chance was 1/k for k analysed sites, and a chance-normalised value, (balanced accuracy − 1/k)/(1 − 1/k), was calculated for cross-cancer description. Patients from rare sites excluded from this site-classification analysis remained in all eligible molecular-endpoint analyses. This analysis tests whether the representation encodes TCGA submitting-site signatures; it does not identify their scanner, laboratory, staining or population causes.")
doc.add_paragraph("For qualitative morphological context, we selected five representative models spanning an H&E-derived TIL quantity, an RNA-derived CIBERSORT fraction, an RNA-expression signature, a copy-number burden and a sequence-supported mutation. Within each model, high and low anchors were chosen from concordant extremes of mean repeated out-of-fold predictions. The closest within-cancer patient was then retrieved by cosine similarity of patient-level mean TITAN representations. For display only, a report-covered slide closest to each patient mean was selected deterministically, and prespecified morphological terms were extracted from TITAN-generated TCGA-Slide-Reports text. Neither report text nor example selection entered model fitting. Only global 768-dimensional pooled representations were available; no patch embeddings, attention maps or slide pixels were retained. The analysis is therefore descriptive nearest-neighbour context, not patch-level relevance, causal attribution or blinded pathologist review.")
doc.add_paragraph("TITANPred retains the binary LDA score but does not convert it to a probability. In the research-software demonstration, its position is labelled 'TCGA out-of-fold score rank (not probability)' rather than percentile or risk. Binary output displays total training n, class counts and prevalence; balanced accuracy, AUROC and PR-AUC; TCGA-prevalence PPV and NPV; fold minima, component-selection variability and repeat stability; and site sensitivity. The 17 screen-positive models with fewer than 50 patients in either class are marked exploratory/limited evidence, omitted from default inference and available only through explicit opt-in with a warning. Probability calibration was not added post hoc because it would require an independently evaluated or fully nested calibration procedure, which is not available in this TCGA benchmark.")
doc.add_paragraph("The fitted PLS and PLS–LDA objects contain the learned transformations and coefficients needed for research inference but no patient-level training rows. They are distributed as prespecified reference models, not as models shown to outperform ridge. The access-controlled inference interface validates the 768-feature order and applies the prespecified patient-level slide aggregation. The secondary PLS1–PLS2 inflammatory comparison, including matched folds and cancer-bootstrap intervals, is described in the Supplementary Methods.")

doc.add_heading("Software, transparency and validation status", level=2)
doc.add_paragraph("Analyses used R 4.6.0 and fastPLS 0.99.20 (Git commit dcf45cc). MC3 objects were loaded with TCGAmutations 0.4.0 from the GitHub source pinned in the software manifest. Version-controlled code, source manifests, target catalogues, out-of-fold prediction summaries, figures and the model registry are organized in the public companion analysis repository [30]. A separate access-controlled GPL-3 R package contains the fitted models, their SHA-256 registry, reference distributions, inference interface and research-software output template [32]. That repository is private and the manuscript does not claim public availability of the fitted artifacts. The artifacts contain learned parameters and training-range summaries but no patient-level training rows. No independent cohort with compatible TITAN features and the required labels was included; every performance estimate is an internally derived TCGA estimate.")
doc.add_heading("Prospectively locked protocol for future independent evaluation", level=2)
doc.add_paragraph("After the internal TCGA screen, but before accessing any external features, outcomes or predictions, we fixed a three-model UCEC subset for future independent evaluation: TP53 mutation, genome doubling and continuous aneuploidy score. The exact fitted-object SHA-256 hashes are c39ba2caeefc148e67412647f3e1835dbcd719088a4d379b53068482ebb2ede0, 8f9fd73578189bab5d5c88bb8005415a07b93a662f2b0adb893f0b04e1cc3604 and 0c520d8a9d57dd5af109416c5b82ffb0f0862e3505cd9888965fdeee4aa60f3b, respectively. Their internal performance is selection-conditioned; locking them now does not convert TCGA estimates into external evidence.")
doc.add_paragraph("A future evaluation must extract every eligible primary-tumour diagnostic slide with the documented CONCH v1.5–TITAN pipeline, apply the same deterministic patient mean pooling, preserve the 768-feature order and use the hashed artifacts without refitting, recalibration, class-prior changes or threshold adjustment. Outcomes must be constructed independently of image predictions using definitions compatible with the TCGA targets; an irreproducible endpoint must be reported as non-evaluable rather than replaced with an externally optimised surrogate. Binary primary performance is balanced accuracy, with AUROC, sensitivity and specificity secondary; continuous primary performance is Q², with RMSE, Spearman correlation, calibration intercept and calibration slope secondary. Patient/slide flow, molecular missingness, exclusions, denominators, uncertainty and all three outcomes—including failures—must be reported. The executable specification and machine-readable lock file are provided in the companion repository.")
doc.add_paragraph("OpenAI Codex was used for analysis-code refactoring, document generation and language editing. All statistical choices, source-data mappings, numerical outputs and manuscript interpretations were reviewed by the human authors, who retain responsibility for the work.")
doc.add_paragraph("Supplementary methods, Tables S1–S14, Figures S1–S4 and the machine-readable-file inventory are provided in Additional file 1. The two post hoc illustrative COAD research-software outputs are supplied separately as Additional files 2 and 3.")

doc.add_heading("Results", level=1)
doc.add_heading("Cohort and analysis coverage", level=2)
doc.add_paragraph(f"The embedding cohort contained {n_patients:,} patients represented by {n_slides:,} slides; {n_multi:,} patients ({100*n_multi/n_patients:.1f}%) had more than one eligible diagnostic slide. The maximum was 30 slides for one patient, and no patient contributed slides from multiple primary sample barcodes. Exact slide-report coverage was {n_exact_reports:,}/{n_slides:,} ({100*n_exact_reports/n_slides:.1f}%). Cancer labels were available for {n_patients-n_missing_cancer:,} patients across {n_cancers} cancers. The benchmark evaluated {len(continuous):,} continuous and {len(binary):,} binary cancer–endpoint pairs (Table 1). Of {len(mutation_eligibility):,} prespecified cancer–gene mutation pairs with a profiled denominator, {n_mutation_eligible:,} met the 20-positive/20-negative rule and {n_mutation_ineligible:,} were reported as ineligible rather than tested-negative. Molecular coverage varied by source and cancer; missing cases were excluded outcome by outcome.")
if participant_overall:
    known_gender = ival(participant_overall.get("female")) + ival(participant_overall.get("male"))
    doc.add_paragraph(
        f'The TCGA Clinical Data Resource [29] matched {ival(participant_overall.get("cdr_matched")):,}/{n_patients:,} participants. '
        f'Age was available for {ival(participant_overall.get("age_available")):,} participants (median {fnum(participant_overall.get("age_median"), 1)} years, IQR {fnum(participant_overall.get("age_q1"), 1)}–{fnum(participant_overall.get("age_q3"), 1)}); '
        f'{ival(participant_overall.get("female")):,}/{known_gender:,} with recorded gender were female. '
        f'Race was recorded for {ival(participant_overall.get("race_available")):,} and broad stage I–IV for {ival(participant_overall.get("stage_available")):,}. '
        "These fields describe cohort composition only; subgroup predictive performance was not evaluated."
    )
doc.add_paragraph("Table 1. Eligible cancer–endpoint models and within-cancer screening categories.", style="Caption")
add_table(doc, ["Family", "Type", "Tests", "Cancers", "Screening tier A", "Screening tier B", "Screen-negative"], family_table,
          [4.8, 2.0, 1.5, 1.5, 1.5, 1.5, 1.5])

doc.add_heading("Endpoint provenance and assay equivalence", level=2)
doc.add_paragraph(
    f"The target-level audit classified the {len(endpoint_dictionary):,} eligible tests as "
    f"{endpoint_class_result_text}. These categories describe the label-generation modality, "
    "not the modality of the TITAN predictor. Only the 13 cancer-specific TIL Regional Fraction "
    "tests used an outcome derived from H&E itself. Thus, apparent predictability of CIBERSORT "
    "fractions, leukocyte fraction or RNA signatures is agreement with another computational "
    "phenotype and cannot be interpreted as recovery of a directly measured immune-cell count. "
    "The complete target-level dictionary and the 194-definition dictionary are Supplementary "
    "Table S13 machine-readable companions."
)

doc.add_heading("Continuous derived immune, genomic-context and instability phenotypes", level=2)
doc.add_paragraph(f"Among continuous targets, {len([r for r in supported_c if r['tier']=='A'])} were assigned to prespecified screening tier A and {len([r for r in supported_c if r['tier']=='B'])} to tier B. The largest primary screening Q² values were {examples(top_c, 'q2', 10)}. These initial nested-CV values determined screening and tier assignment; five-repeat means are reported separately for stability. These results were cancer specific: the same endpoint could be predictable in one tumour type and screen-negative in another.")
doc.add_paragraph("As an explicit example of the two reporting stages, THYM–Th17 had primary screening Q²=0.638 and a five-repeat mean repeated-validation Q²=0.594. This difference reflects independently seeded nested-CV partitions and is not a discrepancy or a change to the primary screen result.")
doc.add_paragraph(f"Of {len(supported_c)} within-cancer screen-positive continuous pairs, {len(global_supported_c)} met q<0.05 under the stricter across-cancer family sensitivity after the 999-permutation refinement. " + ("Continuous findings should therefore be interpreted as cancer-specific candidates rather than globally FDR-supported pan-cancer discoveries. Because the minimum completed-test p-value was 0.001, the global sensitivity was resolution-limited for large endpoint families; failure to pass it is not evidence of absence of biological signal." if not global_supported_c else "Pairs passing this sensitivity are identified explicitly in the figure and machine-readable table; the remainder are cancer-specific candidates. The 0.001 minimum completed-test p-value should be considered when interpreting large global families."))
add_figure(doc, "Figure2_continuous_atlas.png", "Figure 2. Eighteen strongest within-cancer screen-positive continuous cancer–endpoint results, ranked by primary screening Q². The reduced set and larger typography improve readability; complete results and five-repeat estimates are in Supplementary Table S7 and machine-readable files. All Q² values use genuine patient-level outer-fold predictions.")

doc.add_heading("Binary molecular phenotypes", level=2)
doc.add_paragraph(f"The binary screen yielded {len([r for r in supported_b if r['tier']=='A'])} candidates in prespecified screening tier A and {len([r for r in supported_b if r['tier']=='B'])} in tier B. The largest primary screening balanced accuracies were {examples(top_b, 'balanced_accuracy', 12, include_family=True)}. These initial nested-CV values determined screening and tier assignment; five-repeat means are reported separately for stability. The complete table distinguishes tested-negative pairs from outcomes that failed prevalence or sample-size eligibility.")
doc.add_paragraph(f"Of {len(supported_b)} within-cancer screen-positive binary pairs, {len(global_supported_b)} also met the stricter across-cancer family q<0.05, including {len(global_mutation_b)} cancer–gene mutation pairs. " + ("Mutation results are consequently cancer-specific discovery candidates, even where their within-cancer discrimination was strong. The 0.001 permutation resolution also limits the attainable across-cancer q-value for a mutation family spanning many cancer–gene tests." if not global_mutation_b else "Mutation pairs passing this stricter sensitivity are distinguished from those supported only within cancer; the 0.001 permutation resolution remains relevant for the full mutation family."))
doc.add_heading("Binary class-size sensitivity and development stability", level=2)
doc.add_paragraph(
    f"Of {ival(binary_sensitivity_20.get('eligible_binary_targets'))} binary pairs eligible at 20 patients per class, "
    f"{ival(binary_sensitivity_50.get('eligible_binary_targets'))} ({fnum(binary_sensitivity_50.get('eligible_target_retention_percent'), 1)}%) "
    f"also met a 50-per-class rule. Among the {len(supported_b)} screen-positive models, "
    f"{len(binary_standard)} ({100 * len(binary_standard) / len(supported_b):.1f}%) met that stricter standard. "
    f"The other {len(binary_limited_reliability)} were designated exploratory/limited evidence; none of the highlighted "
    "binary models was in this group. The complete atlas retains these results for transparency, whereas default "
    f"TITANPred inference now includes {len(binary_standard)} binary and {len(supported_c)} continuous models and omits "
    "the limited subset unless explicitly requested."
)
doc.add_paragraph(
    f"Across all {len(supported_b)} screen-positive binary models, the smallest repeated outer test fold contained "
    f"{binary_min_outer_positive} positive and {binary_min_outer_negative} negative patients. Within the limited subset, "
    f"exact reconstructed inner folds contained as few as {limited_min_inner_training_positive} positive and "
    f"{limited_min_inner_training_negative} negative training patients and {limited_min_inner_validation_positive} positive and "
    f"{limited_min_inner_validation_negative} negative validation patients. Selected components were therefore reported "
    f"for every one of the 25 repeated outer fits: {sum(ival(r.get('selected_components_minimum')) == 1 for r in binary_limited_reliability)}/{len(binary_limited_reliability)} limited models selected one component in at least one fit, "
    f"{sum(ival(r.get('selected_components_maximum')) == 10 for r in binary_limited_reliability)}/{len(binary_limited_reliability)} reached the ten-component ceiling, and "
    f"{sum(ival(r.get('selected_components_minimum')) == 1 and ival(r.get('selected_components_maximum')) == 10 for r in binary_limited_reliability)} spanned both extremes."
)
doc.add_paragraph(
    f"For the {len(binary_limited_reliability)} limited models, median learning-curve balanced accuracy was "
    f"{fnum(binary_learning_50.get('balanced_accuracy'))}, {fnum(binary_learning_75.get('balanced_accuracy'))} and "
    f"{fnum(binary_learning_100.get('balanced_accuracy'))} at 50%, 75% and 100% of the outer-training data; median AUROC was "
    f"{fnum(binary_learning_50.get('auc'))}, {fnum(binary_learning_75.get('auc'))} and {fnum(binary_learning_100.get('auc'))}, "
    f"and median PR-AUC was {fnum(binary_learning_50.get('pr_auc'))}, {fnum(binary_learning_75.get('pr_auc'))} and "
    f"{fnum(binary_learning_100.get('pr_auc'))}. Median repeat score correlation was "
    f"{fnum(median(r.get('repeat_score_spearman_mean') for r in binary_limited_reliability))} in the limited subset versus "
    f"{fnum(median(r.get('repeat_score_spearman_mean') for r in binary_standard))} in the ≥50-per-class subset. "
    "Class-call agreement was also reported, but can appear high under marked imbalance and is not interpreted alone. "
    "Endpoint-level PR-AUC, TCGA-prevalence PPV/NPV, fold counts, component distributions, learning curves and stability "
    "are provided in Supplementary Table S10g and Figure S3."
)
doc.add_heading("Permutation resolution and atlas-wide multiplicity sensitivity", level=2)
doc.add_paragraph(
    f"The primary within-cancer/family procedure identified {ival(combined_multiplicity.get('within_cancer_family_candidates'))} candidates. "
    f"Of these, {ival(combined_multiplicity.get('across_cancer_family_pass'))} passed BH within outcome type and endpoint family across cancers, "
    f"{ival(combined_multiplicity.get('outcome_wide_pass'))} passed BH across all eligible tests within their outcome type, and "
    f"{ival(combined_multiplicity.get('atlas_wide_pass'))} passed one BH correction across all {ival(combined_multiplicity.get('eligible_tests')):,} eligible continuous and binary atlas tests. "
    "The three local candidates not retaining atlas-wide q<0.05 were READ genome doubling, OV–TP53 mutation and SKCM–BRAF mutation. "
    "The candidate count should therefore be read as a collection of locally controlled cancer-specific questions, with the atlas-wide result as a stricter sensitivity—not as if the primary procedure were one global 5% FDR analysis."
)
doc.add_paragraph(
    f"Among completed 999-permutation tests, {len(zero_999)} had zero null statistics at least as extreme as observed. Their finite empirical p-value was 0.001, but the exact two-sided 95% Monte Carlo interval for the underlying exceedance probability was 0–{fnum(zero_999_upper, 6)}. "
    f"The prespecified high-resolution subset extended {len(targeted_permutation)} leading models to 9,999 full nested permutations; {len(targeted_zero)} had zero exceedances and refined p-values were {targeted_p_summary}. "
    "The refined values are a targeted precision sensitivity and were not inserted into the primary FDR calculations. Because many primary p- and q-values remain tied at their attainable minimum, figures and tables are ordered by the outcome-appropriate predictive metric, with repeated and site-grouped stability reported alongside; q-values are not used for ranking."
)
add_figure(doc, "Figure3_binary_atlas.png", "Figure 3. Eighteen strongest within-cancer screen-positive binary molecular predictions among the 87 models with at least 50 patients per class, ranked by primary screening balanced accuracy. Uniform point size and shortened labels reduce visual load; colour indicates endpoint family and symbols indicate across-cancer correction status. Complete BA, AUROC and PR-AUC results are in Supplementary Tables S6a, S7 and S10g. The 17 limited-evidence models appear in Figure S3.")
doc.add_paragraph(
    f"Precision–recall performance was reported for every screen-positive binary endpoint. Across all {len(binary_reliability)} models, "
    f"median five-repeat PR-AUC was {fnum(median(r.get('repeated_pr_auc_mean') for r in binary_reliability))}; each value should be read against its no-skill reference, the observed TCGA prevalence. "
    f"The {len(low_prevalence_mutation_fusion)} mutation or fusion endpoints with prevalence below 0.20 had median prevalence "
    f"{fnum(median(r.get('observed_tcga_prevalence') for r in low_prevalence_mutation_fusion))}, median PR-AUC "
    f"{fnum(median(r.get('repeated_pr_auc_mean') for r in low_prevalence_mutation_fusion))} (range "
    f"{fnum(min(float(r.get('repeated_pr_auc_mean')) for r in low_prevalence_mutation_fusion))}–"
    f"{fnum(max(float(r.get('repeated_pr_auc_mean')) for r in low_prevalence_mutation_fusion))}), despite median AUROC "
    f"{fnum(median(r.get('repeated_auc_mean') for r in low_prevalence_mutation_fusion))}. This divergence illustrates why AUROC or balanced accuracy alone can overstate positive-class retrieval under low prevalence. Endpoint-level PR-AUC, prevalence, repeat variability and class counts are reported in Supplementary Tables S6a, S7 and S10g and the machine-readable registry."
)
add_figure(doc, "Figure4_prediction_examples.png", "Figure 4. Examples comparing observed outcomes with held-out predictions. Panel A shows the strongest primary screening estimate from the initial patient-level nested-CV continuous screen (TGCT TGF-beta Response); each point is one patient's observed value and genuine held-out prediction. Panel B shows the binary model with the largest five-repeat mean balanced accuracy. Because independently fitted LDA scores need not share a raw scale, held-out scores are standardized within repeat before patient-level averaging for display; repeat-specific AUROCs use untransformed scores. These deliberately strong examples are illustrative, selection-conditioned internal TCGA results rather than external validation or calibration.")
doc.add_paragraph("Complete highlighted-model performance, uncertainty, class-size metrics and tissue-source-site-grouped estimates are reported in Supplementary Table S6a and the machine-readable result files. The main text retains outcome-level summaries and illustrative figures.")
add_figure(doc, "Figure5_supported_counts.png", "Figure 5. Breadth and family of screen-positive targets in the 20 cancers with the largest candidate counts. Restricting the display improves label size; complete cancer-level counts are in Supplementary Table S1. Absence of a target is distinct from a tested screen-negative result.")

doc.add_heading("Metadata-stratified linear-model comparison", level=2)
doc.add_paragraph(
    f"The deterministic representative benchmark contained {len(ridge_binary)} binary and "
    f"{len(ridge_continuous)} continuous targets selected from all eligible tests without "
    f"reference to PLS performance; {len(ridge_screen_positive)}/{len(ridge_comparison)} happened to be screen-positive. "
    f"For binary targets, the median ridge-minus-PLS AUROC difference was "
    f"{fnum(binary_baseline_summary.get('median_delta_ridge_minus_pls'))} "
    f"(IQR {fnum(binary_baseline_summary.get('q1_delta'))} to {fnum(binary_baseline_summary.get('q3_delta'))}); "
    f"ridge had a selection-conditioned paired patient-resampling interval above zero for {sum(r.get('outcome_type') == 'binary' for r in ridge_better)} "
    f"{'model' if sum(r.get('outcome_type') == 'binary' for r in ridge_better) == 1 else 'models'}, "
    f"PLS for {sum(r.get('outcome_type') == 'binary' for r in pls_better)}, and "
    f"{sum(r.get('outcome_type') == 'binary' for r in baseline_uncertain)} AUROC differences were uncertain. "
    f"With the same inner-CV balanced-accuracy threshold rule applied to both methods, the median ridge-minus-PLS balanced-accuracy difference was "
    f"{fnum(binary_baseline_summary.get('median_secondary_delta_ridge_minus_pls'))} "
    f"(IQR {fnum(binary_baseline_summary.get('q1_secondary_delta'))} to {fnum(binary_baseline_summary.get('q3_secondary_delta'))}); "
    f"ridge had a selection-conditioned paired patient-resampling interval above zero for {len(binary_secondary_ridge)} "
    f"{'model' if len(binary_secondary_ridge) == 1 else 'models'}, PLS for {len(binary_secondary_pls)}, and "
    f"{len(binary_secondary_uncertain)} were uncertain. "
    f"For continuous targets, the median ridge-minus-PLS Q² difference was "
    f"{fnum(continuous_baseline_summary.get('median_delta_ridge_minus_pls'))} "
    f"(IQR {fnum(continuous_baseline_summary.get('q1_delta'))} to {fnum(continuous_baseline_summary.get('q3_delta'))}); "
    f"ridge was favoured for {sum(r.get('outcome_type') == 'continuous' for r in ridge_better)}, "
    f"PLS for {sum(r.get('outcome_type') == 'continuous' for r in pls_better)}, and "
    f"{sum(r.get('outcome_type') == 'continuous' for r in baseline_uncertain)} were uncertain. "
    f"For continuous Spearman correlation, ridge was favoured for {len(continuous_secondary_ridge)}, "
    f"PLS for {len(continuous_secondary_pls)} and {len(continuous_secondary_uncertain)} were uncertain. "
    "Thus the benchmark did not establish PLS superiority and more often favoured ridge, especially for continuous Q². These are representative paired comparisons rather than multiplicity-controlled algorithm discoveries or an atlas-wide ridge screen."
)

doc.add_heading("Tissue-source-site sensitivity and submitting-site prediction", level=2)
doc.add_paragraph(
    f"Tissue-source-site-grouped internal validation was feasible for {sum(r.get('feasible')=='TRUE' for r in site_c)+sum(r.get('feasible')=='TRUE' for r in site_b)} screen-positive models. "
    f"Although the median grouped-minus-random performance change was only {fnum(median(site_deltas))}, "
    f"{ival(site_combined.get('below_threshold_models'))}/{ival(site_combined.get('screen_positive_models'))} "
    f"models ({fnum(site_combined.get('below_threshold_percent'), 1)}%) fell below the original prespecified screening threshold: "
    f"58/219 continuous models and 25/104 binary models. Prominent attenuations were "
    f"{site_failure_text(site_named_failures)}. The site-grouped folds contain fewer and less evenly distributed "
    "training and test patients, so the decline cannot be attributed solely to site artefact; nevertheless, the near-chance "
    "colorectal APC results are not adequately summarised by the median and are labelled site-sensitive in the main model table, registry and research-software output. "
)
doc.add_paragraph(
    f"Across {len(site_fold_details):,} outer folds, test partitions contained 1–445 patients and 1–9 tissue-source sites; binary folds contained 0–130 positive and 0–269 negative patients. These imbalances are intrinsic to keeping complete sites together and are reported model by model and fold by fold. The deterministic fold audit confirmed that no site crossed an outer fold and that the same separation was maintained in every inner component-selection split."
)
doc.add_paragraph(
    f"The dedicated within-cancer site-classification analysis was evaluable in {len(site_predictability_eligible)}/32 cancers and included {sum(ival(r.get('analysed_patients')) for r in site_predictability_eligible):,} patients from sites with at least 10 patients. All 27 observed macro balanced accuracies exceeded their cancer-specific 1/k chance reference; the median was {fnum(median([r.get('macro_balanced_accuracy_mean') for r in site_predictability_eligible]))} (range {fnum(min(float(r.get('macro_balanced_accuracy_mean')) for r in site_predictability_eligible))}–{fnum(max(float(r.get('macro_balanced_accuracy_mean')) for r in site_predictability_eligible))}). The strongest chance-normalised site predictability occurred in "
    + "; ".join(
        f"{r.get('tumor_type')} (macro balanced accuracy {fnum(r.get('macro_balanced_accuracy_mean'))}; {ival(r.get('analysed_sites'))} sites)"
        for r in site_predictability_ranked[:5]
    )
    + ". ACC, CHOL, DLBC, MESO and UCS had fewer than two sites meeting the 10-patient requirement and were reported as ineligible. This result shows substantial submitting-site information in the fixed TITAN representation and strengthens the confounding concern; it does not identify a causal scanner, laboratory or staining effect."
)
add_figure(doc, "Figure6_site_grouped_sensitivity.png", "Figure 6. Tissue-source-site sensitivity as a principal result. Panel A compares all 323 random-fold and grouped estimates; Panel B shows the eight largest attenuations; Panel C shows the 15 strongest chance-normalised within-cancer submitting-site results. Larger typography and restricted ranked panels improve readability. Grouping remains internal to TCGA and is not institutional or scanner-level external validation.")

doc.add_heading("Multiple-slide sensitivity", level=2)
doc.add_paragraph(
    f"Replacing each deterministic patient mean with the lexicographically first eligible diagnostic slide produced a median first-slide-minus-mean-pool performance change of {fnum(median(pool_deltas))} across screen-positive models. This patient-level sensitivity is distinct from the tissue-source-site analysis."
)

doc.add_heading("Qualitative morphology context", level=2)
doc.add_paragraph(
    "Five representative models were examined using high/low repeated out-of-fold anchors and "
    "their nearest within-cancer patient-level TITAN neighbours. The anchors were "
    + "; ".join(
        f'{r.get("tumor_type")}–{r.get("endpoint")}: high {r.get("high_anchor")}, low {r.get("low_anchor")}'
        for r in morphology_context_summary
    )
    + ". Exact slide identifiers, observed values, original prediction values, prediction ranks, "
    "embedding cosine similarities and generated-report excerpts are reported in Supplementary "
    "Table S14 and the machine-readable file; Figure S4 places all held-out patients behind the "
    "selected examples. The BLCA TIL Regional Fraction example is explicitly a same-H&E-modality "
    "concordance control. The remaining examples provide descriptive embedding context only. No "
    "patch-level attribution or blinded pathologist interpretation was possible from the released "
    "global embeddings and automated reports."
)

doc.add_heading("External-validation readiness audit", level=2)
doc.add_paragraph("No non-TCGA patient entered model evaluation, and no external performance result is reported. The official TITAN release provided the precomputed TCGA representation used here, but the audit did not identify a compatible non-TCGA TITAN embedding set carrying the required outcomes. CPTAC-UCEC is a plausible future cohort: its official TCIA collection reports 250 subjects, 887 pathology whole-slide images (approximately 154 GB) and links to genomic, proteomic and clinical resources [36]. Those slides were not downloaded or processed in this study, and the availability of linked data does not by itself guarantee exact reconstruction of each endpoint.")
doc.add_paragraph("Three UCEC targets—TP53 mutation, genome doubling and continuous aneuploidy score—were prospectively locked for future independent evaluation, but no external result is available. Target definitions, internal estimates, artifact hashes, endpoint-compatibility rules and mandatory failure reporting are provided in Supplementary Table S6c and the software documentation.")

doc.add_heading("Illustrative COAD output", level=2)
doc.add_paragraph("The fitted-model resource was implemented as research software; construction, model inventory, checksums, score-rank definitions, reliability warnings, input validation and report behaviour are described in Supplementary Methods, Tables S6b–S6c and S13, and the package documentation [32].")
doc.add_paragraph(
    "Two COAD participants, TCGA-AA-A01F and TCGA-AA-3972, are retained solely to show continuous and binary research-software outputs (Figure 7). The pair was chosen post hoc after clinical-text and profile-saturation restrictions and then maximising Euclidean separation across continuous TCGA OOF prediction percentiles; the procedure intentionally enhances visual contrast. Nodal status was not matched: TCGA-AA-A01F was reported as pN1, whereas the available slide summary for TCGA-AA-3972 did not state a nodal category. The cases are not presented as a matched clinical comparison, and treatment, response, recurrence, follow-up and survival are neither shown nor interpreted. This is a visualization example, not performance, prevalence, prognosis or treatment-response evidence."
)
add_figure(
    doc,
    "Figure7_COAD_TITANPred_examples.png",
    "Figure 7. Post hoc TITANPred research-software visualization for two COAD participants. Panel A replaces separate radar plots with a common-scale dumbbell comparison of continuous TCGA OOF ranks; the right column preserves each original prediction. Panel B compares PLS-LDA score ranks and calls with development class counts. All ranks are explicitly not probabilities. The pair was selected after clinical-text and saturation restrictions by maximising continuous-profile separation, intentionally enhancing contrast. Treatment and outcomes are not displayed. This is an interface illustration, not external validation or representative case sampling.",
    width=6.7,
)

doc.add_heading("Discussion", level=1)
doc.add_paragraph(f"This study is best interpreted as a systematic patient-level benchmark and reusable fitted-model resource, not as the first pan-cancer molecular-prediction study or a catalogue of unprecedented mutation targets. It answers a target-by-target, cancer-specific question using one fixed pretrained TITAN representation. Selected immune programmes, mutations and higher-level genomic phenotypes were predictable, but most eligible pairs did not satisfy both multiplicity and prespecified screening-statistic thresholds. That heterogeneity—and transparent negative and ineligible reporting—is the primary empirical result. The primary analysis controls local cancer-specific families; it does not attach one global 5% FDR interpretation to the aggregate {ival(combined_multiplicity.get('within_cancer_family_candidates'))} candidates. In sensitivity analyses, {len(global_supported_c)} continuous and {len(global_mutation_b)} mutation pairs passed the across-cancer family correction, and {ival(combined_multiplicity.get('atlas_wide_pass'))}/{ival(combined_multiplicity.get('within_cancer_family_candidates'))} local candidates passed a single BH correction across all eligible atlas tests. These large counts partly reflect extensive ties at the 0.001 empirical-p floor: a zero-exceedance 999-permutation result still has a 95% Monte Carlo interval extending to {fnum(zero_999_upper, 6)}. The 9,999-permutation refinement sharpened precision for eight leading claims but does not resolve the entire atlas or justify ranking by q-value; outcome-appropriate predictive magnitude and stability remain the appropriate ordering criteria.")
doc.add_paragraph("The closest breadth comparator is Arslan et al., who already trained 12,093 target-specific models for 4,031 multi-omic biomarkers in 8,890 TCGA patients across the same 32 cancers [13]. The present study is smaller in endpoint breadth and lacks external validation. Its differentiating elements are the fixed pretrained representation, deterministic pre-outcome patient aggregation, nested patient-level validation, complete tested-negative and ineligible outputs, target-level site-grouped sensitivity, continuous modelling of selected immune and genomic-context measures, and distribution of compact fitted linear predictors. These distinctions—not breadth or endpoint novelty—motivate the study.")
doc.add_paragraph("The relevant literature spans substantially different validation settings and metrics. For MSI, Kather et al. reported external AUROCs of 0.84 (95% CI 0.72–0.92) in DACHS colorectal cancer and 0.69 (0.52–0.82) in an external gastric cohort, whereas Echle et al. reported an external dMMR AUROC of 0.96 in a 771-case cohort [2,3]. For fusion prediction, Dadhania et al. reported ERG AUROCs of 0.82–0.85, and Mayer et al. reported 100% sensitivity with 100% and 98.6% specificity for ALK and ROS1, respectively, in a small external set [9,10]. Saldanha et al. externally evaluated mutation models in CPTAC, Bergstrom et al. reported external breast HRD AUROC 0.76 (0.71–0.82), and HistoTME reported mean external cell-composition correlation 0.50 and immunotherapy-response AUROC 0.75 (0.61–0.88) [11,14,15]. These externally tested results set a stronger validation standard than the internal TCGA estimates presented here; metric, endpoint and cohort differences preclude claims of direct superiority.")
doc.add_paragraph(
    f"The expanded primary-study audit substantially changes the novelty interpretation. "
    f"{len(prior_supported_mutations)}/{len(mutation_literature_audit)} screen-positive cancer–gene pairs "
    "had already received statistical support in reviewed histology-prediction studies, including pooled "
    "colorectal evidence mapped to both COAD and READ where appropriate. Examples include LGG–CIC "
    f"(prior mean AUROC 0.836, current balanced accuracy {fnum(next(r for r in mutation_literature_audit if r['cancer']=='LGG' and r['gene']=='CIC')['current_balanced_accuracy'])}), "
    "LGG–ATRX (prior mean AUROC 0.816), THCA–NRAS (0.809), LIHC–BAP1 (0.825), "
    "PAAD–TP53 (0.672) and KIRC–VHL (0.631) [13]. Prior studies generally reported AUROC, whereas "
    "the present primary metric is balanced accuracy; these values are contextual and are not estimates "
    "of improvement or inferiority."
)
doc.add_paragraph(
    f"Two current screen-positive pairs had been evaluated but were not statistically supported in the "
    f"reviewed 2024 study: BLCA–PIK3CA (prior mean AUROC 0.588, corrected p=0.058; current balanced "
    f"accuracy {fnum(next(r for r in prior_evaluated_not_supported if r['cancer']=='BLCA')['current_balanced_accuracy'])}) "
    f"and SKCM–BRAF (prior mean AUROC 0.579, corrected p=0.286; current balanced accuracy "
    f"{fnum(next(r for r in prior_evaluated_not_supported if r['cancer']=='SKCM')['current_balanced_accuracy'])}). "
    "Only THYM–GTF2I was not identified in the reviewed predictive-model literature; a strong direct "
    "GTF2I–thymoma morphology association is already established, so this result is presented as a "
    "prediction candidate requiring independent validation rather than a claim of biological novelty. "
    "The OV–TP53 entry is supported by a 2025 preprint reporting AUROC 0.82±0.02 and is labelled as "
    "preliminary evidence rather than peer-reviewed external validation [34]. The GTF2I morphology "
    "association itself was reported by Wells et al. [35]. Supplementary Table S10a and the "
    "machine-readable audit give the evidence scope, prior metric, source "
    "and note for every screen-positive pair."
)
doc.add_paragraph(
    "External results also caution against equating internal discrimination with transportability. For PAAD–TP53, "
    "the current TCGA balanced accuracy was 0.619 and Arslan et al. reported mean internal AUROC 0.672, "
    "whereas Saldanha et al. reported internal AUROC 0.554 and external CPTAC AUROC 0.443±0.064 [11,13]. "
    "Differences can arise from cohort composition, slide preparation, mutation definition, model class and "
    "validation design. The present result therefore supports further locked external testing, not clinical application."
)
doc.add_paragraph("The non-mutation endpoints extend the analysis beyond the exact cancer–gene pairs used for the literature crosswalk, but they must not be collapsed under the phrase 'immune measurements'. Direct sequence-supported alterations, sequencing-derived burdens, inferred immune-cell fractions, transcriptomic signatures, pathology-derived quantities and composite genomic-context scores have different error models and biological meanings. In particular, prediction of a CIBERSORT fraction or RNA signature is agreement with a computational phenotype derived from a bulk specimen; it is not equivalent to flow cytometry, immunohistochemistry, a direct cell count or a clinically certified biomarker. TIL Regional Fraction is an even more specific same-modality case because both predictor and target originate from H&E. Its result is best interpreted as concordance between two image-derived summaries rather than molecular inference.")
doc.add_paragraph(f"For continuous outcomes, prior work predicted RNA expression [6], tumour composition [4], continuous HRD and microenvironment features [12], and externally evaluated tumour-microenvironment composition in non-small cell lung cancer [15]. The present study instead screens the prespecified Thorsson-derived panel, genomic-context scores, MSI scores, fusion burden and aneuploidy burdens across eligible TCGA cancers using the same patient-level validation design. Only {len(global_supported_c)} continuous pairs passed the stricter across-cancer correction; the others should be treated as endpoint- and cancer-specific nominations, not as replications of the external HistoTME results or as evidence of general immune profiling.")
doc.add_paragraph("The contribution is therefore not that TITAN plus PLS is an optimal algorithmic combination. It is the uniform, auditable patient-level screen over prespecified outcomes using one fixed representation, complete negative reporting and compact fitted reference models. PLS1 and PLS–LDA remain the prespecified reference analysis, but the released PLS objects should not be interpreted as the best available models.")
doc.add_paragraph(f"The metadata-stratified benchmark directly weakens any stronger PLS claim. Selection depended only on outcome family, sample size, binary minority-class fraction and a fixed hash, not on PLS performance. Both algorithms used matched outer and inner folds; continuous hyperparameters maximised the same pooled inner Q² and binary thresholds maximised the same inner balanced accuracy. Median ridge-minus-PLS differences were {fnum(binary_baseline_summary.get('median_delta_ridge_minus_pls'))} AUROC and {fnum(continuous_baseline_summary.get('median_delta_ridge_minus_pls'))} Q². Paired intervals favoured ridge for {sum(r.get('outcome_type') == 'binary' for r in ridge_better)}/{len(ridge_binary)} binary and {sum(r.get('outcome_type') == 'continuous' for r in ridge_better)}/{len(ridge_continuous)} continuous targets, PLS for none, and were otherwise uncertain. Portability cannot justify PLS over ridge because both yield compact exportable parameters. We did not replace the locked primary analysis or add post hoc method switching; instead, we report PLS as the prespecified reference and make clear that future model-resource development should compare locked PLS and ridge candidates in independent cohorts.")
doc.add_paragraph("PLS2 addresses a different question. Its mean cancer-level ΔQ² relative to response-by-response PLS1 was 0.066 for inferred cell fractions, 0.017 for infiltration signatures and 0.056 for immune-repertoire measures, supporting shared inflammatory structure. It remains secondary because each joint fit uses only the complete-case intersection for a coherent continuous response block, shares one latent representation across a fixed ordered panel and cannot replace individual mutation, pathway, MSI, aneuploidy or fusion classifiers. The positive internal averages therefore nominate PLS2 for a future locked multiresponse immune resource; they do not establish a generally preferable main model without external validation and a matched multiresponse baseline.")
doc.add_paragraph(f"The prespecified 20-per-class rule supported an inclusive atlas screen, but it was too permissive for presenting every selected binary model as equally mature. Only {len(binary_standard)}/{len(supported_b)} screen-positive binary models met the stricter 50-per-class criterion. The {len(binary_limited_reliability)} smaller-class models had sparse inner validation folds, frequently selected components across a wide range and showed lower median repeat score correlation than the ≥50-per-class group. Their learning curves showed no uniformly monotonic performance gain, illustrating that the available data cannot establish a stable sample-size plateau. We therefore retain them only as transparently reported exploratory results and exclude them from default inference. PR-AUC adds prevalence-aware discrimination information, whereas PPV and NPV are explicitly conditional on each analysed TCGA prevalence and should not be exported to another population. High call agreement under imbalance can be misleading and does not supersede score stability, sensitivity, PR-AUC or external validation.")
doc.add_paragraph(f"Tissue-source-site sensitivity is a central result rather than a generic robustness footnote. Grouping revealed heterogeneity that the median attenuation concealed: {ival(site_combined.get('below_threshold_models'))}/{ival(site_combined.get('screen_positive_models'))} screen-positive models ({fnum(site_combined.get('below_threshold_percent'), 1)}%) fell below their original prespecified screening threshold. READ–APC and COAD–APC declined from balanced accuracies 0.862 and 0.793 to 0.495 and 0.543, respectively. Smaller and less balanced training folds contribute to grouped-validation loss, but performance close to chance requires these models to be labelled site-sensitive rather than robust biological signals. The model registry and inference output therefore expose the grouped metric, delta, number of sites, threshold-retention status and a prominent warning for every affected endpoint.")
doc.add_paragraph(f"The complementary site-classification analysis strengthens this interpretation: the fixed TITAN representation predicted submitting site within {len(site_predictability_eligible)}/32 evaluable cancers, with median macro balanced accuracy {fnum(median([r.get('macro_balanced_accuracy_mean') for r in site_predictability_eligible]))} despite cancer-specific chance references from 0.053 to 0.500. This demonstrates that site information is present in the representation and could support confounding. It does not show that each attenuated molecular model is wholly artifactual, because site correlates with case mix and molecular prevalence as well as technical acquisition. Conversely, successful tissue-source-site grouping cannot prove transportability: the barcode-derived site is an imperfect proxy for institutions, scanners, staining laboratories and batch effects.")
doc.add_paragraph("A further practical feature is data-minimising model portability. Preprocessing values, latent weights and coefficients are sufficient for continuous prediction, with compact LDA parameters added for binary classification. The access-controlled TITANPred package applies cancer-matched models to correctly ordered TITAN features without requiring the original patient embeddings or outcomes [32]. Its current private status is stated explicitly; the manuscript makes no public-artifact availability claim. This design reduces the patient-level data that would need to be exchanged in authorised research, although it does not itself confer privacy, satisfy local governance requirements or establish transportability. The synthetic smoke-test vector supports routine interface testing without creating a patient narrative. The two main-text COAD examples remain explicitly post hoc and cannot provide new accuracy, prognosis or response-prediction evidence.")
doc.add_paragraph("Strengths include primary-tumour matching across every data source, deterministic mean pooling of multiple slides, nested patient-level validation, FDR control within cancer and endpoint family, exact negative-result reporting, site-grouped sensitivity, saved research models and executable inference code. The supplied Bonneville spreadsheets contained only an ACC/CESC/MESO subset; cBioPortal PanCancer Atlas MANTIS fields were therefore used for full coverage and agreed exactly for all 387 overlapping cases.")
doc.add_paragraph("The principal limitation is absence of independent external validation. TCGA resampling—including the analysis termed 'TCGA tissue-source-site-grouped internal validation'—cannot establish transportability to another institution, scanner, stain distribution or patient population. Tissue-source site is not a complete scanner or laboratory identifier. The study is retrospective and exploratory; thresholds are prioritisation rules, not clinical operating points. Although TCGA was excluded from TITAN's Mass-340K pretraining corpus, it was used in the original paper's downstream evaluation; the present cohort is therefore neither a pretraining cohort nor an independent validation cohort. Predictability does not establish biological causality, and image-derived estimates cannot justify omitting a molecular assay. The three-model UCEC subset is locked only for future independent evaluation. Until that untouched analysis reports every target, the work should be evaluated as a computational pathology benchmark and research-software resource rather than a translational prediction study.")
doc.add_paragraph("CPTAC-UCEC illustrates both feasibility and the remaining work: the official collection provides pathology images and links to molecular data, but the approximately 154-GB slide archive requires de novo TITAN extraction and exact outcome harmonisation [36]. Neither the slides nor external labels were inspected here. The locked protocol prohibits model refitting, recalibration and threshold changes; mandates the same patient-level pooling; and requires failure or non-evaluability to be reported. This protocol improves readiness and guards against post hoc target selection, but it is not external validation evidence.")
doc.add_paragraph("The reported uncertainty is deliberately labelled a 95% selection-conditioned patient-resampling interval for repeated out-of-fold predictions. It represents patient sampling variability conditional on the five fitted nested-CV prediction sets and incorporates the five repeat-specific metrics into each bootstrap calculation. It does not repeat the initial screen or highlighting decision, generate new fold partitions, re-estimate scaling, reselect components, refit a model, correct winner's-curse optimism, or sample a new institution, scanner, staining process or population. It is therefore not a conventional confidence interval for model generalisation or an external-performance interval; locked-model evaluation in an untouched cohort is required for transportability assessment.")
doc.add_paragraph("Arithmetic mean pooling gives each eligible slide equal weight and prevents multiple-slide leakage, but it does not model variable tissue area or within-patient morphological heterogeneity. The first-slide sensitivity measures dependence on one deterministic alternative and should not be interpreted as a comparison with learned or tissue-area-weighted aggregation.")
doc.add_paragraph("The 768-dimensional TITAN representation is not morphologically self-explanatory. The new high/low anchor and within-cancer nearest-neighbour analysis places representative predictions in the context of globally similar patient embeddings and preserves exact slide/report provenance, but it cannot determine which patch or tissue compartment drove a prediction. The available TCGA-Slide-Reports text was generated by TITAN from the same slides and is neither an independent annotation nor a blinded pathologist review. Consequently, descriptive terms such as lymphoid infiltrate, stromal pattern or necrosis are hypotheses for future review, not validated mechanisms. Tile-level representations or relevance maps retained prospectively, followed by blinded pathologist assessment, are required before morphological explanations can be claimed.")
doc.add_paragraph("Demographic subgroup performance and algorithmic fairness were not evaluated. The TCGA Clinical Data Resource provides descriptive age, recorded gender, race and stage fields, but coverage and subgroup sizes vary across cancers and would fragment further under endpoint-specific molecular missingness. Tissue-source-site grouping probes one form of institutional heterogeneity but cannot establish fairness across ancestry, sex, age or access-to-care groups. Any external evaluation should prespecify representative sampling and subgroup performance, calibration and failure analysis.")

doc.add_heading("Conclusions", level=1)
doc.add_paragraph(f"Using mean-pooled fixed pretrained TITAN representations, this study provides a systematic patient-level TCGA benchmark of direct genomic alterations and derived immune, genomic-context, MSI, aneuploidy and fusion phenotypes. The access-controlled TITANPred package demonstrates how fitted PLS and PLS–LDA objects can be applied without distributing patient-level training data. Its default interface includes all {len(supported_c)} continuous models and the {len(binary_standard)} binary models meeting at least 50 patients per class; {len(binary_limited_reliability)} smaller-class binary models require explicit exploratory opt-in. All reported performance and example outputs are internally derived TCGA estimates. Prediction of a computationally inferred target is not evidence that a direct assay has been reproduced, and the qualitative neighbour examples are not patch-level explanation. The contribution is a transparent computational pathology benchmark and research-software framework—not broad endpoint novelty, external performance evidence or clinical validation—and locked independent testing remains essential.")

doc.add_heading("Abbreviations", level=1)
doc.add_paragraph(
    "AUROC, area under the receiver operating characteristic curve; BA, balanced "
    "accuracy; CV, cross-validation; FDR, false discovery rate; H&E, "
    "haematoxylin and eosin; LDA, linear discriminant analysis; MC3, Multi-Center "
    "Mutation Calling in Multiple Cancers; MSI, microsatellite instability; OOF, "
    "out-of-fold; PLS, partial least squares; PLS1, single-response PLS; PLS2, "
    "multi-response PLS; Q², cross-validated coefficient of determination; RMSE, "
    "root-mean-square error; TCGA, The Cancer Genome Atlas."
)

doc.add_heading("Declarations", level=1)
for h, text in [
    ("Ethics approval and consent to participate", "This secondary analysis used publicly available, de-identified TCGA data. No participants were recruited and no new tissue was collected."),
    ("Consent for publication", "Not applicable."),
    ("Study registration and protocol", "This retrospective secondary analysis was not registered. The locked executable analysis plan is available in the companion repository [30]."),
    ("Patient and public involvement", "Patients and members of the public were not involved in the design, conduct, interpretation or reporting of this secondary analysis."),
    ("Availability of data and materials", "The analysis used the gated public TITAN TCGA feature artifact TCGA_TITAN_features.pkl, downloaded from https://huggingface.co/MahmoodLab/TITAN/blob/main/TCGA_TITAN_features.pkl and converted to the documented 768-feature CSV schema with tools/convert_tcga_titan_pickle.py in the companion repository [30]. The analysed CSV contained 11,658 slide rows (SHA-256 d3d91fb0f83a6de440eda5ff437a63e3ca13f50095e6841fb8efcc40e58763f0). TCGA molecular data and the cited public supplementary records remain available from their original repositories. Exact source locations, access conditions and checksums are recorded in the companion repository. Controlled-access whole-slide images are not redistributed. The model registry, hashes, inference example and future external-evaluation protocol are in the public analysis repository. The separate TITANPred package and 323 fitted research objects are currently maintained in a private access-controlled repository [32]; this manuscript does not claim that those artifacts are publicly downloadable."),
    ("Competing interests", "[Author confirmation required before submission.]"),
    ("Funding", "[Author confirmation required before submission.]"),
    ("Authors’ contributions", "[Author initials and contribution statement required before submission.]"),
    ("Acknowledgements", "OpenAI Codex assisted with code refactoring and language editing. Human authors remain responsible for verification, interpretation and the submitted text."),
]:
    doc.add_heading(h, level=2); doc.add_paragraph(text)

doc.add_heading("Additional files", level=1)
doc.add_paragraph(
    "Additional file 1 (.docx): Supplementary material. Contains Supplementary "
    "Methods, Tables S1–S14, Figures S1–S4 and an inventory of the companion "
    "machine-readable result files."
)
doc.add_paragraph(
    "Additional file 2 (.pdf): COAD example A (TCGA-AA-A01F) TITANPred research-software output. "
    "Post hoc supplementary interface illustration containing internally derived full-cohort TCGA estimates; not probability-calibrated or externally validated."
)
doc.add_paragraph(
    "Additional file 3 (.pdf): COAD example B (TCGA-AA-3972) TITANPred research-software output. "
    "Post hoc supplementary interface illustration containing internally derived full-cohort TCGA estimates; not probability-calibrated or externally validated."
)

doc.add_heading("References", level=1)
references = [
"1. Coudray N, et al. Classification and mutation prediction from non-small cell lung cancer histopathology images using deep learning. Nat Med. 2018;24:1559–1567. doi:10.1038/s41591-018-0177-5.",
"2. Kather JN, et al. Deep learning can predict microsatellite instability directly from histology in gastrointestinal cancer. Nat Med. 2019;25:1054–1056. doi:10.1038/s41591-019-0462-y.",
"3. Echle A, et al. Clinical-grade detection of microsatellite instability in colorectal tumors by deep learning. Gastroenterology. 2020;159:1406–1416.e11. doi:10.1053/j.gastro.2020.06.021.",
"4. Fu Y, et al. Pan-cancer computational histopathology reveals mutations, tumor composition and prognosis. Nat Cancer. 2020;1:800–810. doi:10.1038/s43018-020-0085-8.",
"5. Kather JN, et al. Pan-cancer image-based detection of clinically actionable genetic alterations. Nat Cancer. 2020;1:789–799. doi:10.1038/s43018-020-0087-6.",
"6. Schmauch B, et al. A deep learning model to predict RNA-Seq expression of tumours from whole slide images. Nat Commun. 2020;11:3877. doi:10.1038/s41467-020-17678-4.",
"7. Howard FM, et al. The impact of site-specific digital histology signatures on deep learning model accuracy and bias. Nat Commun. 2021;12:4423. doi:10.1038/s41467-021-24698-1.",
"8. Loeffler CML, et al. Predicting mutational status of driver and suppressor genes directly from histopathology with deep learning. Front Genet. 2022;12:806386. doi:10.3389/fgene.2021.806386.",
"9. Dadhania V, et al. Leveraging artificial intelligence to predict ERG gene fusion status in prostate cancer. BMC Cancer. 2022;22:494. doi:10.1186/s12885-022-09559-4.",
"10. Mayer C, et al. Direct identification of ALK and ROS1 fusions in non-small cell lung cancer from hematoxylin and eosin-stained slides using deep learning algorithms. Mod Pathol. 2022;35:1882–1887. doi:10.1038/s41379-022-01141-4.",
"11. Saldanha OL, et al. Self-supervised attention-based deep learning for pan-cancer mutation prediction from histopathology. npj Precis Oncol. 2023;7:35. doi:10.1038/s41698-023-00365-0.",
"12. El Nahhas OSM, et al. Regression-based deep-learning predicts molecular biomarkers from pathology slides. Nat Commun. 2024;15:1253. doi:10.1038/s41467-024-45589-1.",
"13. Arslan S, et al. A systematic pan-cancer study on deep learning-based prediction of multi-omic biomarkers from routine pathology images. Commun Med. 2024;4:48. doi:10.1038/s43856-024-00471-5.",
"14. Bergstrom EN, et al. Deep learning artificial intelligence predicts homologous recombination deficiency and platinum response from histologic slides. J Clin Oncol. 2024;42:3550–3560. doi:10.1200/JCO.23.02641.",
"15. Patkar S, et al. Predicting the tumor microenvironment composition and immunotherapy response in non-small cell lung cancer from digital histopathology images. npj Precis Oncol. 2024;8:280. doi:10.1038/s41698-024-00765-w.",
"16. Ding T, et al. A multimodal whole-slide foundation model for pathology. Nat Med. 2025;31:3749–3761. doi:10.1038/s41591-025-03982-3.",
"17. Wold S, Sjöström M, Eriksson L. PLS-regression: a basic tool of chemometrics. Chemometr Intell Lab Syst. 2001;58:109–130. doi:10.1016/S0169-7439(01)00155-1.",
"18. Szymańska E, et al. Double-check: validation of diagnostic statistics for PLS-DA models. Metabolomics. 2012;8(Suppl 1):3–16. doi:10.1007/s11306-011-0330-3.",
"19. Thorsson V, et al. The Immune Landscape of Cancer. Immunity. 2018;48:812–830.e14. doi:10.1016/j.immuni.2018.03.023.",
"20. Taylor AM, et al. Genomic and functional approaches to understanding cancer aneuploidy. Cancer Cell. 2018;33:676–689.e3. doi:10.1016/j.ccell.2018.03.007.",
"21. Gao Q, et al. Driver fusions and their implications in the development and treatment of human cancers. Cell Rep. 2018;23:227–238.e3. doi:10.1016/j.celrep.2018.03.050.",
"22. Bonneville R, et al. Landscape of microsatellite instability across 39 cancer types. JCO Precis Oncol. 2017;2017:PO.17.00073. doi:10.1200/PO.17.00073.",
"23. Cerami E, et al. The cBio cancer genomics portal: an open platform for exploring multidimensional cancer genomics data. Cancer Discov. 2012;2:401–404. doi:10.1158/2159-8290.CD-12-0095.",
"24. Ellrott K, et al. Scalable open science approach for mutation calling of tumor exomes. Cell Syst. 2018;6:271–281.e7. doi:10.1016/j.cels.2018.03.002.",
"25. Bailey MH, et al. Comprehensive characterization of cancer driver genes and mutations. Cell. 2018;173:371–385.e18. doi:10.1016/j.cell.2018.02.060.",
"26. Sanchez-Vega F, et al. Oncogenic signaling pathways in The Cancer Genome Atlas. Cell. 2018;173:321–337.e10. doi:10.1016/j.cell.2018.03.035.",
"27. Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J R Stat Soc B. 1995;57:289–300. doi:10.1111/j.2517-6161.1995.tb02031.x.",
"28. Collins GS, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.",
"29. Liu J, et al. An integrated TCGA pan-cancer clinical data resource to drive high-quality survival outcome analytics. Cell. 2018;173:400–416.e11. doi:10.1016/j.cell.2018.02.052.",
f"30. TITAN patient-level benchmark repository. GitHub. {REPO}. Accessed 15 Aug 2026.",
"31. International Agency for Research on Cancer. Global Cancer Observatory: GLOBOCAN 2022 world fact sheet. Lyon: IARC; 2024. https://gco.iarc.who.int/media/globocan/factsheets/populations/900-world-fact-sheet.pdf. Accessed 16 Aug 2026.",
f"32. TITANPred R package and fitted-model repository. GitHub. {MODEL_REPO}. Accessed 16 Aug 2026.",
"33. National Cancer Institute. Genomic Data Commons Cases API. https://api.gdc.cancer.gov/cases. Accessed 16 Aug 2026.",
"34. Fernandes G. Morpho-genomic deep learning for ovarian cancer subtype and gene mutation prediction from histopathology. arXiv. 2025;arXiv:2511.03365. doi:10.48550/arXiv.2511.03365.",
"35. Wells K, Lamrca A, Papaxoinis G, et al. Unique correlation between GTF2I mutation and spindle cell morphology in thymomas (type A and AB thymomas). J Clin Pathol. 2023;76:463–466. doi:10.1136/jclinpath-2021-207837.",
"36. National Cancer Institute Clinical Proteomic Tumor Analysis Consortium. CPTAC-UCEC: The Clinical Proteomic Tumor Analysis Consortium Uterine Corpus Endometrial Carcinoma Collection. The Cancer Imaging Archive. doi:10.7937/K9/TCIA.2018.3R3JUISW. Accessed 24 Aug 2026.",
"37. Newman AM, et al. Robust enumeration of cell subsets from tissue expression profiles. Nat Methods. 2015;12:453–457. doi:10.1038/nmeth.3337.",
"38. Saltz J, et al. Spatial organization and molecular correlation of tumor-infiltrating lymphocytes using deep learning on pathology images. Cell Rep. 2018;23:181–193.e7. doi:10.1016/j.celrep.2018.03.086.",
]
for ref in references: doc.add_paragraph(ref)
doc.save(OUT / "manuscript_JTM_patient_level_TITAN.docx")


# Supplementary document
sup = setup(Document(), "Supplementary material — TITAN benchmark")
sup.add_heading("Supplementary material", 0)
sup.add_paragraph(MANUSCRIPT_TITLE)
sup.add_paragraph("This document is Additional file 1. The COAD example A (TCGA-AA-A01F) and example B (TCGA-AA-3972) TITANPred research-software outputs are supplied as separate PDF attachments (Additional files 2 and 3, respectively).")
sup.add_heading("Supplementary Methods", 1)
sup.add_paragraph("The executable analysis plan, source manifest, eligibility catalogues, checkpoint-capable scripts and released out-of-fold prediction tables are available in the public companion analysis repository. The separate access-controlled TITANPred R package contains the fitted research models, model registry, reference distributions and research-software output template; it is currently private and is not claimed as a public artifact release. Large local restart checkpoints are not redistributed. Tables below are concise views; complete machine-readable CSV files are authoritative.")
sup.add_heading("Binary class-size, fold-composition and stability analyses", 2)
sup.add_paragraph("The primary screening estimate is the initial nested-CV value used for candidate selection, permutation testing and tier assignment. The repeated-validation estimate is the mean across five additional independently seeded nested-CV partitions and is reported as a post-selection stability estimate. The values need not be identical because the fold partitions and rSVD seeds differ; THYM–Th17, for example, had primary screening Q²=0.638 and five-repeat mean Q²=0.594. The prespecified screen retained a minimum of 20 positive and 20 negative participants so that rare but potentially informative cancer-specific endpoints were not silently discarded. We additionally evaluated a 50-per-class threshold without redefining the original multiplicity analysis. For all screen-positive binary models, every repeated outer fold was audited for training/test class counts, and the exact inner partitions used to select 1–10 components were reconstructed and audited for class counts. Component distributions summarize the 25 outer fits generated by five repeats of five-fold nested validation. PR-AUC is non-interpolated average precision; PPV and NPV use held-out calls and the observed TCGA prevalence and are not calibrated or transportable probabilities.")
sup.add_paragraph("Prediction stability used standardized held-out LDA scores because independently fitted discriminant scales are not directly comparable. For each pair of repeats, Spearman correlation was calculated over patient scores and agreement over class calls, then averaged. Limited-evidence learning curves used the original outer test fold at every fraction and fitted the model to deterministic stratified 50%, 75% and 100% subsets of that fold's training patients. The full-data point exactly reproduces the original repeated nested-CV result. Small-class screen-positive models remain in the complete analysis registry but are excluded from default TITANPred inference; an explicit include_limited_evidence=TRUE opt-in emits a warning.")
sup.add_heading("Endpoint provenance and qualitative morphology context", 2)
sup.add_paragraph("Every eligible cancer–endpoint target was assigned a label-generation class before interpretation: directly observed genomic alteration, sequencing-derived continuous burden, computationally inferred immune-cell fraction, transcriptomic signature, pathology-associated quantity or composite genomic-context score. The machine-readable endpoint_dictionary.csv is the complete Supplementary Table S13 dictionary and contains one row for every target, including source modality, direct/inferred status, derivation algorithm, original scale, analysis transformation, missingness, expected measurement error, biological interpretation and an assay-equivalence caveat. endpoint_definition_dictionary.csv collapses these rows to unique endpoint definitions. TIL Regional Fraction alone is flagged as a same-H&E-modality target.")
sup.add_paragraph("For the morphology-context sensitivity, five models were selected to represent different label modalities rather than to imply exhaustive interpretability. High and low anchors were restricted to report-covered concordant extremes of mean repeated out-of-fold predictions. Their nearest within-cancer neighbours were selected by cosine similarity of patient-level mean TITAN representations within the corresponding high/low prediction stratum. A representative report-covered slide was chosen by minimum squared distance to the patient mean. Exact prediction values and ranks were retained without dichotomising continuous outputs. Prespecified terms and excerpts were extracted from TITAN-generated slide reports only for contextual audit. This analysis cannot localise patches or support causal morphology because tile representations, pixels and independent blinded pathology annotations were unavailable.")
sup.add_heading("Secondary PLS1–PLS2 inflammatory comparison", 2)
sup.add_paragraph("The secondary comparison jointly modelled three prespecified inflammatory blocks: infiltration/signatures, immune repertoire and inferred immune-cell fractions. Response-by-response PLS1 and joint PLS2 used identical complete-case patients and outer folds, training-fold outcome scaling, separately selected component counts and three independently seeded nested-CV repeats. Changes in Q² were aggregated first within cancer; uncertainty was estimated by resampling cancers. Endpoint win counts were not treated as inferential evidence. PLS2 was not substituted into the main resource because it applies only to coherent continuous panels, requires the intersection of patients with every response observed, shares a latent representation across a fixed ordered response set and changes the estimand from one endpoint to a joint panel. It cannot replace the individual binary PLS–LDA models. Its positive average results motivate a future locked multiresponse immune resource, but no external validation or matched multiresponse ridge baseline was available.")
sup.add_heading("Table S1. Analysis coverage", 1)
add_table(sup, ["Family", "Type", "Tests", "Cancers", "Screening tier A", "Screening tier B", "Screen-negative"], family_table)
if participant_characteristics:
    sup.add_heading("Table S2. Participant characteristics from the TCGA Clinical Data Resource", 1)
    sup.add_paragraph("Counts are descriptive source fields. Gender and race are reported using the CDR categories; W/B/A/O denote White, Black or African American, Asian and other recorded race. Stage combines available pathologic stage with clinical stage only when pathologic stage is unavailable. The full machine-readable table retains availability counts and broad race/stage components.")
    add_table(sup, ["Cancer", "n", "CDR matched", "Age, median (IQR)",
                    "Gender F/M/missing", "Race W/B/A/O/missing",
                    "Stage I/II/III/IV/missing"],
              [(r.get("tumor_type"), r.get("patients"), r.get("cdr_matched"),
                f'{fnum(r.get("age_median"), 1)} ({fnum(r.get("age_q1"), 1)}–{fnum(r.get("age_q3"), 1)})',
                f'{r.get("female")}/{r.get("male")}/{r.get("gender_missing")}',
                f'{r.get("race_white")}/{r.get("race_black_or_african_american")}/{r.get("race_asian")}/{r.get("race_other_recorded")}/{r.get("race_missing")}',
                f'{r.get("stage_I")}/{r.get("stage_II")}/{r.get("stage_III")}/{r.get("stage_IV")}/{r.get("stage_missing_or_other")}')
               for r in participant_characteristics])
sup.add_heading("Table S3. Slide multiplicity and report coverage by cancer", 1)
add_table(sup, ["Cancer", "Patients", "Slides", "Multi-slide patients", "Maximum slides", "Multiple primary barcodes", "Exact report matches"],
          [(r.get("tumor_type") or "Unresolved", r.get("patients"), r.get("eligible_slides"),
            r.get("multi_slide_patients"), r.get("maximum_slides_per_patient"),
            r.get("patients_with_multiple_primary_sample_barcodes"), r.get("exact_report_matched_slides"))
           for r in slide_multiplicity])
sup.add_heading("Table S4. Mutation molecular coverage by cancer", 1)
sup.add_paragraph(f"Across {len(mutation_eligibility):,} prespecified cancer–gene pairs, {n_mutation_eligible:,} were eligible, {sum(r.get('eligibility') == 'insufficient_positive' for r in mutation_eligibility):,} had fewer than 20 positive patients and {sum(r.get('eligibility') == 'insufficient_negative' for r in mutation_eligibility):,} had fewer than 20 negative patients. These ineligible pairs were not entered into model fitting or multiplicity correction.")
sup.add_paragraph("The machine-readable mutation_variant_classification_audit.csv lists every MC3 variant class observed in the modelling slot and whether it was retained under the explicit protein-altering definition.")
add_table(sup, ["Cancer", "Embedding patients", "MC3-profiled", "Missing MC3", "Multiple primary aliquots"],
          [(r.get("tumor_type"), r.get("titan_embedding_patients"), r.get("matched_profiled_patients"),
            r.get("embedding_patients_without_mc3_profile"), r.get("profiled_patients_with_multiple_primary_aliquots"))
           for r in mutation_coverage])
sup.add_heading("Table S5. Non-mutation molecular-source coverage", 1)
add_table(sup, ["Source", "Cohort patients", "Covered", "Missing", "Multiple primary aliquots", "Maximum aliquots", "Aggregation rule"],
          source_coverage_summary)
if highlighted:
    sup.add_heading("Table S6a. Highlighted-model performance and uncertainty", 1)
    sup.add_paragraph("The initial nested-CV screening estimate and the mean across five additional independently seeded nested-CV repeats are labelled separately. SC interval denotes the 95% selection-conditioned patient-resampling interval for repeated out-of-fold predictions. Patients are resampled from five fixed held-out prediction sets; screening, new fold generation, tuning and model fitting are not repeated within bootstrap replicates.")
    performance_rows = []
    deployment_rows = []
    for r in highlighted:
        if r.get("outcome_type") == "binary":
            primary = (f'Initial screen BA {fnum(r.get("primary_screen_balanced_accuracy"))}; '
                       f'five-repeat BA {fnum(r.get("balanced_accuracy"))} '
                       f'({fnum(r.get("balanced_accuracy_ci_low"))}–{fnum(r.get("balanced_accuracy_ci_high"))})')
            class_metrics = (f'Se {fnum(r.get("sensitivity"))} '
                             f'({fnum(r.get("sensitivity_ci_low"))}–{fnum(r.get("sensitivity_ci_high"))}); '
                             f'Sp {fnum(r.get("specificity"))} '
                             f'({fnum(r.get("specificity_ci_low"))}–{fnum(r.get("specificity_ci_high"))})')
            other = (f'AUROC {fnum(r.get("auc"))} '
                     f'({fnum(r.get("auc_ci_low"))}–{fnum(r.get("auc_ci_high"))}); '
                     f'PR-AUC {fnum(r.get("pr_auc"))} '
                     f'({fnum(r.get("pr_auc_ci_low"))}–{fnum(r.get("pr_auc_ci_high"))}); '
                     f'PPV/NPV at TCGA prevalence {fnum(r.get("ppv_tcga_prevalence"))}/'
                     f'{fnum(r.get("npv_tcga_prevalence"))}')
        else:
            primary = (f'Initial screen Q² {fnum(r.get("primary_screen_q2"))}; '
                       f'five-repeat Q² {fnum(r.get("q2"))} '
                       f'({fnum(r.get("q2_ci_low"))}–{fnum(r.get("q2_ci_high"))})')
            class_metrics = "Not applicable"
            other = (f'RMSE {fnum(r.get("rmse"))} '
                     f'({fnum(r.get("rmse_ci_low"))}–{fnum(r.get("rmse_ci_high"))}); '
                     f'ρ {fnum(r.get("spearman"))} '
                     f'({fnum(r.get("spearman_ci_low"))}–{fnum(r.get("spearman_ci_high"))})')
        performance_rows.append((
            r.get("outcome_type"), r.get("tumor_type") + "–" + r.get("endpoint"),
            r.get("n"), primary, class_metrics, other,
            f'{r.get("site_grouped_metric_name")} {fnum(r.get("site_grouped_metric"))}; '
            f'{r.get("site_robustness_status")}'
        ))
        model_metadata = (f'{r.get("model_id")}; {r.get("ncomp")} components; '
                          f'fit seed {r.get("model_fit_seed")}; '
                          f'{r.get("feature_dimension")} features; '
                          f'{r.get("aggregation")}; transform: '
                          f'{r.get("endpoint_transform")}; output: '
                          f'{r.get("output_units")}')
        classification_metadata = "Not applicable"
        if r.get("outcome_type") == "binary":
            classification_metadata = (
                f'coding: {r.get("class_labels")}; priors: '
                f'{r.get("class_priors")}; rule: {r.get("prediction_rule")}; '
                f'calibration: {r.get("calibration_status")}'
            )
        provenance_metadata = (
            f'model SHA-256 prefix: {str(r.get("model_sha256"))[:12]}; '
            f'feature-schema SHA-256 prefix: '
            f'{str(r.get("feature_schema_sha256"))[:12]}; TITAN-file SHA-256 '
            f'prefix: {str(r.get("titan_feature_file_sha256"))[:12]}; fastPLS '
            f'{r.get("model_fastPLS_version")} '
            f'({str(r.get("model_fastPLS_remote_sha"))[:7]}), '
            f'{r.get("model_backend")}/{r.get("model_svd_method")}; '
            f'rSVD oversampling {r.get("model_rsvd_oversample")}, power '
            f'{r.get("model_rsvd_power")}; '
            f'analysis fingerprint prefix: '
            f'{str(r.get("model_analysis_fingerprint"))[:12]}; '
            f'external validation: '
            f'{r.get("external_validation")}; patient training rows retained: '
            f'{r.get("contains_patient_level_training_rows")}; '
            f'TSS-grouped sites: {r.get("site_grouped_n_sites")}; '
            f'delta: {fnum(r.get("site_performance_delta"))}; '
            f'status: {r.get("site_robustness_status")}; '
            f'inner site separation: {r.get("site_grouped_inner_site_separation")}; '
            f'{r.get("redistribution_status")}'
        )
        deployment_rows.append((
            r.get("tumor_type") + "–" + r.get("endpoint"), model_metadata,
            classification_metadata, provenance_metadata
        ))
    add_table(sup, ["Type", "Cancer–endpoint", "n", "Primary metric (95% SC interval)",
                    "Sensitivity/specificity (95% SC interval)", "Other metrics (95% SC interval)",
                    "TSS-grouped metric/status"],
              performance_rows)
    sup.add_heading("Table S6b. Highlighted-model research-use and provenance metadata", 1)
    sup.add_paragraph(
        "Checksum prefixes are shown for readability; complete SHA-256 values are "
        "provided in highlighted_model_performance.csv and models/model_registry.csv."
    )
    add_table(sup, ["Cancer–endpoint", "Model and input", "Classification",
                    "Provenance and release status"], deployment_rows)
sup.add_heading("Table S6c. Prospectively locked subset for future independent evaluation", 1)
sup.add_paragraph("These three models were selected after the internal TCGA screen but before inspection of any external features, outcomes or predictions. The external analysis has not been performed. No refitting, recalibration, class-prior change, threshold adjustment or outcome optimisation is permitted. Every target must be reported, including failure or non-evaluability.")
add_table(
    sup,
    ["Order", "Target/type", "Model SHA-256", "Internal TCGA evidence", "External metrics", "Compatibility and reporting rule"],
    [(
        r["evaluation_order"], f'{r["cancer_type"]}–{r["endpoint"]} ({r["target_type"]})',
        r["model_sha256"],
        f'{r["internal_repeated_cv_primary"]}; {r["internal_repeated_cv_secondary"]}; {r["site_grouped_metric"]}',
        f'Primary: {r["locked_external_primary_metric"]}; secondary: {r["locked_external_secondary_metrics"]}',
        f'{r["endpoint_compatibility_rule"]}. {r["reporting_rule"]}. Status: {r["external_validation_status"]}.'
    ) for r in external_locked_targets],
)
sup.add_heading("Table S7. Top continuous results from the initial nested-CV screen", 1)
add_table(sup, ["Cancer", "Endpoint", "Family", "n", "Primary screen Q²", "q", "Category"],
          [(r["tumor_type"], r["endpoint"], r["family"], r["n"], fnum(r["q2"]), fnum(r["q_value"]), category(r["tier"])) for r in top_c[:60]])
sup.add_heading("Table S8. Top binary results from the initial nested-CV screen", 1)
add_table(sup, ["Cancer", "Endpoint", "Family", "n", "Positive", "Primary screen BA", "Five-repeat PR-AUC", "q", "Evidence/default"],
          [(r["tumor_type"], r["endpoint"], r["family"], r["n"], r["positive"], fnum(r["balanced_accuracy"]),
            fnum(binary_reliability_by_key.get((r["family"], r["tumor_type"], r["endpoint"]), {}).get("repeated_pr_auc_mean")),
            fnum(r["q_value"]),
            ("standard; included" if binary_reliability_by_key.get((r["family"], r["tumor_type"], r["endpoint"]), {}).get("default_inference") == "TRUE" else "exploratory/limited; excluded"))
           for r in top_b[:60]])
sup.add_heading("Table S9. Published exact cancer–gene results compared with the current screen", 1)
sup.add_paragraph("Prior studies predominantly reported AUROC, whereas the current prespecified metric is balanced accuracy. These values are displayed side by side for context but are not directly subtractable and do not constitute a head-to-head model comparison.")
add_table(sup, ["Study", "Cancer", "Gene", "Prior result", "Current primary-screen BA", "Current q", "Current status"],
          [(r["study"], r["cancer"], r["gene"], r["prior_metric"], fnum(r["current_balanced_accuracy"]), fnum(r["current_q"]), r["current_status"].replace("_", " ")) for r in lit_accuracy])
sup.add_heading("Table S10a. Expanded literature audit for every screen-positive mutation predictor", 1)
sup.add_paragraph("The audit maps pooled colorectal cohorts to COAD and READ where appropriate and distinguishes prior statistical support, prior evaluation without support, and a result not identified in the reviewed predictive-model literature. The last category is not a claim of biological novelty or an exhaustive proof of bibliographic novelty.")
add_table(sup, ["Cancer", "Gene", "Evidence class", "Prior study/scope", "Prior result", "Current primary-screen BA", "Current q", "Current category"],
          [(r["cancer"], r["gene"], r["evidence_class"],
            f'{r.get("prior_study")} / {r.get("prior_scope")}', r.get("prior_result"),
            fnum(r["current_balanced_accuracy"]), fnum(r["current_q"]), category(r["current_tier"]))
           for r in mutation_literature_audit])
sup.add_heading("Table S10b. Screen-positive models below the original prespecified screening threshold under site-grouped validation", 1)
add_table(sup, ["Type", "Family", "Cancer", "Endpoint", "n", "Sites", "Primary screen metric", "Site-grouped metric", "Delta"],
          [(r.get("outcome_type"), r.get("family"), r.get("tumor_type"), r.get("endpoint"),
            r.get("n"), r.get("n_sites"), fnum(r.get("original_metric")),
            fnum(r.get("site_grouped_metric")), fnum(r.get("delta")))
           for r in site_threshold_failures])
sup.add_heading("Table S10c. Tissue-source-site-grouped outer-fold composition for the two colorectal APC examples", 1)
sup.add_paragraph(f"The table shows every outer fold for the two near-chance APC examples. The complete {len(site_fold_details):,}-row file reports test and training patients, tissue-source sites, positive and negative cases, site identifiers, seeds and inner-site-separation diagnostics for every screen-positive model. Inner component selection used the same tissue-source-site grouping and had no site overlap.")
add_table(
    sup,
    ["Model", "Fold", "Test patients", "Test sites", "Test positive/negative", "Training patients", "Training sites", "Training positive/negative"],
    [(
        f'{r.get("tumor_type")}–{r.get("endpoint")}', r.get("outer_fold"),
        r.get("test_patients"), r.get("test_sites"),
        f'{r.get("test_positive")}/{r.get("test_negative")}',
        r.get("training_patients"), r.get("training_sites"),
        f'{r.get("training_positive")}/{r.get("training_negative")}'
    ) for r in site_fold_details
      if r.get("outcome_type") == "binary"
      and (r.get("tumor_type"), r.get("endpoint")) in {("READ", "APC"), ("COAD", "APC")}]
)
sup.add_heading("Table S10d. Within-cancer prediction of TCGA tissue-source site from TITAN representations", 1)
sup.add_paragraph("Sites with fewer than 10 patients were excluded only from this dedicated confounding analysis. Macro balanced accuracy is the mean per-site recall; its chance reference is 1/k for k analysed sites. Good site prediction demonstrates that the representation carries submitting-site information, but tissue-source site is not a complete scanner, laboratory or staining-batch identifier.")
add_table(
    sup,
    ["Cancer", "Eligible", "Analysed patients", "Analysed sites", "Macro BA, mean (SD)", "Chance 1/k", "Chance-normalised BA", "Reason if ineligible"],
    [(
        r.get("tumor_type"), r.get("eligible"), r.get("analysed_patients"),
        r.get("analysed_sites"),
        f'{fnum(r.get("macro_balanced_accuracy_mean"))} ({fnum(r.get("macro_balanced_accuracy_sd"))})',
        fnum(r.get("chance_macro_balanced_accuracy")),
        fnum(r.get("normalized_macro_balanced_accuracy")), r.get("error")
    ) for r in site_predictability]
)
sup.add_heading("Table S10e. Metadata-stratified representative PLS–ridge benchmark", 1)
sup.add_paragraph("All 2,073 eligible tests entered the sampling frame. A fixed salted SHA-256 rank selected one endpoint per non-empty outcome-family × empirical sample-size-tercile cell; binary endpoints were additionally stratified by empirical minority-class-fraction tercile. Selection did not use PLS performance and yielded 12 continuous and 35 binary targets. Both algorithms used identical patients, outer folds and inner folds. Continuous hyperparameters maximised the same pooled inner out-of-fold Q²; binary operating thresholds maximised the same inner out-of-fold balanced accuracy. For repeat r, d_r=M_r(ridge)−M_r(PLS), with Δ=(1/5)Σ_r d_r. Each of 2,000 paired patient-resampling replicates retained both methods and all five predictions. The interval does not repeat target sampling, generate partitions or refit models, and it does not represent external transportability. This is representative rather than atlas-wide. Positive differences favour ridge.")
sup.add_paragraph("Panel A. Sampling-stratum summary for the primary metric.", style="Caption")
add_table(
    sup,
    ["Type", "Stratifier", "Stratum", "Models", "Median Δ", "IQR", "PLS/ridge/uncertain"],
    [(
        benchmark_type_label(r.get("outcome_type")), r.get("stratifier"),
        benchmark_family_label(r.get("stratum")), r.get("models"),
        fnum(r.get("median_delta_ridge_minus_pls")),
        f'{fnum(r.get("q1_delta"))} to {fnum(r.get("q3_delta"))}',
        f'{r.get("pls_better")}/{r.get("ridge_better")}/{r.get("difference_uncertain")}'
    ) for r in ridge_stratified]
)
sup.add_paragraph("Panel B. Target-level primary metric: AUROC for binary outcomes and Q² for continuous outcomes.", style="Caption")
add_table(sup, ["Type/family", "Size/imbalance", "Cancer–endpoint", "Metric", "PLS", "Ridge", "Δ Ridge−PLS", "95% paired patient interval", "Result"],
          [(f'{benchmark_type_label(r.get("outcome_type"))}/{benchmark_family_label(r.get("family"))}',
            f'{r.get("size_stratum")}/{r.get("imbalance_stratum")}',
            f'{r.get("tumor_type")}–{r.get("endpoint")}',
            benchmark_metric_label(r.get("primary_metric")), fnum(r.get("pls_primary_mean")),
            fnum(r.get("ridge_primary_mean")), fnum(r.get("delta_ridge_minus_pls")),
            f'{fnum(r.get("delta_ci_low"))} to {fnum(r.get("delta_ci_high"))}',
            benchmark_interpretation(r.get("selected_method"))) for r in ridge_comparison])
sup.add_paragraph("Panel C. Target-level secondary metric; binary balanced accuracy uses operating thresholds selected identically for both methods.", style="Caption")
add_table(sup, ["Type/family", "Size/imbalance", "Cancer–endpoint", "Metric", "PLS", "Ridge", "Δ Ridge−PLS", "95% paired patient interval", "Result"],
          [(f'{benchmark_type_label(r.get("outcome_type"))}/{benchmark_family_label(r.get("family"))}',
            f'{r.get("size_stratum")}/{r.get("imbalance_stratum")}',
            f'{r.get("tumor_type")}–{r.get("endpoint")}',
            benchmark_metric_label(r.get("secondary_metric")), fnum(r.get("pls_secondary_mean")),
            fnum(r.get("ridge_secondary_mean")), fnum(r.get("delta_secondary_ridge_minus_pls")),
            f'{fnum(r.get("delta_secondary_ci_low"))} to {fnum(r.get("delta_secondary_ci_high"))}',
            ("ridge" if float(r.get("delta_secondary_ci_low")) > 0 else
             "PLS" if float(r.get("delta_secondary_ci_high")) < 0 else
             "uncertain")) for r in ridge_comparison])
sup.add_heading("Table S10f. Permutation precision and atlas-wide multiplicity sensitivity", 1)
sup.add_paragraph("Panel A. Candidate retention under increasingly broad Benjamini–Hochberg sensitivities. The primary candidate definition is within cancer and endpoint family. The atlas-wide column applies one correction to every eligible continuous and binary cancer–endpoint test; it is a sensitivity analysis rather than a replacement of the prespecified local question.", style="Caption")
add_table(
    sup,
    ["Outcome", "Eligible", "Screening-threshold eligible", "Local candidates", "Across-cancer family", "Outcome-wide", "Atlas-wide"],
    [(
        r.get("outcome_type"), r.get("eligible_tests"), r.get("effect_eligible"),
        r.get("within_cancer_family_candidates"), r.get("across_cancer_family_pass"),
        r.get("outcome_wide_pass"), r.get("atlas_wide_pass")
    ) for r in multiplicity_summary]
)
sup.add_paragraph("Panel B. Targeted high-resolution refinement. All models reuse the saved first 999 permutations and continue the same deterministic permutation-index sequence to 9,999. Every permutation repeats training-fold centering, inner component selection, outer refitting and held-out prediction. Refined p-values and exact Monte Carlo intervals are precision sensitivities and were not substituted into the primary FDR screen.", style="Caption")
add_table(
    sup,
    ["Type", "Cancer", "Endpoint", "b/999", "p (999)", "b/9,999", "p (9,999)", "95% MC interval"],
    [(
        benchmark_type_label(r.get("outcome_type")), r.get("tumor_type"),
        r.get("endpoint"), r.get("primary_exceedances_999"),
        fnum(r.get("primary_p_999"), 4), r.get("refined_exceedances_9999"),
        fnum(r.get("refined_p_9999"), 4),
        f'{fnum(r.get("refined_mc_lower_95"), 6)} to {fnum(r.get("refined_mc_upper_95"), 6)}'
    ) for r in targeted_permutation]
)
# The table helper appends an empty paragraph. When this table exactly fills the
# portrait page, that paragraph can spill to a blank page before the landscape
# section break in LibreOffice. Remove only that known trailing spacer.
if sup.paragraphs and not sup.paragraphs[-1].text.strip():
    spacer = sup.paragraphs[-1]._element
    spacer.getparent().remove(spacer)
add_landscape_section(sup)
sup.add_heading("Table S10g. Binary class-size sensitivity and development reliability", 1)
sup.add_paragraph("Panel A. Eligibility and screen-positive retention. The 20-per-class row is the prespecified inclusive atlas screen; the 50-per-class row defines standard internal evidence and default inference eligibility.", style="Caption")
add_table(
    sup,
    ["Minimum/class", "Eligible binary pairs", "Screen-positive models", "Eligible retention", "Candidate retention", "Interpretation"],
    [(
        r.get("minimum_per_class"), r.get("eligible_binary_targets"),
        r.get("screen_positive_binary_models"),
        f'{fnum(r.get("eligible_target_retention_percent"), 1)}%',
        f'{fnum(r.get("screen_positive_retention_percent"), 1)}%',
        r.get("interpretation"),
    ) for r in binary_class_sensitivity],
    [2.2, 3.0, 3.2, 2.7, 2.7, 9.0],
)
sup.add_paragraph("Panel B. All 17 exploratory/limited-evidence screen-positive binary models. PR-AUC is average precision; PPV and NPV are conditional on the observed endpoint prevalence in TCGA.", style="Caption")
add_table(
    sup,
    ["Cancer", "Endpoint", "+/−", "BA", "AUROC", "PR-AUC", "PPV", "NPV", "Min outer +/−", "Evidence/default"],
    [(
        r.get("tumor_type"), r.get("endpoint"), f'{r.get("positive")}/{r.get("negative")}',
        fnum(r.get("repeated_balanced_accuracy_mean")), fnum(r.get("repeated_auc_mean")),
        fnum(r.get("repeated_pr_auc_mean")), fnum(r.get("ppv_tcga_prevalence_mean")),
        fnum(r.get("npv_tcga_prevalence_mean")),
        f'{r.get("minimum_outer_test_positive")}/{r.get("minimum_outer_test_negative")}',
        "limited; excluded",
    ) for r in binary_limited_reliability],
    [1.6, 5.0, 1.7, 1.5, 1.7, 1.7, 1.5, 1.5, 2.4, 2.7],
)
sup.add_paragraph("Panel C. Inner-fold composition, component selection and repeat stability for the limited-evidence models. Score ρ is the mean pairwise Spearman correlation of standardized held-out LDA scores; call agreement is shown descriptively and can be inflated by imbalance.", style="Caption")
add_table(
    sup,
    ["Cancer", "Endpoint", "Min inner train +/−", "Min inner validation +/−", "Components median (IQR; range)", "Score ρ", "Call agreement"],
    [(
        r.get("tumor_type"), r.get("endpoint"),
        f'{r.get("minimum_inner_training_positive")}/{r.get("minimum_inner_training_negative")}',
        f'{r.get("minimum_inner_validation_positive")}/{r.get("minimum_inner_validation_negative")}',
        f'{fnum(r.get("selected_components_median"), 1)} ({fnum(r.get("selected_components_q1"), 1)}–{fnum(r.get("selected_components_q3"), 1)}; {r.get("selected_components_minimum")}–{r.get("selected_components_maximum")})',
        fnum(r.get("repeat_score_spearman_mean")),
        fnum(r.get("repeat_class_agreement_mean")),
    ) for r in binary_limited_reliability],
    [1.8, 5.5, 3.2, 3.4, 5.0, 2.0, 2.4],
)
add_portrait_section(sup)
add_figure(sup, "FigureS3_binary_class_reliability.png", "Figure S3. Binary class-size reliability. Panel A shows fixed-test-fold learning curves for all 17 exploratory/limited-evidence models at 50%, 75% and 100% of each outer-training set. Panel B shows selected components across the 25 repeated outer fits per model. Panel C compares repeat score correlation and class-call agreement; high agreement under imbalance should not be interpreted without the continuous-score stability and class-specific metrics.")
sup.add_heading("Table S11. Prior histology-based molecular and derived immune-feature prediction landscape", 1)
add_table(sup, ["Study", "Year", "Scope", "Endpoints", "Development cohort", "External validation", "Reported performance", "DOI"],
          [(r.get("study"), r.get("year"), r.get("scope"), r.get("endpoints"),
            r.get("development_cohort"), r.get("external_validation"),
            r.get("reported_performance"), r.get("doi")) for r in literature_landscape])
sup.add_heading("Table S12. TRIPOD+AI reporting map", 1)
sup.add_paragraph("Items are mapped to the revised manuscript and repository using the official TRIPOD+AI checklist (version 11 January 2024). Pending entries require author or institutional information and are not statistical-analysis omissions.")
add_table(sup, ["Item", "Topic", "Reported location", "Status"],
          [(r.get("item"), r.get("topic"), r.get("reported_location"), r.get("status"))
           for r in tripod_map])
add_landscape_section(sup)
sup.add_heading("Table S13. Endpoint provenance, derivation and assay equivalence", 1)
sup.add_paragraph("Panel A. High-level label-generation classes. 'Definitions' counts unique outcome type–family–endpoint combinations; tests are cancer-specific. The full 2,073-row target-level dictionary is the authoritative table and reports all requested fields for every target.", style="Caption")
add_table(
    sup,
    ["Measurement class", "Cancer–endpoint tests", "Unique definitions", "Same-H&E-modality tests"],
    endpoint_class_summary,
    [8.0, 4.0, 4.0, 4.5],
)
sup.add_paragraph("Panel B. Compact derivation-group view. Expected measurement error, biological interpretation, transformation, missingness and source reference remain target-specific in endpoint_dictionary.csv and endpoint_definition_dictionary.csv.", style="Caption")
add_table(
    sup,
    ["Definition group", "Measurement class", "Tests/definitions", "Source modality", "Direct/inferred", "Derivation and source scale", "Assay-equivalence caveat"],
    [(
        group, measurement_class, f"{tests}/{definitions}", modality,
        directness, f"{derivation}; scale: {scale}", caveat
    ) for group, measurement_class, tests, definitions, modality, directness,
          derivation, scale, caveat in definition_group_summary],
    [4.0, 4.2, 2.2, 5.0, 4.6, 8.2, 6.0],
)
sup.add_heading("Table S14. Qualitative high/low prediction anchors and within-cancer TITAN neighbours", 1)
sup.add_paragraph("Panel A. Five representative modality classes. Prediction values retain their original analysed units. TIL Regional Fraction is a same-H&E-modality concordance example; the remaining rows are descriptive cross-modal contexts.", style="Caption")
add_table(
    sup,
    ["Cancer–endpoint", "Target class", "High patient: observed/predicted", "Low patient: observed/predicted", "Interpretation"],
    [(
        f'{r.get("tumor_type")}–{r.get("endpoint")}', r.get("context_class"),
        f'{r.get("high_anchor")}: {fnum(r.get("high_observed"), 4)}/{fnum(r.get("high_prediction"), 4)}',
        f'{r.get("low_anchor")}: {fnum(r.get("low_observed"), 4)}/{fnum(r.get("low_prediction"), 4)}',
        r.get("interpretation"),
    ) for r in morphology_context_summary],
    [5.0, 5.0, 5.2, 5.2, 10.5],
)
sup.add_paragraph("Panel B. Anchor and nearest-neighbour records. H/L denote high/low prediction strata. Cosine similarity is to the anchor; anchor values equal 1. The complete CSV additionally includes report excerpts, slide-to-patient-mean distance, patient slide count and explicit provenance/interpretability warnings.", style="Caption")
add_table(
    sup,
    ["Model", "Stratum/role", "Patient", "Representative slide", "Observed", "Prediction", "Prediction rank", "Cosine", "Generated-report terms"],
    [(
        f'{r.get("tumor_type")}–{r.get("endpoint")}',
        f'{r.get("stratum")}/{r.get("role")}', r.get("patient"),
        r.get("representative_slide"), fnum(r.get("observed"), 4),
        fnum(r.get("predicted_value"), 4), fnum(r.get("predicted_rank"), 3),
        fnum(r.get("embedding_cosine_similarity_to_anchor"), 3),
        r.get("report_terms"),
    ) for r in morphology_context],
    [4.2, 3.2, 3.0, 8.0, 2.2, 2.4, 2.6, 2.0, 7.0],
)
add_figure(sup, "FigureS4_morphology_context.png", "Figure S4. Qualitative morphology context for five representative label modalities. Grey points are all repeated out-of-fold patient predictions. H/L are concordant high/low anchors and triangles are their closest within-cancer patient-level mean-embedding neighbours. The BLCA TIL Regional Fraction target was itself inferred from H&E. Neighbour retrieval is not patch-level attribution; exact patient/slide identifiers, original predictions and automated-report provenance are in Table S14.", width=8.8)
add_portrait_section(sup)
sup.add_heading("Secondary PLS1–PLS2 results", 1)
if pls2:
    text = []
    for r in pls2:
        text.append(f'{r["block"].replace("_", " ")}: mean cancer-level ΔQ² {fnum(r["mean_cancer_delta"])} (95% interval {fnum(r["ci_low"])} to {fnum(r["ci_high"])})')
    sup.add_paragraph("; ".join(text) + ". The intervals describe variation across cancers; interpretation is based on effect magnitude rather than endpoint win counts.")
add_figure(sup, "Figure6a_pls1_vs_pls2_targets.png", "Figure S1. Matched target-level PLS1 and PLS2 Q² on identical held-out patients.")
add_figure(sup, "Figure6b_pls1_vs_pls2_cancers.png", "Figure S2. Cancer-level mean Q² change for joint PLS2 relative to response-by-response PLS1.")
sup.add_heading("Machine-readable additional files", 1)
for name in ["continuous_screen.csv", "binary_screen.csv", "endpoint_dictionary.csv", "endpoint_definition_dictionary.csv", "endpoint_dictionary_summary.csv", "morphology_context_examples.csv", "morphology_context_model_summary.csv", "permutation_monte_carlo_uncertainty.csv", "targeted_permutation_refinement.csv", "targeted_permutation_refinement_targets.csv", "multiplicity_sensitivity_by_endpoint.csv", "multiplicity_sensitivity_summary.csv", "screen_positive_performance_summary.csv", "highlighted_model_performance.csv", "continuous_repeated_nested_cv.csv", "binary_repeated_nested_cv.csv", "continuous_repeated_oof_predictions.csv.gz", "binary_repeated_oof_predictions.csv.gz", "binary_minimum_class_sensitivity.csv", "binary_class_reliability_summary.csv", "binary_limited_evidence_models.csv", "binary_reliability_by_repeat.csv", "binary_outer_fold_class_counts.csv", "binary_selected_components_by_fold.csv", "binary_selected_component_distribution.csv", "binary_prediction_stability.csv", "binary_limited_evidence_learning_curve_repeats.csv", "binary_limited_evidence_learning_curve_folds.csv", "binary_limited_evidence_learning_curve_summary.csv", "continuous_site_grouped_sensitivity.csv", "binary_site_grouped_sensitivity.csv", "site_grouped_retention_summary.csv", "site_grouped_models_below_effect_threshold.csv", "site_grouped_outer_fold_composition.csv", "site_grouped_fold_composition_summary.csv", "tissue_source_site_predictability_repeats.csv", "tissue_source_site_predictability_summary.csv", "continuous_slide_pooling_sensitivity.csv", "binary_slide_pooling_sensitivity.csv", "pls1_vs_pls2_inflammation.csv", "pls_vs_ridge_representative_sampling_frame.csv", "pls_vs_ridge_representative_jobs.csv", "pls_vs_ridge_representative_repeated_nested_cv.csv", "binary_symmetric_pls_ridge_representative_repeated_nested_cv.csv", "pls_vs_ridge_representative_paired_repeat_metrics.csv", "pls_vs_ridge_representative_matched_oof_predictions.csv.gz", "pls_vs_ridge_representative_models.csv", "pls_vs_ridge_representative_summary.csv", "pls_vs_ridge_representative_stratified_summary.csv", "prior_mutation_literature_crosswalk.csv", "prior_mutation_accuracy_comparison.csv", "supported_mutation_literature_audit.csv", "pan_cancer_benchmark_comparison.csv", "external_validation_locked_targets.csv", "slide_report_coverage_audit.csv", "patient_slide_multiplicity_by_cancer.csv", "participant_characteristics_by_cancer.csv", "tcga_cdr_match_audit.csv", "molecular_source_coverage_audit.csv", "mutation_coverage_audit.csv", "mutation_target_eligibility_audit.csv", "mutation_variant_classification_audit.csv", "source_manifest.csv", "software_manifest.csv", "models/model_registry.csv"]:
    p = sup.add_paragraph(name, style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
sup.save(OUT / "supplementary_material_JTM.docx")


# Point-by-point response
resp = setup(Document(), "Response to reviewer — TITAN benchmark")
resp.add_heading("Response to reviewer", 0)
resp.add_paragraph("Manuscript: " + MANUSCRIPT_TITLE)
resp.add_paragraph("We thank the reviewer for identifying validation and reproducibility as the principal issues. The analysis has been rebuilt from the original files with patient-first slide aggregation and primary-tumour molecular matching.")

responses = [
    ("1. The novelty relative to existing pan-cancer studies must be defined much more sharply",
     "Agreed. We changed the title and abstract to frame the work as a systematic patient-level benchmark and reusable model resource, not a discovery atlas or a first pan-cancer molecular screen. The Background now quantifies the closest precedents: Fu et al. analysed 17,355 slides across 28 cancers; Kather et al. used more than 5,000 patients across 14 cancers; Saldanha et al. externally tested mutation models in seven matched TCGA/CPTAC cancers; and Arslan et al. trained 12,093 models for 4,031 biomarkers in 8,890 TCGA patients across the same 32 cancers. To keep the research article focused, the comparative review-style table has been removed from the main manuscript; the literature context is stated concisely in prose, while the detailed audit remains in Supplementary Table S11 and data/reference/pan_cancer_benchmark_comparison.csv. The abstract and Discussion state that 38/41 screen-positive cancer–gene pairs had prior statistical support. We make no general mutation-target novelty claim; the distinctive contribution is the fixed pretrained TITAN representation, deterministic pre-outcome patient aggregation, nested patient-level validation, complete negative and ineligible outputs, target-level site sensitivity, continuous immune/genomic-context modelling and compact fitted-model distribution."),
    ("Audit comment: continuous repeated cross-validation returned negative Q² and near-zero correlation",
     "Confirmed and corrected. fastPLS returns continuous predictions as an n-by-1-by-1 array. The previous repeated-validation helper used Ypred[[1]], which extracted one scalar and silently recycled it across every patient in the held-out fold. The revised helper drops only singleton dimensions, verifies that the prediction length equals the held-out patient count, and then assigns one prediction per patient. A targeted COAD TIL Regional Fraction check agreed with pls.double.cv (corrected Q² 0.430 and Spearman 0.701 versus 0.377 and 0.655 with the independent routine). We invalidated the affected analysis fingerprint and regenerated all repeated continuous predictions, summaries, uncertainty intervals, reference distributions, figures and reports. Across the 219 screen-positive continuous models, corrected mean repeated-CV Q² values are now positive, with median Q² 0.300 and median Spearman correlation 0.553."),
    ("Audit comment: continuous radar reference positions inherited the broken predictions",
     "Confirmed. The fitted full-cohort predictions were not affected, but the TCGA out-of-fold reference distributions used for radar positions were. TITANPred's prediction_reference.rds was rebuilt from the corrected repeated held-out predictions, the package was reinstalled, and the COAD reports and Figure 7 were regenerated. The original model prediction remains printed at every radar corner."),
    ("3. Site sensitivity is a central result, not a supplementary robustness check",
     f"Addressed as a principal analysis. Main Figure 6 now has three panels: all 323 random-fold versus tissue-source-site-grouped results; the largest target-level attenuations, including READ–APC (balanced accuracy 0.862 to 0.495) and COAD–APC (0.793 to 0.543); and the new within-cancer tissue-source-site prediction analysis. Complete highlighted-model grouped metrics and warning status are in Supplementary Table S6a rather than an oversized main-text table. The public analysis registry and access-controlled TITANPred registry contain the grouped metric, delta, number of sites, threshold-retention flag, near-chance flag, validation scope and warning for all 323 models; inference outputs and reports display these fields and prominently warn about site-sensitive models. We report that {ival(site_combined.get('below_threshold_models'))}/{ival(site_combined.get('screen_positive_models'))} models ({fnum(site_combined.get('below_threshold_percent'), 1)}%) fell below the original threshold: 58/219 continuous and 25/104 binary. The fold audit provides all {len(site_fold_details):,} outer-fold compositions with test/training patients, sites, positive and negative cases, site identifiers and seeds; selected APC folds appear in Table S10c. We confirmed from both the implementation and deterministic fold reconstruction that the tissue-source-site constraint was passed to inner component selection as well as outer evaluation, with no site overlap. In the new five-repeat nested PLS–LDA analysis, TITAN predicted tissue-source site in {len(site_predictability_eligible)}/32 evaluable cancers; median multiclass macro balanced accuracy was {fnum(median([r.get('macro_balanced_accuracy_mean') for r in site_predictability_eligible]))}. We consistently call the endpoint sensitivity 'TCGA tissue-source-site-grouped internal validation' and state that tissue-source site is an imperfect proxy; neither retained performance nor site classification proves institutional or scanner-level transportability."),
    ("Audit comment: mutation novelty was overstated",
     f"Confirmed. The exact-code crosswalk missed pooled colorectal studies and did not include newer pan-cancer source data. We expanded the audit using Kather et al., Saldanha et al., Arslan et al. and a clearly labelled ovarian preprint. Of {len(mutation_literature_audit)} screen-positive mutation pairs, {len(prior_supported_mutations)} were previously supported, {len(prior_evaluated_not_supported)} had been evaluated without statistical support, and only THYM–GTF2I was not identified in the reviewed predictive-model literature. All 'atlas-nominated' novelty language was removed. The revised Table S10a records evidence class, cancer scope, prior metric, source and note for every pair."),
    ("Audit comment: PLS lacked a simple exportable baseline",
     f"Superseded by a stronger representative benchmark selected without PLS performance. All {len(continuous)+len(binary):,} eligible tests entered a metadata-only sampling frame, and a fixed salted SHA-256 rank selected one target from every non-empty outcome-family × sample-size-tercile cell, with an additional minority-class-fraction tercile for binary outcomes. The resulting {len(ridge_comparison)} targets comprise {len(ridge_continuous)} continuous and {len(ridge_binary)} binary endpoints, of which only {len(ridge_screen_positive)} were screen-positive. Median ridge-minus-PLS differences were {fnum(continuous_baseline_summary.get('median_delta_ridge_minus_pls'))} Q² and {fnum(binary_baseline_summary.get('median_delta_ridge_minus_pls'))} AUROC. We make no PLS-superiority claim; Table S10e and machine-readable files report the sampling frame, hashes, jobs, folds, predictions and paired intervals."),
    ("4. The binary modelling and ridge comparison use asymmetric decision rules",
     f"Confirmed and corrected. The representative benchmark contains {len(ridge_binary)} binary comparisons across five repeats. PLS–LDA and ridge receive identical outer and inner folds; each operating threshold is selected from that method's inner out-of-fold score by the identical balanced-accuracy rule. Threshold-independent AUROC is primary and symmetrically thresholded balanced accuracy secondary. The median ridge-minus-PLS AUROC difference was {fnum(binary_baseline_summary.get('median_delta_ridge_minus_pls'))} (IQR {fnum(binary_baseline_summary.get('q1_delta'))} to {fnum(binary_baseline_summary.get('q3_delta'))}); {sum(r.get('outcome_type') == 'binary' for r in ridge_better)} targets favoured ridge, none favoured PLS and {sum(r.get('outcome_type') == 'binary' for r in baseline_uncertain)} were uncertain. For balanced accuracy, {len(binary_secondary_ridge)} favoured ridge, {len(binary_secondary_pls)} PLS and {len(binary_secondary_uncertain)} were uncertain. Continuous PLS components and ridge penalties likewise use identical folds and maximise the same pooled inner out-of-fold Q². Complete repeat-level results record fold and tuning symmetry."),
    ("5. The multiplicity framework is thoughtful but remains resolution-limited",
     f"Addressed at the analysis and reporting levels. We now state unambiguously that every patient-label permutation reruns the complete nested modelling process: training-fold centering, inner five-fold selection of 1–10 components, outer-fold refitting and held-out prediction. No scaling parameter, selected component count or outer prediction is reused from the observed-label model. For every completed permutation test we added exact two-sided 95% Clopper–Pearson Monte Carlo bounds for the underlying null exceedance probability. Among completed 999-permutation tests, {len(zero_999)} had zero exceedances; their finite empirical p-value is 0.001 but the interval is 0–{fnum(zero_999_upper, 6)}. Conservatively early-stopped tests retain p=1 and are explicitly marked as censored rather than given a precision interval. Before generating high-resolution results, we locked TGCT TGF-beta Response, THYM Th17 Cells, THCA Lymphocyte Infiltration Signature Score, BRCA Wound Healing, COAD APC, THYM GTF2I, COAD strict MSI-H and UCEC any-called-fusion; the registry was recorded in Git commit ac30ccb before the result table existed. We continued their saved deterministic permutation streams from 999 to 9,999 complete-process permutations. {len(targeted_zero)}/{len(targeted_permutation)} had zero exceedances; refined p-values were {targeted_p_summary}. These refined values are reported as a targeted precision sensitivity and were not substituted into the prespecified FDR screen. We also added outcome-wide and single atlas-wide BH sensitivities. Of {ival(combined_multiplicity.get('within_cancer_family_candidates'))} locally controlled candidates, {ival(combined_multiplicity.get('atlas_wide_pass'))} passed one BH correction across all {ival(combined_multiplicity.get('eligible_tests')):,} eligible tests. The manuscript now says explicitly that the aggregate candidate count is not itself a single global 5% FDR result. Figures and tables are ordered by the outcome-appropriate predictive metric, with repeated and site-grouped stability reported alongside; tied permutation q-values are not used for ranking. Full endpoint-level uncertainty, atlas-wide q-values, the locked target registry and 9,999-permutation results are supplied in machine-readable files and Supplementary Table S10f."),
    ("6. The minimum binary class requirement is too permissive for headline and distributed models",
     f"Agreed. We retained the prespecified 20-per-class rule only for an inclusive, fully reported atlas screen, and added a separate ≥50-per-class evidence standard. Of {ival(binary_sensitivity_20.get('eligible_binary_targets'))} eligible binary cancer–endpoint pairs, {ival(binary_sensitivity_50.get('eligible_binary_targets'))} ({fnum(binary_sensitivity_50.get('eligible_target_retention_percent'), 1)}%) met the stricter rule. Of {len(supported_b)} screen-positive binary models, {len(binary_standard)} ({100 * len(binary_standard) / len(supported_b):.1f}%) met it; the remaining {len(binary_limited_reliability)} are now labelled exploratory/limited evidence and excluded from default TITANPred inference. None of the main highlighted binary models is limited evidence. For all 104 candidates we released all {len(binary_fold_counts):,} repeated outer-fold class counts, all {len(binary_component_summary) * 25:,} outer-fit component selections with reconstructed inner-fold class minima, repeat-specific PR-AUC, TCGA-prevalence PPV/NPV, and score/call stability. The smallest outer test fold contained {binary_min_outer_positive} positive and {binary_min_outer_negative} negative patients; within the limited subset the exact inner partitions contained as few as {limited_min_inner_training_positive} positive training and {limited_min_inner_validation_positive} positive validation patients. For all 17 limited models we ran fixed-outer-test learning curves at 50%, 75% and 100% of the training data. Median balanced accuracy was {fnum(binary_learning_50.get('balanced_accuracy'))}, {fnum(binary_learning_75.get('balanced_accuracy'))} and {fnum(binary_learning_100.get('balanced_accuracy'))}; median AUROC was {fnum(binary_learning_50.get('auc'))}, {fnum(binary_learning_75.get('auc'))} and {fnum(binary_learning_100.get('auc'))}. Supplementary Table S10g and Figure S3 show endpoint-level counts, PR-AUC, predictive values, components, learning curves and repeat stability. The registry and inference output expose the evidence tier, default status and reliability metadata; explicit opt-in to a limited model emits a warning. PPV/NPV are labelled cohort-specific and the binary score remains explicitly uncalibrated and not a probability."),
    ("7. Reported uncertainty is conditional and does not capture the entire modelling process",
     f"Agreed. Every main and supplementary performance-table heading now uses '95% SC interval', defined in full as a '95% selection-conditioned patient-resampling interval for repeated out-of-fold predictions'. For highlighted atlas models, 1,000 bootstrap replicates retain five existing held-out prediction sets; the Methods enumerate the omitted screening, repartitioning, tuning, refitting, winner's-curse and external-site components. The PLS–ridge comparison uses 2,000 paired patient-resampling replicates and the exact d_r and Δ construction requested. Its interval is now conditional on the metadata-stratified benchmark sample—not PLS-based highlighting—and does not regenerate partitions or refit models. Matched patient-level predictions and all {len(ridge_comparison)*5} repeat-specific differences are released. Paired primary-metric intervals favoured ridge for {sum(r.get('outcome_type') == 'binary' for r in ridge_better)} binary and {sum(r.get('outcome_type') == 'continuous' for r in ridge_better)} continuous comparisons, PLS for none."),
    ("8. The justification for retaining PLS as the central model is currently insufficient",
     f"Agreed with the premise: the revision no longer describes PLS as the best or uniquely portable method. We replaced the PLS-highlighted comparison with a {len(ridge_comparison)}-target representative benchmark selected solely by outcome family, empirical sample-size tercile, binary minority-class-fraction tercile and a fixed salted hash. PLS and ridge use identical folds and identical tuning objectives. Median ridge-minus-PLS differences were {fnum(binary_baseline_summary.get('median_delta_ridge_minus_pls'))} AUROC and {fnum(continuous_baseline_summary.get('median_delta_ridge_minus_pls'))} Q²; paired intervals favoured ridge for {sum(r.get('outcome_type') == 'binary' for r in ridge_better)}/{len(ridge_binary)} binary and {sum(r.get('outcome_type') == 'continuous' for r in ridge_better)}/{len(ridge_continuous)} continuous targets, PLS for none. We therefore retain PLS only as the prespecified reference analysis and label the fitted PLS resource accordingly; portability is explicitly stated to apply to ridge as well. PLS2 remains secondary because it covers only coherent continuous inflammatory panels, uses complete-case intersections and a fixed ordered response block, shares latent structure across responses, and cannot replace individual binary mutation/pathway/MSI/aneuploidy/fusion models. Its mean cancer-level ΔQ² values of 0.066, 0.017 and 0.056 support future locked multiresponse immune modelling, but not replacement of the endpoint-specific resource without external validation and a matched multiresponse baseline."),
    ("9. Derived immune phenotypes, endpoint provenance and morphological context",
     f"Agreed and addressed at the analysis, manuscript and supplementary-data levels. We no longer use 'immune measurements' as an undifferentiated label. Every one of the {len(endpoint_dictionary):,} eligible cancer–endpoint tests is classified as a directly observed genomic alteration, sequencing-derived continuous burden, computationally inferred immune-cell fraction, transcriptomic signature, pathology-associated quantity or composite genomic-context score. The new endpoint_dictionary.csv records source modality, direct/inferred status, derivation algorithm, original scale, analysis transformation, analysed and missing patients, expected measurement error, biological interpretation, assay-equivalence caveat and source reference for every target; endpoint_definition_dictionary.csv provides {len(endpoint_definitions)} unique definitions. The Methods and Discussion now distinguish RNA-seq CIBERSORT deconvolution, methylation-derived leukocyte fraction, RNA signatures, repertoire reconstruction, predicted neoantigens, copy-number/purity scores and directly called alterations. TIL Regional Fraction is explicitly labelled a Saltz H&E-derived same-modality concordance target. We state that predicting an inferred score is not equivalent to flow cytometry, IHC, a direct immune-cell count or a clinically certified biomarker. Supplementary Table S13 summarises the classes and derivations while designating the full 2,073-row CSV as the authoritative target-level dictionary. For morphology, we added five representative high/low repeated-OOF anchor analyses spanning BLCA TIL Regional Fraction, TGCT Macrophages M2, BRCA Wound Healing, UCEC Aneuploidy score and THYM GTF2I. Each anchor is paired with its nearest within-cancer patient-level mean TITAN neighbour, with exact patient and representative-slide IDs, observed and original predicted values, prediction ranks, cosine similarity and TITAN-generated report context in Table S14 and morphology_context_examples.csv; Figure S4 shows all held-out patients behind the selected points. Because only global pooled embeddings and automated same-slide reports were available, we explicitly describe this as qualitative neighbour retrieval—not patch relevance, causal morphology or blinded pathologist review—and identify tile-level attribution plus independent pathology review as required future work."),
    ("10. The COAD single-sample demonstration risks overinterpretation",
     "Partly agreed. At the authors' request, the two COAD participants and radar comparison remain in the main manuscript solely as a software-visualization example. We removed the clinical narrative and made the post hoc construction explicit in both Results and the Figure 7 caption: clinical-text and profile-saturation restrictions were followed by maximisation of Euclidean separation across the continuous TCGA OOF prediction-percentile profile, intentionally enhancing contrast. We removed the claim that the cases are clinically comparable and state that TCGA-AA-A01F was reported as pN1, whereas the available slide summary for TCGA-AA-3972 did not state a nodal category. All treatment, response, recurrence, follow-up and survival narrative has been deleted from the manuscript, generated case table, separate case reports and TITANPred report interface because these variables were neither inputs nor validation outcomes. Figure 7 is labelled an interface illustration—not calibrated output, external validation, representative sampling, prognosis or treatment-response evidence. The package tutorial additionally retains an explicitly non-patient synthetic smoke-test vector. The report has a prominent banner stating 'UNVALIDATED RESEARCH OUTPUT', 'TCGA OOF score rank (not probability)', 'No independent external validation' and 'Not for diagnosis, treatment selection, or clinical risk estimation'. Every binary figure and table places '(not probability)' immediately beside the score-rank label; inference rows expose rank_interpretation, rank_is_probability=FALSE and calibration_status='uncalibrated; no probability estimate'."),
    ("Audit comment: the Introduction did not state the gap",
     "Addressed. The Background now states explicitly that breadth is not the gap after the large pan-cancer studies by Fu, Kather, Saldanha and especially Arslan. It defines the narrower gap as a reproducible patient-level benchmark using one fixed pretrained representation, deterministic slide aggregation, one nested validation framework, complete negative and ineligible reporting, target-level site sensitivity and compact fitted research models."),
    ("Audit comment: binary percentiles could be mistaken for probabilities and small classes were not visible",
     "Addressed by relabelling rather than adding an unevaluated post-hoc calibration layer. Binary displays use 'TCGA out-of-fold score rank (not probability)', retain the raw uncalibrated LDA score and class call, and show training class counts and prevalence, balanced accuracy, AUROC, PR-AUC, TCGA-prevalence PPV/NPV, fold minima, component variability and repeat stability. Models below 50 patients in either class are no longer merely warned: they are excluded from default inference, marked exploratory/limited evidence throughout, and require explicit opt-in that emits a warning. Proper probability calibration remains future work requiring independently evaluated or fully nested calibration."),
    ("1. Across-cancer multiplicity and permutation resolution",
     f"Addressed. Every endpoint meeting the prespecified screening-statistic threshold is refined toward 999 complete-process patient-label permutations with a conservative stopping boundary of 49 exceedances, minimum completed-test resolution 0.001 and exact attempted counts retained. The prespecified within-cancer/family q-value, across-cancer family q-value, outcome-wide q-value and single atlas-wide q-value are now reported in separate machine-readable fields. Exact Monte Carlo intervals and the targeted 9,999-permutation refinement make the remaining resolution limitation explicit. The abstract, figures and Discussion distinguish locally controlled candidates from the {ival(combined_multiplicity.get('atlas_wide_pass'))}/{ival(combined_multiplicity.get('within_cancer_family_candidates'))} that also pass the single atlas-wide sensitivity; non-passage is not treated as evidence of absence."),
    ("2. External validation is the principal unresolved limitation",
     "We agree that this is the central unresolved limitation. The official TITAN release provided the TCGA embeddings used here, but our audit found no compatible non-TCGA precomputed TITAN representation with the required outcomes; no independent patient entered analysis. CPTAC-UCEC is a realistic future cohort because its official TCIA collection reports 250 subjects, 887 pathology slides (approximately 154 GB) and linked molecular resources, but de novo TITAN extraction and exact endpoint harmonisation are required and were not performed. We therefore adopted the reviewer's computational-resource fallback throughout: 'deployment', 'patient molecular profile' and 'prediction report' language was removed; Figure 7 is retained only as a post hoc research-software visualization with an explicit no-validation caption; and all numerical outputs are labelled internally derived TCGA estimates. The abstract, intended use, Discussion and Conclusions state that there is no external performance evidence or basis for clinical interpretation. To prevent future target or threshold cherry-picking, we prospectively locked three UCEC artifacts before inspecting any external features or outcomes: TP53 mutation, genome doubling and continuous aneuploidy score. Their exact SHA-256 hashes, endpoint-compatibility rules and metrics are in Supplementary Table S6c, docs/EXTERNAL_VALIDATION_PROTOCOL.md and data/reference/external_validation_locked_targets.csv rather than the main Results. The future protocol requires exact TITAN extraction, identical patient pooling, no refitting, recalibration or threshold adjustment, and reporting of every target including failures and non-evaluable endpoints. We explicitly state that locking a protocol is not external validation."),
    ("Additional site-sensitivity coverage across endpoint families",
     f"Tissue-source-site-grouped validation is attempted for every within-cancer screen-positive continuous and binary endpoint, not only mutations. It was feasible for {sum(r.get('feasible')=='TRUE' for r in site_c)+sum(r.get('feasible')=='TRUE' for r in site_b)} models. Figure 6, Table S10b and the complete machine-readable tables report target-level attenuation across immune, genomic-context, MSI, aneuploidy, fusion, pathway and mutation families."),
    ("4. Reproducibility materials must be deposited before submission",
     f"The analysis repository ({REPO}) remains public and contains the analysis plan, source URLs/checksums, software commit, eligibility tables, code, out-of-fold outputs, literature crosswalk, model registry and locked external-evaluation protocol. At the authors' request, the TITANPred repository ({MODEL_REPO}) is now private. The manuscript no longer claims that the package or 323 fitted objects are publicly downloadable. A public model release, if permitted, and a persistent DOI remain future release decisions; reviewers can be granted controlled access when authorised."),
    ("5. Distinguish cancer genes from functional driver alleles",
     "Addressed throughout. Mutation targets are described as qualifying protein-altering PASS mutations in tissue-specific consensus cancer genes; the manuscript explicitly states that not every allele is functionally validated."),
    ("6. Preserve effect-size-first interpretation of PLS2",
     "Addressed. The PLS1–PLS2 comparison is now confined to the Supplementary Methods, results and Figures S1–S2. It remains restricted to coherent inflammatory blocks, uses identical patients and folds, and is interpreted by cancer-level ΔQ² with bootstrap intervals rather than win counts. Binary molecular endpoints retain one-at-a-time PLS–LDA as the primary analysis."),
    ("7. Complete submission-specific fields",
     "Partly outstanding. Aamilah Ismail and Martin Ocharo are listed as shared co-first authors, with Martin Ocharo second in the author order and assigned to affiliations 1 and 2. Brendon Price is included in the middle of the author list with the Division of Anatomical Pathology, University of Cape Town and National Health Laboratory Service affiliation. Silvano Piazza and Dinesh Gupta are listed immediately before Stefano Cacciatore. Silvano Piazza has dual affiliations with the ICGEB Computational Biology Group in Trieste and the Bioinformatics Facility, CIBIO, University of Trento; Dinesh Gupta has the ICGEB New Delhi affiliation. The remaining supplied author names, affiliations, available email addresses and corresponding-author details have been entered. Email addresses for Martin Ocharo, Brendon Price, Ekene Emmanuel Nweke, Silvano Piazza and Dinesh Gupta were not supplied. Funding, competing interests and contribution statements still require author confirmation and remain visibly marked where applicable. No scientific values are placeholder text."),
    ("8. Presentation and algorithm-comparison claims",
     "Addressed. High-resolution figures and machine-readable tables accompany the Word documents. Runtime and speed claims were removed. A focused same-fold ridge benchmark was added because it directly tests whether PLS materially outperforms a simpler portable linear model; it is reported with effect magnitudes, uncertainty and selection-conditioning caveats."),
    ("Additional change: distinguish replication from predictors not identified in the reviewed literature",
     "The Discussion uses an expanded primary-study audit, maps pooled colorectal evidence to COAD and READ, gives prior AUROC and current balanced accuracy side by side without treating the metrics as directly comparable, and distinguishes previously supported, previously evaluated-but-not-supported, and not-identified categories. Only THYM–GTF2I was not identified in the reviewed predictive-model literature, and even that result is not claimed as biological or definitive bibliographic novelty. The broader landscape audit covers MSI, continuous biomarkers, expression, gene fusion, HRD and tumour-microenvironment prediction, including external-validation examples."),
    ("Additional change: clarify model portability without training-data release",
     "The Background, Methods, Discussion and Conclusions explain that fitted PLS and PLS–LDA models contain compact learned transformations and coefficients and can be applied without distributing patient-level training embeddings or outcomes. Licensing, privacy, governance and external-validation limitations are retained."),
    ("Additional change: multiple slides and molecular specimen matching",
     f"The primary predictor is now the feature-wise mean of every eligible diagnostic slide per patient; {n_multi:,} multi-slide patients are retained in a single validation fold. We also audited every molecular source and excluded non-primary TCGA sample types before aggregation. A matched first-slide sensitivity quantifies dependence on the pooling rule."),
    ("Additional change: molecular missingness, wild-type status and aliquots",
     f"Mutation wild type is now assigned only among {n_mc3_profiled:,} patients matched to an MC3 primary-tumour profile; {n_mc3_missing:,} embedding patients without a profile are excluded. The nine accepted protein-altering MC3 variant classes are explicit and a cancer-level variant-class audit is released. Fusion negatives are defined only within the study sample list. Source-specific missingness, multiple aliquots and aggregation rules are written to audit tables and reported in Methods and Supplementary Material."),
    ("Additional change: TITAN pretraining and TCGA relationship",
     "The manuscript now documents that the published Mass-340K pretraining corpus excluded TCGA and PANDA, while TCGA was used for downstream evaluation in the original TITAN study. We therefore describe this work as a secondary TCGA benchmark and model resource, neither a pretraining-overlap analysis nor independent external validation."),
    ("Additional change: complete performance and prediction examples",
     "Highlighted binary models now report sensitivity, specificity, balanced accuracy and AUROC; highlighted continuous models report Q², RMSE and Spearman correlation, with explicitly labelled selection-conditioned patient-resampling intervals for repeated out-of-fold predictions. Research-use provenance includes feature checksums and ranges, class coding and priors, decision rule, calibration status, external-validation status and intended use. Figure 4 compares held-out predictions directly with observed data. Figure 7 and the two separate PDF attachments (Additional files 2 and 3) are restricted to post hoc software-interface visualization and labelled internally derived, uncalibrated and not externally validated."),
    ("Additional change: rSVD-only PLS decomposition",
     "All primary, permutation, repeated, site-grouped, slide-pooling, PLS1–PLS2 and final-model fits use CPU rSVD with 10 oversampling vectors, two power iterations and fixed seeds. Solver identity and controls are recorded in screening, sensitivity and model-registry metadata and verified against saved-model diagnostics."),
    ("Effect terminology",
     f"Agreed. We replaced 'higher effect' and 'moderate effect' throughout the manuscript, figures, tables, software report and machine-readable performance summary with the neutral labels 'prespecified screening tier A' and 'prespecified screening tier B'. The Methods define tier A as primary q<0.05 with Q²≥0.40 for continuous outcomes or balanced accuracy≥0.70 for binary outcomes, and tier B as primary q<0.05 with Q² from 0.20 to <0.40 or balanced accuracy from 0.60 to <0.70. We state explicitly that these are prioritisation rules, not established clinical or statistical effect categories, and that Q² and 2×balanced accuracy−1 are not interpreted as directly commensurate. The abstract, Results and Discussion also report that {len(global_supported_c)} continuous and {len(global_mutation_b)} cancer–mutation pairs passed the across-cancer family correction and that {ival(combined_multiplicity.get('atlas_wide_pass'))}/{ival(combined_multiplicity.get('within_cancer_family_candidates'))} local candidates passed a single BH correction across all eligible atlas tests."),
    ("Primary versus repeated estimates",
     "Agreed and clarified throughout. We now reserve 'primary screening estimate' for the initial nested-CV value used for candidate selection, permutation testing and tier assignment, and 'repeated-validation estimate' for the arithmetic mean across five additional independently seeded nested-CV partitions. The Methods explain that different folds and rSVD seeds make numerical differences expected and that the repeated estimate describes post-selection stability without replacing the primary screen. THYM–Th17 is stated explicitly as primary screening Q²=0.638 versus five-repeat mean Q²=0.594. Supplementary Table S6a presents both estimates side by side; the abstract, Results, figure captions and Supplementary Tables S6a–S10b label the relevant stage. Machine-readable summaries include explicit primary-screen fields, repeated fields, estimate labels and the repeated-minus-primary difference."),
    ("Manuscript length and table density",
     "Agreed. The main-text comparison table was removed because the manuscript is a research article rather than a review; its essential message remains in concise Background and Discussion prose, while detailed literature comparisons remain in Supplementary Table S11 and the machine-readable crosswalk. The oversized highlighted-model performance table was also removed from the main text, with all metrics retained in Supplementary Table S6a and machine-readable outputs. After further streamlining, the locked-target implementation table was moved to Supplementary Table S6c; the main manuscript now contains only Table 1 (analysis coverage)."),
    ("Report precision–recall metrics",
     f"Agreed. PR-AUC is calculated as non-interpolated average precision from held-out continuous LDA scores in each repeated nested-CV run, and its no-skill reference is explicitly defined as the observed endpoint prevalence. The Abstract now reports PR-AUC for the highlighted binary model. The main Results summarize all {len(binary_reliability)} screen-positive binary models and show the prevalence-dependent contrast directly: the {len(low_prevalence_mutation_fusion)} mutation or fusion endpoints below 20% prevalence had median prevalence {fnum(median(r.get('observed_tcga_prevalence') for r in low_prevalence_mutation_fusion))}, median PR-AUC {fnum(median(r.get('repeated_pr_auc_mean') for r in low_prevalence_mutation_fusion))} and median AUROC {fnum(median(r.get('repeated_auc_mean') for r in low_prevalence_mutation_fusion))}. Endpoint-level PR-AUC, repeat variability, prevalence and class counts remain in Supplementary Tables S6a, S7 and S10g, the model registry and TITANPred output."),
    ("Figure readability",
     "Agreed. Figures 2 and 3 now show the 18 strongest models rather than 30, use larger fonts and points, remove redundant encodings and direct readers to complete supplementary results. Figure 5 is restricted to the 20 cancers with the largest candidate counts and uses larger axis and legend typography. Figure 6 retains all 323 models in the overview but limits the labelled attenuation panel to eight endpoints and the submitting-site panel to the 15 strongest chance-normalised results, with larger fonts throughout. Figure 7 no longer uses two radar plots: its continuous panel is a common-scale dumbbell comparison that permits direct between-case reading while preserving every original prediction in a separate value column; the simplified binary panel remains below it. Exact values, sample IDs and all non-probability and internal-validation warnings are retained."),
    ("Move implementation detail out of the main Results",
     "Agreed. The detailed TITANPred inventory, input schema, pooling behaviour, score-rank construction, checksums, class counts, reliability warnings, report-generation behaviour, synthetic smoke test and locked-artifact metadata have been removed from the main Results. These materials remain in Supplementary Methods, Tables S6b–S6c and S13, the machine-readable registries and the package documentation. The former software-interface subsection is now a short 'Illustrative COAD output' section containing only the purpose and selection caveats needed to interpret Figure 7. The locked UCEC targets are stated in one sentence, with all hashes and protocol mechanics moved to Supplementary Table S6c."
    ),
    ("Additional change: TRIPOD+AI and fairness reporting",
     "Supplementary Table S12 maps every TRIPOD+AI item to the manuscript or repository. Participant characteristics from the TCGA Clinical Data Resource are now reported overall and by cancer; the Methods and Discussion state that no resampling or calibrated probability output was used, no formal power calculation was performed, treatment endpoints were not modelled, and demographic subgroup performance was not evaluated. Tissue-source-site sensitivity is not presented as a substitute for representative external subgroup evaluation."),
    ("Additional change: selection-conditioned uncertainty",
     "The Methods, tables and Discussion now use the explicit label 'selection-conditioned patient-resampling interval for repeated out-of-fold predictions' and enumerate included and excluded uncertainty components. These intervals condition on endpoint selection and five fixed nested-CV prediction sets; they do not remove winner's-curse optimism, repeat screening or refitting, or represent external performance."),
]
for title, answer in responses:
    resp.add_heading(title, 1); resp.add_paragraph(answer)
resp.save(OUT / "response_to_reviewer_JTM.docx")

for source_name, output_name in (
    ("COAD_example_A.pdf", "Additional_file_2_COAD_example_A_TITANPred_report.pdf"),
    ("COAD_example_B.pdf", "Additional_file_3_COAD_example_B_TITANPred_report.pdf"),
):
    shutil.copyfile(ROOT / "results" / "reports" / source_name, OUT / output_name)

print(OUT / "manuscript_JTM_patient_level_TITAN.docx")
print(OUT / "supplementary_material_JTM.docx")
print(OUT / "response_to_reviewer_JTM.docx")
