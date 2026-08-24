#!/usr/bin/env python3
"""Fail closed on structural requirements of the final JTM DOCX package."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript"
EXPECTED = {
    "manuscript_JTM_patient_level_TITAN.docx",
    "supplementary_material_JTM.docx",
    "response_to_reviewer_JTM.docx",
    "reviewer_report_JTM.docx",
}
AUTHORS = (
    "Aamilah Ismail",
    "Martin Ocharo",
    "Moussa Kassim",
    "Dalia Ahmed",
    "Brendon Price",
    "Dupe Ojo",
    "Ekene Emmanuel Nweke",
    "Silvano Piazza",
    "Dinesh Gupta",
    "Stefano Cacciatore",
)
REMOVED = ("Alessia Vignoli", "Leonardo Tenori", "CERM", "University of Florence")


def require(ok: bool, message: str) -> None:
    if not ok:
        raise RuntimeError(message)


def full_text(document: Document) -> str:
    chunks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


found = {p.name for p in OUT.glob("*.docx") if not p.name.startswith("~$")}
require(found == EXPECTED, f"Expected exactly four final DOCX files; found {sorted(found)}")

documents = {name: Document(OUT / name) for name in EXPECTED}
texts = {name: full_text(document) for name, document in documents.items()}
combined = "\n".join(texts.values())
for removed in REMOVED:
    require(removed not in combined, f"Removed author/affiliation remains in DOCX output: {removed}")

main_name = "manuscript_JTM_patient_level_TITAN.docx"
main = documents[main_name]
main_text = texts[main_name]
supplement_text = texts["supplementary_material_JTM.docx"]
for author in AUTHORS:
    require(author in main_text, f"Required author missing from manuscript: {author}")
author_positions = [main_text.index(author) for author in AUTHORS]
require(
    author_positions == sorted(author_positions),
    "Author order is inconsistent with the approved ten-author sequence",
)
require(
    "Aamilah Ismail3,*" in main_text and "Martin Ocharo1,2,*" in main_text,
    "Shared co-first authorship markers are missing for Aamilah Ismail and Martin Ocharo",
)
require(
    "Department of Surgery, Faculty of Health Sciences, University of the Witwatersrand, "
    "Johannesburg, Gauteng, South Africa" in main_text,
    "Ekene Emmanuel Nweke's affiliation is missing",
)
require(
    "Division of Anatomical Pathology, University of Cape Town and National Health "
    "Laboratory Service, Observatory, Cape Town, South Africa" in main_text,
    "Brendon Price's affiliation is missing",
)
require(
    "Computational Biology Group, International Centre for Genetic Engineering and "
    "Biotechnology (ICGEB), Trieste, Italy" in main_text,
    "Silvano Piazza's ICGEB affiliation is missing",
)
require(
    "Bioinformatics Facility, Department of Cellular, Computational and Integrative "
    "Biology - CIBIO, University of Trento, Trento, Italy" in main_text,
    "Silvano Piazza's CIBIO affiliation is missing",
)
require(
    "Translational Bioinformatics Group, International Centre for Genetic Engineering "
    "and Biotechnology (ICGEB), Aruna Asaf Ali Marg, New Delhi 110067, India" in main_text,
    "Dinesh Gupta's affiliation is missing",
)
require(
    "† Corresponding author: Stefano Cacciatore" in main_text,
    "Sole corresponding-author statement is missing",
)
require(
    "Silvano Piazza6,8" in main_text,
    "Silvano Piazza's dual-affiliation markers are missing",
)
for label in (
    "Additional file 2 (.pdf): COAD example A (TCGA-AA-A01F) TITANPred research-software output",
    "Additional file 3 (.pdf): COAD example B (TCGA-AA-3972) TITANPred research-software output",
):
    require(label in main_text, f"Separate supplementary report is not declared: {label}")
for pdf_name in (
    "Additional_file_2_COAD_example_A_TITANPred_report.pdf",
    "Additional_file_3_COAD_example_B_TITANPred_report.pdf",
):
    pdf_path = OUT / pdf_name
    require(pdf_path.is_file(), f"Separate supplementary PDF is missing: {pdf_name}")
    require(pdf_path.read_bytes().startswith(b"%PDF"), f"Invalid PDF signature: {pdf_name}")
for heading in (
    "Abstract",
    "Background",
    "Methods",
    "Results",
    "Discussion",
    "Conclusions",
    "Abbreviations",
    "Declarations",
    "References",
):
    require(heading in main_text, f"Required manuscript section missing: {heading}")

paragraphs = [p.text.strip() for p in main.paragraphs]
abstract_start = paragraphs.index("Abstract") + 1
keywords_index = next(
    i for i in range(abstract_start, len(paragraphs))
    if paragraphs[i].startswith("Keywords:")
)
abstract = " ".join(paragraphs[abstract_start:keywords_index])
abstract_words = re.findall(r"\b[\w²×–-]+\b", abstract, flags=re.UNICODE)
require(len(abstract_words) <= 350, f"Abstract has {len(abstract_words)} words (>350)")
keyword_count = len([x for x in paragraphs[keywords_index][9:].split(";") if x.strip()])
require(3 <= keyword_count <= 10, f"Keyword count is {keyword_count}; expected 3–10")
require(len(main.inline_shapes) >= 7, "Main manuscript contains fewer than seven embedded figures")
for forbidden in (
    "IRLBA",
    "Liver and pancreatic cancer examples",
    "Previously reported and atlas-nominated predictors",
    "Practical contribution of the PLS framework",
    "discovery atlas",
    "patient molecular profile",
    "prediction report",
    "fluorouracil",
    "oxaliplatin",
    "bevacizumab",
    "complete response",
    "progressive disease",
):
    require(forbidden.lower() not in main_text.lower(), f"Removed manuscript text remains: {forbidden}")
require("frozen" not in main_text.lower(), "Ambiguous 'frozen' terminology remains")
for obsolete_effect_label in ("higher effect", "moderate effect", "higher-effect", "moderate-effect"):
    require(
        obsolete_effect_label not in (main_text + "\n" + supplement_text).lower(),
        f"Obsolete effect-sounding screening label remains: {obsolete_effect_label}",
    )
for required_revision in (
    "A systematic patient-level benchmark and reusable model resource",
    "Table 1. Eligible cancer–endpoint models and within-cancer screening categories",
    "12,093 target-specific models for 4,031 multi-omic biomarkers",
    "fixed pretrained representation",
    "tested-negative and sample-size-ineligible results",
    "83/323 models (25.7%) fell below the original prespecified screening threshold",
    "Threshold-independent AUROC was the primary binary comparison metric",
    "the same inner-CV balanced-accuracy threshold rule applied to both methods",
    "targets were selected from all 2,073 eligible cancer–endpoint tests without using PLS performance",
    "The fixed rule yielded 12 continuous and 35 binary targets",
    "PLS is the prespecified reference method, not an empirically optimal algorithm",
    "Portability cannot justify PLS over ridge because both yield compact exportable parameters",
    "Its mean cancer-level ΔQ² relative to response-by-response PLS1 was 0.066",
    "every permutation, patient labels were reassigned and the complete modelling process was repeated",
    "single BH correction across all eligible continuous and binary cancer–endpoint tests",
    "zero exceedances among 999 permutations this interval is 0–0.003686",
    "prespecified high-resolution subset extended 8 leading models to 9,999",
    "figures and tables are ordered by the outcome-appropriate predictive metric, with repeated and site-grouped stability reported alongside; q-values are not used for ranking",
    "prespecified screening tier A",
    "prespecified screening tier B",
    "The tiers are prespecified prioritisation rules, not established clinical or statistical effect categories",
    "their numerical values are not interpreted as directly commensurate across outcome types",
    "The initial nested-CV estimate used for candidate screening, permutation testing and tier assignment is termed the primary screening estimate",
    "THYM–Th17 had primary screening Q²=0.638 and a five-repeat mean repeated-validation Q²=0.594",
    "This difference reflects independently seeded nested-CV partitions and is not a discrepancy",
    "38/41 screen-positive cancer–gene pairs",
    "TCGA out-of-fold score rank (not probability)",
    "Probability calibration was not added",
    "95% selection-conditioned patient-resampling interval for repeated out-of-fold predictions",
    "It includes neither the initial screen and endpoint highlighting nor new partition generation",
    "For repeat r, d_r=M_r(ridge)−M_r(PLS)",
    "generate new partitions or refit either algorithm",
    "A stricter ≥50-per-class sensitivity retained 87/104 binary candidates",
    "Binary class-size sensitivity and development stability",
    "17 screen-positive models with fewer than 50 patients in either class",
    "omitted from default inference and available only through explicit opt-in",
    "PR-AUC was non-interpolated average precision",
    "Precision–recall performance was reported for every screen-positive binary endpoint",
    "mutation or fusion endpoints with prevalence below 0.20",
    "This divergence illustrates why AUROC or balanced accuracy alone can overstate positive-class retrieval under low prevalence",
    "cohort-specific descriptive estimates, not calibrated probabilities",
    "Tissue-source-site sensitivity and submitting-site prediction",
    "the same separation was maintained in every inner component-selection split",
    "27/32 cancers",
    "TCGA tissue-source-site-grouped internal validation",
    "every performance estimate is an internally derived TCGA estimate",
    "Prospectively locked protocol for future independent evaluation",
    "Table 2. Prospectively locked subset for future independent evaluation",
    "CPTAC-UCEC is a plausible future cohort",
    "research-software demonstration",
    "The package provides a synthetic feature vector constructed from stored training-feature means",
    "Figure 7. Post hoc TITANPred research-software visualization for two COAD participants",
    "Nodal status was not matched",
    "treatment, response, recurrence, follow-up and survival are neither shown nor interpreted",
    "TCGA OOF score rank (not probability)",
    "currently maintained in a private access-controlled repository",
    "Endpoint provenance and assay equivalence",
    "2,073-row endpoint dictionary",
    "TIL Regional Fraction inferred from H&E",
    "same-histology-modality concordance task",
    "equivalent to flow cytometry, immunohistochemistry, a directly counted immune-cell assay",
    "Qualitative morphology context",
    "Figure S4 places all held-out patients behind the selected examples",
    "not patch-level relevance, causal attribution or blinded pathologist review",
):
    require(
        required_revision in main_text,
        f"A required audit-driven revision is missing: {required_revision}",
    )
require("Table 3." not in main_text, "Oversized Table 3 remains in the main manuscript")
require("Table 4." not in main_text, "Obsolete Table 4 numbering remains in the main manuscript")
require("Comparison with major pan-cancer histology–molecular prediction studies" not in main_text,
        "Review-style comparison table remains in the main manuscript")
require("Complete highlighted-model performance, uncertainty, class-size metrics and tissue-source-site-grouped estimates are reported in Supplementary Table S6a" in main_text,
        "Main text does not direct readers to the complete supplementary performance table")

supplement = documents["supplementary_material_JTM.docx"]
supplement_text = texts["supplementary_material_JTM.docx"]
response_text = texts["response_to_reviewer_JTM.docx"]
reviewer_text = texts["reviewer_report_JTM.docx"]
require("Table S6a. Highlighted-model performance and uncertainty" in supplement_text,
        "Complete highlighted-model performance table is missing from the supplement")
require("Manuscript length and table density" in response_text,
        "Response does not address manuscript length and table density")
require("Report precision–recall metrics" in response_text,
        "Response does not address precision-recall reporting")
require("Manuscript density" in reviewer_text,
        "Reviewer report does not acknowledge the streamlined table structure")
require("Precision–recall reporting" in reviewer_text,
        "Reviewer report does not acknowledge precision-recall reporting")
for item in (
    "Table S10a. Expanded literature audit",
    "Table S10b. Screen-positive models below the original prespecified screening threshold",
    "Table S10c. Tissue-source-site-grouped outer-fold composition",
    "Table S10d. Within-cancer prediction of TCGA tissue-source site",
    "Table S10e. Metadata-stratified representative PLS–ridge benchmark",
    "Table S10f. Permutation precision and atlas-wide multiplicity sensitivity",
    "Table S10g. Binary class-size sensitivity and development reliability",
    "Table S6c. Prospectively locked subset for future independent evaluation",
    "Table S12",
    "Figure S1",
    "Figure S2",
    "Figure S3",
    "Table S13. Endpoint provenance, derivation and assay equivalence",
    "Table S14. Qualitative high/low prediction anchors and within-cancer TITAN neighbours",
    "Figure S4",
    "Machine-readable additional files",
    "95% paired patient interval",
    "The interval does not repeat target sampling, generate partitions or refit models",
    "THYM–Th17, for example, had primary screening Q²=0.638 and five-repeat mean Q²=0.594",
):
    require(item in supplement_text, f"Supplementary item missing: {item}")
for item in (
    "7. Reported uncertainty is conditional and does not capture the entire modelling process",
    "does not regenerate partitions or refit models",
    "9. Derived immune phenotypes, endpoint provenance and morphological context",
    "Every one of the 2,073 eligible cancer–endpoint tests is classified",
    "qualitative neighbour retrieval—not patch relevance, causal morphology or blinded pathologist review",
    "10. The COAD single-sample demonstration risks overinterpretation",
    "remain in the main manuscript solely as a software-visualization example",
    "All treatment, response, recurrence, follow-up and survival narrative has been deleted",
    "rank_is_probability=FALSE",
    "Primary versus repeated estimates",
    "THYM–Th17 is stated explicitly as primary screening Q²=0.638 versus five-repeat mean Q²=0.594",
):
    require(item in response_text, f"Reviewer-response item missing: {item}")
for item in (
    "7. Selection-conditioned uncertainty",
    "2,000-replicate paired patient bootstrap",
    "9. Endpoint provenance and morphological interpretation",
    "complete 2,073-row endpoint dictionary",
    "No additional internal interpretability claim is requested",
    "10. COAD software-interface illustration",
    "remain in main Figure 7 solely to show the software output",
    "No further case-based clinical interpretation is justified",
    "THYM–Th17 is transparently labelled Q² 0.638 in the screen and 0.594 in repeated validation",
    "neutral labels prespecified screening tier A/B are appropriate",
):
    require(item in reviewer_text, f"Reviewer-report item missing: {item}")
require(len(supplement.inline_shapes) >= 4, "Supplement contains fewer than four figures")
require(
    "Panel C. Target-level secondary metric; binary balanced accuracy uses operating thresholds selected identically for both methods."
    in supplement_text,
    "Symmetric binary decision-rule panel is missing from Table S10e",
)

require(
    "6. The minimum binary class requirement is too permissive for headline and distributed models"
    in response_text
    and "87 (83.7%) met it" in response_text
    and "excluded from default TITANPred inference" in response_text
    and "fixed-outer-test learning curves" in response_text,
    "Response to reviewer does not fully address binary class-size reliability",
)
require(
    "5. The multiplicity framework is thoughtful but remains resolution-limited"
    in response_text
    and "Clopper–Pearson Monte Carlo bounds" in response_text
    and "9,999 complete-process permutations" in response_text
    and "single atlas-wide BH sensitivities" in response_text,
    "Response to reviewer does not fully address permutation resolution and multiplicity",
)

require(
    "4. The binary modelling and ridge comparison use asymmetric decision rules"
    in response_text
    and "Confirmed and corrected" in response_text
    and "35 binary comparisons" in response_text,
    "Response to reviewer does not fully address the binary decision-rule asymmetry",
)
require(
    "8. The justification for retaining PLS as the central model is currently insufficient"
    in response_text
    and "47-target representative benchmark" in response_text
    and "portability is explicitly stated to apply to ridge as well" in response_text
    and "0.066, 0.017 and 0.056" in response_text,
    "Response to reviewer does not fully address the rationale for retaining PLS",
)
require(
    "one BH correction across all" in reviewer_text
    and "9,999 full permutations" in reviewer_text
    and "No additional multiplicity correction is requested" in reviewer_text,
    "Reviewer report does not reflect the corrected multiplicity analysis",
)

require(
    "Metadata-stratified PLS–ridge benchmark" in reviewer_text
    and "47-target benchmark comprises 12 continuous and 35 binary endpoints" in reviewer_text
    and "These results do not establish PLS as the best model" in reviewer_text
    and "PLS2 remains a secondary, hypothesis-generating analysis" in reviewer_text,
    "Reviewer report does not reflect the representative algorithm benchmark",
)

require(
    "Binary class size and development stability" in reviewer_text
    and "17 are now designated exploratory/limited evidence" in reviewer_text
    and "excludes the 17 limited models from default inference" in reviewer_text,
    "Reviewer report does not reflect binary class-size reliability changes",
)

for name, text in texts.items():
    require("Error! Reference source not found" not in text, f"Broken Word field in {name}")

print(
    f"Document audit passed: four DOCX files; {len(AUTHORS)}-author manuscript; "
    f"abstract {len(abstract_words)} words; {len(main.inline_shapes)} main figures; "
    f"{len(supplement.inline_shapes)} supplementary figures"
)
