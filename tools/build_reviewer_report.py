#!/usr/bin/env python3
"""Build a data-aware, fresh JTM-style review of the revised manuscript."""
from __future__ import annotations

import csv
from datetime import date
from pathlib import Path
from statistics import median

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
OUTPUT = ROOT / "manuscript" / "reviewer_report_JTM.docx"
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
BLACK = RGBColor(0, 0, 0)


def read_rows(name):
    with (TABLES / name).open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def set_font(run, size=11, bold=False, italic=False, color=BLACK):
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(paragraph, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)


def add_body(doc, text, lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if lead and text.startswith(lead):
        set_font(p.add_run(lead), bold=True)
        set_font(p.add_run(text[len(lead):]))
    else:
        set_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


continuous = read_rows("continuous_screen.csv")
binary = read_rows("binary_screen.csv")
highlighted = read_rows("highlighted_model_performance.csv")
mutation_coverage = read_rows("mutation_coverage_audit.csv")
slide_audit = read_rows("patient_slide_multiplicity_by_cancer.csv")
participant_characteristics = read_rows("participant_characteristics_by_cancer.csv")
ridge_summary = read_rows("pls_vs_ridge_representative_summary.csv")
ridge_comparison = read_rows("pls_vs_ridge_representative_models.csv")
ridge_jobs = read_rows("pls_vs_ridge_representative_jobs.csv")
multiplicity_summary = read_rows("multiplicity_sensitivity_summary.csv")
permutation_uncertainty = read_rows("permutation_monte_carlo_uncertainty.csv")
targeted_permutation = read_rows("targeted_permutation_refinement.csv")
binary_class_sensitivity = read_rows("binary_minimum_class_sensitivity.csv")
binary_reliability = read_rows("binary_class_reliability_summary.csv")
binary_learning = read_rows("binary_limited_evidence_learning_curve_summary.csv")
endpoint_dictionary = read_rows("endpoint_dictionary.csv")
endpoint_definitions = read_rows("endpoint_definition_dictionary.csv")
morphology_context = read_rows("morphology_context_examples.csv")
combined_multiplicity = next(
    r for r in multiplicity_summary if r["outcome_type"] == "combined"
)
zero_999 = [
    r for r in permutation_uncertainty
    if r["permutations"] == "999"
    and r["permutation_exceedances"] == "0"
    and r["permutation_stopped_early"] == "FALSE"
]
zero_999_upper = max(float(r["p_mc_upper_95"]) for r in zero_999)
targeted_zero = [r for r in targeted_permutation if r["zero_exceedances"] == "TRUE"]
targeted_p = [float(r["refined_p_9999"]) for r in targeted_permutation]
targeted_p_summary = (
    f"all {targeted_p[0]:.4f}"
    if targeted_p and max(targeted_p) - min(targeted_p) < 1e-12
    else f"{min(targeted_p):.4f}–{max(targeted_p):.4f}"
)
binary_ridge_summary = next(
    r for r in ridge_summary if r["outcome_type"] == "binary"
)
continuous_ridge_summary = next(
    r for r in ridge_summary if r["outcome_type"] == "continuous"
)
binary_primary_ridge = [
    r for r in ridge_comparison
    if r["outcome_type"] == "binary" and float(r["delta_ci_low"]) > 0
]
binary_primary_uncertain = [
    r for r in ridge_comparison
    if r["outcome_type"] == "binary"
    and not (float(r["delta_ci_low"]) > 0 or float(r["delta_ci_high"]) < 0)
]
binary_secondary_ridge = [
    r for r in ridge_comparison
    if r["outcome_type"] == "binary"
    and float(r["delta_secondary_ci_low"]) > 0
]
binary_secondary_pls = [
    r for r in ridge_comparison
    if r["outcome_type"] == "binary"
    and float(r["delta_secondary_ci_high"]) < 0
]
binary_secondary_uncertain = [
    r for r in ridge_comparison
    if r["outcome_type"] == "binary"
    and not (
        float(r["delta_secondary_ci_low"]) > 0
        or float(r["delta_secondary_ci_high"]) < 0
    )
]
continuous_secondary_ridge = [
    r for r in ridge_comparison
    if r["outcome_type"] == "continuous"
    and float(r["delta_secondary_ci_low"]) > 0
]
continuous_secondary_pls = [
    r for r in ridge_comparison
    if r["outcome_type"] == "continuous"
    and float(r["delta_secondary_ci_high"]) < 0
]
continuous_secondary_uncertain = [
    r for r in ridge_comparison
    if r["outcome_type"] == "continuous"
    and not (
        float(r["delta_secondary_ci_low"]) > 0
        or float(r["delta_secondary_ci_high"]) < 0
    )
]
ridge_screen_positive = [r for r in ridge_jobs if r["screen_tier"] in ("A", "B")]
participant_overall = next(
    r for r in participant_characteristics if r["tumor_type"] == "Overall"
)

screen_c = [r for r in continuous if r["tier"] in ("A", "B")]
screen_b = [r for r in binary if r["tier"] in ("A", "B")]
binary_standard = [
    r for r in binary_reliability
    if r["model_evidence_tier"] == "standard_internal_evidence"
]
binary_limited = [
    r for r in binary_reliability
    if r["model_evidence_tier"] == "exploratory_limited_evidence"
]
binary_sensitivity_50 = next(
    r for r in binary_class_sensitivity if r["minimum_per_class"] == "50"
)
binary_learning_100 = [r for r in binary_learning if float(r["training_fraction"]) == 1.0]
global_c = [r for r in screen_c if float(r["q_value_global"]) < 0.05]
global_mut = [
    r for r in screen_b
    if r["family"] == "driver_mutation" and float(r["q_value_global"]) < 0.05
]
n_mc3_profiled = sum(int(r["matched_profiled_patients"]) for r in mutation_coverage)
n_mc3_missing = sum(int(r["embedding_patients_without_mc3_profile"]) for r in mutation_coverage)
n_multi = sum(int(r["multi_slide_patients"]) for r in slide_audit)

doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.9)
section.left_margin = section.right_margin = Inches(1.0)
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = BLUE
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
set_font(header.add_run("Independent internal review | Journal of Translational Medicine"),
         size=9, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("Page "), size=9, color=MUTED)
add_page_field(footer)

p = doc.add_paragraph()
set_font(p.add_run("REVIEWER REPORT"), size=23, bold=True, color=BLUE)
p = doc.add_paragraph()
set_font(
    p.add_run(
        "A systematic patient-level benchmark and reusable model resource for "
        "molecular and derived immune-feature prediction from pretrained TITAN whole-slide "
        "representations across 32 TCGA cancers"
    ), size=14, color=MUTED
)
for label, value in (
    ("Target journal", "Journal of Translational Medicine — Molecular Pathology"),
    ("Article type", "Research article"),
    ("Review basis", "Fresh review of the revised manuscript, supplement and reproducibility materials"),
    ("Recommendation", "Major revision; consider only under an explicit computational-resource framing unless external validation is added"),
    ("Date", date.today().strftime("%d %B %Y")),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(label + ": "), bold=True)
    set_font(p.add_run(value))

callout = doc.add_paragraph()
callout.paragraph_format.space_before = Pt(12)
callout.paragraph_format.space_after = Pt(10)
callout.paragraph_format.left_indent = Inches(0.12)
callout.paragraph_format.right_indent = Inches(0.12)
shade(callout, "EEF4F8")
set_font(callout.add_run("Overall assessment. "), bold=True, color=BLUE)
set_font(callout.add_run(
    "This is a carefully reconstructed, patient-level TCGA benchmark and research-software resource. "
    "The revised work addresses the major scientific concerns that would "
    "otherwise prevent interpretation: multiple slides are aggregated before "
    "validation; molecular missingness is not labelled wild type; binary "
    "models use PLS–LDA; continuous and binary performance is reported with "
    "uncertainty; multiplicity is shown locally, across cancers and atlas-wide; "
    "and the absence of independent external validation is explicit. A metadata-stratified, "
    "performance-independent comparison now shows that ridge often matches or exceeds PLS, so the "
    "manuscript retains PLS only as the prespecified reference analysis and makes no optimality claim. The revised wording no longer "
    "presents internal TCGA estimates as translational evidence. Nevertheless, independent validation "
    "remains the decisive missing element for publication as a translational prediction study."
))

add_heading(doc, "Confidential comments to the editor", 1)
add_body(doc,
    "The manuscript fits the Molecular Pathology remit as a reproducible study "
    "of image-derived molecular and derived immune-feature signals in human tumours. Its main "
    "contribution is a transparent endpoint-by-endpoint benchmark using a fixed "
    "pretrained representation, complete negative reporting and portable fitted linear models. "
    "The authors appropriately refrain from clinical claims and label every value as an internally "
    "derived TCGA estimate. Without an independent cohort, the work should be evaluated as a "
    "computational pathology resource rather than a translational prediction study."
)

add_heading(doc, "Comments to the authors", 1)

add_heading(doc, "1. Scientific design and analysis", 2)
add_body(doc,
    f"The patient is now the analysis unit and {n_multi:,} patients with multiple "
    "eligible slides are pooled before outcome matching and fold assignment. "
    "The first-slide sensitivity is an appropriate patient-aggregation check. "
    "Tissue-source-site sensitivity is now treated separately as a central finding. Mutation denominators are restricted to "
    f"{n_mc3_profiled:,} MC3-profiled patients, while {n_mc3_missing:,} patients "
    "without a matched profile remain missing; the nine retained protein-altering "
    "variant classes are now explicit and audited. The source-level aliquot and "
    "fusion-coverage audits make the outcome construction reproducible. "
    f"The TCGA Clinical Data Resource now supplies descriptive characteristics "
    f"for {int(participant_overall['cdr_matched']):,} matched participants, "
    "including age, recorded gender, race and broad stage overall and by cancer."
)
add_body(doc,
    "No additional analysis requested. The manuscript should retain the exact "
    "wild-type, fusion-negative and patient-aggregation definitions during copyediting, "
    "and should not imply that descriptive demographics constitute subgroup model validation.",
    lead="No additional analysis requested."
)

add_heading(doc, "2. Multiplicity and interpretation", 2)
add_body(doc,
    f"The final screen reports {len(screen_c):,} continuous and {len(screen_b):,} "
    "binary within-cancer candidates. The stricter across-cancer family correction "
    f"retains {len(global_c):,} continuous and {len(global_mut):,} mutation pairs, while "
    f"{int(combined_multiplicity['atlas_wide_pass']):,}/{int(combined_multiplicity['within_cancer_family_candidates']):,} "
    f"local candidates pass one BH correction across all {int(combined_multiplicity['eligible_tests']):,} eligible atlas tests. "
    "The manuscript correctly distinguishes this single atlas-wide sensitivity from its primary local cancer-specific questions. "
    "It also now states explicitly that every permutation reruns training-fold centering, inner component selection, outer refitting and held-out prediction. "
    f"Among completed 999-permutation tests, {len(zero_999)} had zero exceedances; the exact two-sided 95% Monte Carlo interval is 0–{zero_999_upper:.6f}. "
    f"The locked eight-model refinement continues the same streams to 9,999 full permutations, with {len(targeted_zero)}/{len(targeted_permutation)} zero-exceedance results and refined p-values {targeted_p_summary}. "
    "The refinement is appropriately presented as a precision sensitivity rather than being inserted post hoc into the prespecified FDR screen. "
    "Models are ordered by the outcome-appropriate predictive metric, repeated and site-grouped stability are reported alongside, and tied minimum q-values are not used for ranking. "
    "The former effect-sounding categories have been replaced by the neutral labels 'prespecified screening tier A' and 'prespecified screening tier B'. The Methods state that these are prioritisation rules based on Q² for continuous outcomes and 2×balanced accuracy−1 for binary outcomes, not established clinical or statistical effect categories or directly commensurate scales."
)
add_body(doc,
    "No additional multiplicity correction is requested. Preserve the numerical distinction between "
    "within-cancer, across-cancer-family and atlas-wide results, and retain the Monte Carlo interval and targeted-refinement caveats.",
    lead="No additional multiplicity correction is requested."
)

add_heading(doc, "3. Performance reporting and research-use provenance", 2)
add_body(doc,
    f"The {len(highlighted):,} highlighted models report sensitivity, specificity, "
    "balanced accuracy, AUROC, PR-AUC and cohort-specific PPV/NPV for binary outcomes and Q², RMSE and Spearman "
    "correlation for continuous outcomes, with 95% selection-conditioned patient-resampling intervals for repeated out-of-fold predictions. "
    "The Methods and Discussion now state that these intervals resample patients from five fixed prediction sets and do not repeat screening, fold generation, tuning or fitting. They do not remove winner's-curse optimism or estimate external performance. "
    "The observed-versus-predicted figure is useful. Model artifacts record ordered "
    "feature schema, checksums, training ranges, pooling, endpoint transformations, "
    "output units, class coding, priors, decision behaviour, calibration status "
    "and external-validation status. The registry also reports the exact fastPLS "
    "commit, computation backend and exclusive seeded rSVD configuration. Removal of training latent-score and "
    "fitted-value rows strengthens the portability claim."
)
add_body(doc,
    "No additional analysis requested. Continue to describe LDA scores as uncalibrated "
    "and all fitted objects as research-only TCGA models; PPV and NPV must remain explicitly tied to observed TCGA prevalence.",
    lead="No additional analysis requested."
)

add_heading(doc, "4. Tissue-source-site sensitivity and confounding", 2)
add_body(doc,
    "The revised manuscript now makes the site result visible rather than reducing it to a median robustness statistic. "
    "Main Figure 6 shows every target, the largest attenuations and the new within-cancer submitting-site classifiers; "
    "Supplementary Table S6a reports grouped performance beside highlighted models. Overall, 83/323 screen-positive models (25.7%) "
    "fell below their original prespecified screening threshold. READ–APC and COAD–APC declined to balanced accuracies 0.495 and 0.543. "
    "The public analysis registry and inference output expose the grouped metric, delta, site count, threshold-retention "
    "status and prominent warning for every model. The complete fold audit reports patients, sites and class counts in "
    "all 1,613 outer folds and confirms that inner component selection also kept tissue-source sites intact."
)
add_body(doc,
    "The dedicated repeated nested analysis predicted tissue-source site from TITAN representations in 27/32 evaluable "
    "cancers, with median multiclass macro balanced accuracy 0.628. This quantifies substantial confounding potential. "
    "The manuscript appropriately calls the endpoint analysis 'TCGA tissue-source-site-grouped internal validation' and "
    "does not interpret retained performance as proof of scanner, laboratory or institutional transportability. Tissue-source "
    "site remains an imperfect proxy, and smaller grouped folds can account for part—but not all—of the attenuation."
)

add_heading(doc, "5. Metadata-stratified PLS–ridge benchmark", 2)
add_body(doc,
    f"The revised comparison is no longer conditional on successful PLS screening. All 2,073 eligible tests entered "
    f"a metadata-only sampling frame, and a fixed salted SHA-256 rank selected one target from each non-empty outcome-family "
    f"× sample-size-tercile cell, with an additional minority-class-fraction tercile for binary outcomes. The resulting "
    f"47-target benchmark comprises 12 continuous and 35 binary endpoints; only {len(ridge_screen_positive)} were PLS "
    "screen-positive and 33 were screen-negative. This is a representative benchmark, not an atlas-wide ridge screen. "
    "Both methods use the same patients, five repeated outer partitions and inner partitions. Continuous PLS components "
    "and ridge penalties are each selected by maximising pooled inner out-of-fold Q². Binary PLS–LDA components and ridge "
    "penalties are tuned within the same inner folds, and both operating thresholds are selected from method-specific inner "
    "out-of-fold scores by maximising balanced accuracy. Threshold-independent AUROC is the primary binary comparison "
    f"metric: the median ridge-minus-PLS difference is {float(binary_ridge_summary['median_delta_ridge_minus_pls']):.3f} "
    f"(IQR {float(binary_ridge_summary['q1_delta']):.3f} to {float(binary_ridge_summary['q3_delta']):.3f}); "
    f"{int(binary_ridge_summary['ridge_better'])} endpoints favour ridge, {int(binary_ridge_summary['pls_better'])} favour PLS and "
    f"{int(binary_ridge_summary['difference_uncertain'])} are uncertain. Under symmetric thresholding, the median "
    f"balanced-accuracy difference is {float(binary_ridge_summary['median_secondary_delta_ridge_minus_pls']):.3f} "
    f"(IQR {float(binary_ridge_summary['q1_secondary_delta']):.3f} to {float(binary_ridge_summary['q3_secondary_delta']):.3f}); "
    f"{len(binary_secondary_ridge)} favour ridge, {len(binary_secondary_pls)} favour PLS and {len(binary_secondary_uncertain)} are uncertain."
)
add_body(doc,
    f"For continuous outcomes, the median ridge-minus-PLS Q² difference is "
    f"{float(continuous_ridge_summary['median_delta_ridge_minus_pls']):.3f} "
    f"(IQR {float(continuous_ridge_summary['q1_delta']):.3f} to {float(continuous_ridge_summary['q3_delta']):.3f}); "
    f"{int(continuous_ridge_summary['ridge_better'])}/12 favour ridge, none favour PLS and "
    f"{int(continuous_ridge_summary['difference_uncertain'])} are uncertain. The secondary Spearman difference has median "
    f"{float(continuous_ridge_summary['median_secondary_delta_ridge_minus_pls']):.3f}; "
    f"{len(continuous_secondary_ridge)} favour ridge, {len(continuous_secondary_pls)} favour PLS and "
    f"{len(continuous_secondary_uncertain)} are uncertain. The 2,000-replicate paired patient-resampling intervals are "
    "conditional on the fixed representative sample, folds and predictions and do not refit either model. These results do "
    "not establish PLS as the best model. The primary atlas appropriately retains PLS only as its prespecified reference "
    "analysis; the released PLS models should not be described as algorithmically optimal."
)

add_heading(doc, "6. Binary class size and development stability", 2)
add_body(doc,
    f"The authors have now separated screening inclusiveness from model maturity. Of 459 binary pairs eligible under "
    f"the prespecified 20-per-class rule, {int(binary_sensitivity_50['eligible_binary_targets'])} "
    f"({float(binary_sensitivity_50['eligible_target_retention_percent']):.1f}%) meet a stricter 50-per-class rule. "
    f"Among 104 screen-positive binary models, {len(binary_standard)} ({100 * len(binary_standard) / len(screen_b):.1f}%) meet "
    f"the stricter standard and {len(binary_limited)} are now designated exploratory/limited evidence. None of the main "
    "highlighted binary models is limited evidence. The package retains the complete atlas for transparency but excludes "
    "the 17 limited models from default inference; explicit opt-in emits a warning."
)
add_body(doc,
    "The reliability supplement now provides class counts for every repeated outer fold, exact reconstructed inner-fold "
    "class minima, selected-component distributions across 25 fits per model, average precision, TCGA-prevalence PPV/NPV, "
    "repeat score and call stability, and fixed-test-fold learning curves at 50%, 75% and 100% of the outer-training data. "
    f"The limited models had as few as 4 positive cases in an outer test fold, 12 positive cases in an inner training fold "
    f"and 3 in an inner validation fold. At full training size, their median balanced accuracy was "
    f"{median(float(r['balanced_accuracy_mean']) for r in binary_learning_100):.3f}, median AUROC "
    f"{median(float(r['auc_mean']) for r in binary_learning_100):.3f} and median PR-AUC "
    f"{median(float(r['pr_auc_mean']) for r in binary_learning_100):.3f}. Component selection frequently reached the "
    "prespecified ceiling, supporting the decision not to present these models beside substantially better-supported ones."
)
add_body(doc,
    "No further internal sample-size analysis is requested. Retain these models only in the clearly separated exploratory "
    "tier, keep them out of default inference, and do not interpret learning curves or repeat stability as substitutes for "
    "independent validation.",
    lead="No further internal sample-size analysis is requested."
)

add_heading(doc, "7. Selection-conditioned uncertainty", 2)
add_body(doc,
    "The revised manuscript no longer presents generic '95% bootstrap intervals' or '95% confidence intervals' for the highlighted TCGA results. The abstract and table captions define these as 95% selection-conditioned patient-resampling intervals for repeated out-of-fold predictions. Each replicate samples patients as clusters, retains all five existing held-out predictions, recalculates the metric within repeat and averages across repeats. The Methods now distinguish represented patient-sampling variability from excluded screening/highlighting uncertainty, new-fold variability, retuning/refitting variability, winner's-curse correction and external cohort/site/scanner/population variation."
)
add_body(doc,
    f"The PLS–ridge comparison has also been strengthened. For repeat r, the paired difference is d_r=M_r(ridge)−M_r(PLS), with Δ=(1/5)Σ_r d_r. A 2,000-replicate paired patient bootstrap samples patients with replacement while retaining both methods and all five matched repeat predictions; percentile limits are calculated from the resulting Δ values. The matched predictions and all {len(ridge_comparison) * 5} repeat-specific differences are machine-readable. The interval conditions on the metadata-stratified representative sample and does not regenerate partitions or refit either algorithm inside a bootstrap replicate, which is now stated explicitly."
)
add_body(doc,
    "No additional internal uncertainty analysis is requested. Preserve the full interval label in the abstract and define the SC abbreviation in every applicable table; do not relabel these intervals as generalisation confidence intervals during copyediting.",
    lead="No additional internal uncertainty analysis is requested."
)

add_heading(doc, "8. Relation to prior work and translational positioning", 2)
add_body(doc,
    "The expanded review now covers mutation, MSI, gene expression, continuous "
    "biomarkers, fusions, homologous-recombination deficiency and tumour-microenvironment "
    "prediction, and it distinguishes internal current metrics from prior external "
    "AUROCs and correlations. The manuscript correctly documents that TCGA was "
    "excluded from TITAN's Mass-340K pretraining but used in TITAN's downstream "
    "evaluation. It is therefore accurate to call this a secondary systematic TCGA "
    "benchmark rather than independent validation."
)
add_body(doc,
    "The authors now document that CPTAC-UCEC is a feasible future cohort but that its "
    "whole-slide archive requires de novo TITAN extraction. They lock UCEC TP53 mutation, "
    "genome doubling and continuous aneuploidy score before inspecting external data, with "
    "artifact hashes, no-refitting rules and mandatory reporting of failures. This is useful "
    "prospective discipline, but it is not external validation evidence."
)
add_body(doc,
    "PLS2 remains a secondary, hypothesis-generating analysis because it answers a different multiresponse question. "
    "Within complete-case, ordered inflammatory panels it improved mean cancer-level Q² by 0.066 for inferred cell "
    "fractions, 0.017 for infiltration scores and 0.056 for immune-repertoire features. However, the shared latent "
    "solution changes the estimand and eligible complete-case population, has no binary analogue in this resource, and "
    "has not been compared with a matched multiresponse ridge baseline or independently validated. It should therefore "
    "not replace the endpoint-wise primary benchmark post hoc. A future locked multiresponse study could evaluate PLS2 "
    "and alternative multitask models on identical cohorts and external data. Portability is not unique to PLS; both PLS "
    "and ridge yield compact exportable linear parameters."
)

add_heading(doc, "9. Endpoint provenance and morphological interpretation", 2)
add_body(doc,
    f"The revision now distinguishes the label-generating assays and algorithms rather than presenting all continuous targets as equivalent immune measurements. The complete {len(endpoint_dictionary):,}-row endpoint dictionary covers every eligible cancer–endpoint test and records source modality, direct/inferred status, derivation algorithm, original scale, analysis transformation, target-specific missingness, expected measurement error, biological interpretation and an assay-equivalence caveat. Its {len(endpoint_definitions)} unique definitions are also supplied separately. The high-level classes are direct sequence-supported alterations, sequencing-derived burdens, inferred immune-cell fractions, transcriptomic signatures, pathology-associated quantities and composite genomic-context scores. CIBERSORT fractions are correctly described as bulk-RNA deconvolution outputs, leukocyte fraction as methylation-derived and TIL Regional Fraction as an H&E-derived Saltz quantity. The text appropriately states that prediction of these inferred phenotypes is not equivalent to a direct cell count, IHC, flow cytometry or a clinically certified biomarker."
)
add_body(doc,
    f"The authors also add a qualitative context analysis with {len(morphology_context)} records across five representative models. High/low repeated-OOF anchors are paired with their nearest within-cancer patient-level mean TITAN neighbour, and the table preserves exact patient/slide identifiers, original predictions, ranks, cosine similarities and generated-report context. This is a useful transparency step, especially because the BLCA TIL Regional Fraction panel is explicitly identified as same-H&E-modality concordance. However, global pooled embeddings do not localise relevant patches, and TCGA-Slide-Reports text generated by TITAN from the same images is not independent evidence or blinded pathology review. The manuscript now states those limits plainly."
)
add_body(doc,
    "No additional internal interpretability claim is requested. Retain Figure S4 and Tables S13–S14 as qualitative context, not mechanistic evidence. Any future morphological explanation should preserve tile-level representations or relevance maps and undergo blinded pathologist review.",
    lead="No additional internal interpretability claim is requested."
)

add_heading(doc, "10. COAD software-interface illustration", 2)
add_body(doc,
    "At the authors' request, the two COAD participants and radar comparison remain in main Figure 7 solely to show the software output. The Results and caption state the full post hoc construction: clinical-text and profile-saturation restrictions followed by selection of the pair with maximum Euclidean separation across continuous TCGA OOF prediction percentiles. This makes clear that the visual contrast was designed after inspecting model outputs and is not representative sampling or performance evidence."
)
add_body(doc,
    "The case language is now appropriately limited. TCGA-AA-A01F is reported as pN1, whereas the available slide summary for TCGA-AA-3972 contains no nodal category; no claim of clinical comparability remains. Treatment, response, recurrence, follow-up and survival narratives have been removed because these variables were neither model inputs nor validation outcomes. The package tutorial separately uses a synthetic non-patient vector for routine smoke testing."
)
add_body(doc,
    "The package report makes the principal limitations visually prominent: its banner states 'UNVALIDATED RESEARCH OUTPUT', 'TCGA OOF score rank (not probability)', 'No independent external validation' and 'Not for diagnosis, treatment selection, or clinical risk estimation'. Binary plots and tables place '(not probability)' directly beside every score-rank label, and machine-readable inference rows expose an explicit false probability flag and uncalibrated status. No further case-based clinical interpretation is justified.",
    lead="No further case-based clinical interpretation is justified."
)

add_heading(doc, "11. Reproducibility and model redistribution", 2)
add_body(doc,
    "The reproducibility resource is split between the public analysis repository and a "
    "separate GPL-3 TITANPred R package. The package contains all 323 fitted research models, with 306 in default inference "
    "and 17 limited-evidence binary models requiring explicit opt-in, "
    "their SHA-256 registry, repeated out-of-fold reference distributions, cancer-vector "
    "inference interface and HTML/PDF research-software template. The package repository is "
    "currently private, and the manuscript now states this rather than claiming public release. The "
    "artifacts contain learned parameters and training-range summaries but no patient-level "
    "training rows. This materially strengthens the portability and transparency claim."
)
add_body(doc,
    "Required before submission: create a versioned release and persistent DOI for the public "
    "analysis repository. If fitted-model redistribution is authorised, document the model-package "
    "access terms and archive it separately; otherwise retain the private/access-controlled status "
    "throughout the manuscript. Keep the research-only and no-external-validation notices and cite "
    "the upstream TITAN work and feature source in the package documentation.",
    lead="Required before submission:"
)

add_heading(doc, "12. Administrative completion", 2)
add_body(doc,
    "Required before submission: replace the remaining placeholders for funding, "
    "competing interests and author contributions; confirm the institutional ethics/waiver "
    "wording. Aamilah Ismail and Martin Ocharo are now the shared co-first authors, "
    "with Martin Ocharo second in the author order and assigned to affiliations 1 and 2. "
    "Brendon Price is included in the middle of the author list with the Division of "
    "Anatomical Pathology, University of Cape Town and National Health Laboratory Service "
    "affiliation. Ekene Emmanuel Nweke's University of the Witwatersrand affiliation is "
    "included. Silvano Piazza and Dinesh Gupta are listed immediately before Stefano "
    "Cacciatore. Silvano Piazza has dual affiliations with the ICGEB Computational Biology "
    "Group in Trieste and the Bioinformatics Facility, CIBIO, University of Trento; Dinesh "
    "Gupta has the ICGEB New Delhi affiliation. "
    "Email addresses for Martin Ocharo, Brendon Price, Ekene Emmanuel Nweke, Silvano "
    "Piazza and Dinesh Gupta were not supplied. The current author list contains Aamilah "
    "Ismail, Martin Ocharo, Moussa Kassim, Dalia Ahmed, Brendon Price, Dupe Ojo, Ekene "
    "Emmanuel Nweke, Silvano Piazza, Dinesh Gupta and Stefano Cacciatore."
)

add_heading(doc, "Minor presentation points", 1)
for title, body in (
    ("Reporting checklist", "The TRIPOD+AI item-to-location map in Supplementary Table S12 and the endpoint dictionary in Table S13 should accompany the submission."),
    ("Fusion caveat", "Keep the technical-coverage caveat for fusion-negative status and fusion burden."),
    ("Figures", "Figure 1 is no longer clipped, and Figure 4A now shows the strongest primary-screen continuous prediction with explicit held-out metrics. These corrected production figures should be retained."),
    ("Primary versus repeated estimates", "The revision now distinguishes the initial nested-CV primary screening estimate from the mean across five additional nested-CV repeats. Supplementary Table S6a reports both; THYM–Th17 is transparently labelled Q² 0.638 in the screen and 0.594 in repeated validation. The difference is expected resampling variation, not an inconsistency."),
    ("Manuscript density", "The main paper is appropriately streamlined as a research article: the review-style literature comparison table has been removed, and the oversized highlighted-model performance table has moved to Supplementary Table S6a without loss of metrics. The main manuscript now contains only two compact tables."),
    ("Precision–recall reporting", "PR-AUC is now visible in the Abstract and main Results, not only in supplementary reliability fields. It is correctly defined as non-interpolated average precision and interpreted against endpoint prevalence. The explicit low-prevalence mutation/fusion summary usefully shows that apparently strong AUROC can coexist with substantially more modest positive-class retrieval."),
    ("Screening tiers", "The neutral labels prespecified screening tier A/B are appropriate. They are clearly defined as prioritisation rules rather than established clinical or statistical effect categories."),
):
    add_heading(doc, title, 3)
    add_body(doc, body)

add_heading(doc, "Strengths", 1)
for body in (
    "Patient-first aggregation and validation prevent multiple-slide leakage.",
    "Wild type, molecular missingness, fusion coverage and aliquot aggregation are auditable.",
    "The benchmark distinguishes inferred immune and inflammatory features from direct alterations and other derived genomic phenotypes.",
    "Negative and ineligible endpoints are preserved rather than selectively omitted.",
    "PLS1–PLS2 comparison is appropriately secondary because it changes the estimand and complete-case population; PLS is presented as the prespecified reference rather than the best method.",
    "The literature discussion reports that 38 of 41 screen-positive mutation pairs had prior support and avoids broad endpoint-novelty claims.",
):
    add_body(doc, "• " + body)

add_heading(doc, "Recommended editorial decision", 1)
decision = doc.add_paragraph()
shade(decision, "E8EEF5")
set_font(decision.add_run("MAJOR REVISION / RESOURCE FRAMING"), size=12, bold=True, color=BLUE)
add_body(doc,
    "The scientific analysis may be considered as a TCGA patient-level computational pathology benchmark and research-software resource. "
    "It is not yet supported as a translational prediction study because no independent cohort has been evaluated. "
    "The revised framing, locked future protocol and complete internal reporting are appropriate, but they do not replace external evidence. "
    "Declarations and author metadata also remain to be completed."
)

add_heading(doc, "Review standards consulted", 1)
add_body(doc,
    "Warren S. Molecular Pathology. J Transl Med. 2024;22:91. "
    "https://doi.org/10.1186/s12967-024-04868-7"
)
add_body(doc,
    "Collins GS, et al. TRIPOD+AI statement. BMJ. 2024;385:e078378. "
    "https://doi.org/10.1136/bmj-2023-078378"
)

doc.core_properties.title = "Fresh reviewer report — TITAN TCGA benchmark and model resource"
doc.core_properties.subject = "Journal of Translational Medicine-style internal peer review"
doc.core_properties.author = "Internal scientific review"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
